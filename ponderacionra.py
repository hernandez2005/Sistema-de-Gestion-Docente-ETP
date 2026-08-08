import streamlit as st
import google.generativeai as genai
from openai import OpenAI, RateLimitError as OpenAIRateLimitError
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from io import BytesIO
import PyPDF2
import json
import re
import unicodedata
from google.api_core.exceptions import ResourceExhausted
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type

# --- ESTILOS CSS ---
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    html, body, [class*="css"] { font-family: 'Inter', sans-serif; background-color: #F8FAFC; color: #1E293B; }
    .main-header { font-size: 2.2rem; font-weight: 700; color: #0F172A; text-align: center; margin-bottom: 5px; line-height: 1.2; }
    .sub-header { text-align: center; color: #475569; font-size: 1.1rem; font-weight: 400; margin-bottom: 35px; }
    [data-testid="stForm"] { background-color: #FFFFFF; border: 1px solid #E2E8F0; border-radius: 8px; box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05); padding: 30px; }
    .section-title { color: #1D4ED8; font-weight: 600; font-size: 1.2rem; border-bottom: 2px solid #DBEAFE; padding-bottom: 8px; margin-top: 20px; margin-bottom: 18px; }
    div.stButton > button:first-child, div.stFormSubmitButton > button:first-child { background-color: #2563EB !important; color: #FFFFFF !important; border: none !important; border-radius: 6px !important; font-weight: 600 !important; padding: 10px 24px !important; width: 100%; }
    div.stButton > button:first-child:hover, div.stFormSubmitButton > button:first-child:hover { background-color: #1D4ED8 !important; }
    .ra-card { background: #F0F9FF; border-left: 4px solid #0EA5E9; padding: 12px 16px; border-radius: 6px; margin-bottom: 10px; }
    .ra-id { font-weight: 700; color: #0C4A6E; font-size: 0.95rem; }
    .ra-text { color: #334155; font-size: 0.88rem; margin-top: 4px; }
</style>
""", unsafe_allow_html=True)

# ===========================================================================
# SCHEMAS JSON
# ===========================================================================
SCHEMA_EXTRACCION = {
    "type": "OBJECT",
    "properties": {
        "MODULO_ENCONTRADO": {"type": "BOOLEAN"},
        "MOTIVO_NO_ENCONTRADO": {"type": "STRING"},
        "NOMBRE_MODULO_HALLADO": {"type": "STRING"},
        "CODIGO_MODULO": {"type": "STRING"},
        "ESQUEMA_DETECTADO": {"type": "STRING"},
        "CANTIDAD_RA": {"type": "INTEGER"},
        "RESULTADOS_APRENDIZAJE": {
            "type": "ARRAY",
            "items": {
                "type": "OBJECT",
                "properties": {
                    "ID": {"type": "STRING"},
                    "TEXTO": {"type": "STRING"},
                    "CR": {"type": "STRING"},
                    "CE": {"type": "STRING"},
                    "CONTENIDOS": {"type": "STRING"}
                },
                "required": ["ID", "TEXTO", "CR", "CE", "CONTENIDOS"]
            }
        }
    },
    "required": ["MODULO_ENCONTRADO", "NOMBRE_MODULO_HALLADO", "CANTIDAD_RA", "RESULTADOS_APRENDIZAJE"]
}

SCHEMA_PONDERACION = {
    "type": "OBJECT",
    "properties": {
        "RESUMEN": {"type": "STRING"},
        "METODOLOGIA": {"type": "STRING"},
        "SECUENCIA": {"type": "STRING"},
        "TABLA_GENERAL": {
            "type": "ARRAY",
            "items": {
                "type": "OBJECT",
                "properties": {
                    "RA": {"type": "STRING"},
                    "BLOOM": {"type": "STRING"},
                    "FASE": {"type": "STRING"},
                    "VALOR": {"type": "STRING"},
                    "SEMANAS": {"type": "STRING"},
                    "INSTRUMENTO": {"type": "STRING"}
                },
                "required": ["RA", "BLOOM", "FASE", "VALOR", "SEMANAS", "INSTRUMENTO"]
            }
        },
        "MATRICES": {
            "type": "ARRAY",
            "items": {
                "type": "OBJECT",
                "properties": {
                    "TEXTO": {"type": "STRING"},
                    "BLOOM": {"type": "STRING"},
                    "CONTENIDOS": {"type": "STRING"},
                    "CR": {"type": "STRING"},
                    "CE": {"type": "STRING"},
                    "RECURSOS": {"type": "STRING"},
                    "VALOR": {"type": "STRING"},
                    "SEMANAS": {"type": "STRING"},
                    "INDICADORES": {"type": "STRING"},
                    "DEPENDENCIAS": {"type": "STRING"},
                    "DESEMPENIO_L": {"type": "STRING"},
                    "DESEMPENIO_EP": {"type": "STRING"},
                    "DESEMPENIO_NA": {"type": "STRING"}
                },
                "required": ["TEXTO", "BLOOM", "CONTENIDOS", "CR", "CE", "RECURSOS", "VALOR", "SEMANAS", "INDICADORES", "DEPENDENCIAS", "DESEMPENIO_L", "DESEMPENIO_EP", "DESEMPENIO_NA"]
            }
        },
        "PLAN_CONTINGENCIA": {"type": "STRING"},
        "CONEXIONES_INTERCURRICULARES": {"type": "STRING"},
        "GLOSARIO": {
            "type": "ARRAY",
            "items": {
                "type": "OBJECT",
                "properties": {
                    "TERMINO": {"type": "STRING"},
                    "DEFINICION": {"type": "STRING"}
                },
                "required": ["TERMINO", "DEFINICION"]
            }
        }
    },
    "required": ["RESUMEN", "METODOLOGIA", "SECUENCIA", "TABLA_GENERAL", "MATRICES", "PLAN_CONTINGENCIA", "CONEXIONES_INTERCURRICULARES", "GLOSARIO"]
}

# ===========================================================================
# MARCADORES SEGUROS
# ===========================================================================
MARKER_NL = "<<NL>>"
MARKER_DQ = "<<DQ>>"
MARKER_TAB = "<<TAB>>"

def codificar_marcadores_texto(texto):
    texto = texto.replace('\r\n', MARKER_NL)
    texto = texto.replace('\n', MARKER_NL)
    texto = texto.replace('\r', MARKER_NL)
    texto = texto.replace('\t', MARKER_TAB)
    texto = texto.replace('"', MARKER_DQ)
    return texto

def decodificar_marcadores(obj):
    if isinstance(obj, str):
        return obj.replace(MARKER_NL, '\n').replace(MARKER_DQ, '"').replace(MARKER_TAB, '\t')
    elif isinstance(obj, dict):
        return {decodificar_marcadores(k): decodificar_marcadores(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [decodificar_marcadores(item) for item in obj]
    return obj

# ===========================================================================
# PDF
# ===========================================================================
def limpiar_texto_pdf(texto):
    texto = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f-\x9f]', '', texto)
    texto = texto.replace('\u201c', '"').replace('\u201d', '"')
    texto = texto.replace('\u2018', "'").replace('\u2019', "'")
    texto = re.sub(r'[ \t]{3,}', ' ', texto)
    texto = re.sub(r'\n{3,}', '\n\n', texto)
    return texto

def extraer_texto_pdf(archivo_pdf):
    pdf_reader = PyPDF2.PdfReader(archivo_pdf)
    total_paginas = len(pdf_reader.pages)
    texto = "".join([pagina.extract_text() or "" for pagina in pdf_reader.pages])
    if not texto.strip():
        raise ValueError("No se pudo extraer texto del PDF.")
    alnum_count = sum(1 for c in texto if c.isalnum())
    ratio = alnum_count / max(len(texto), 1)
    if ratio < 0.4:
        raise ValueError(f"Texto corrupto (ratio: {ratio:.1%}).")
    texto = limpiar_texto_pdf(texto)
    if len(texto) > 60000:
        texto = texto[:60000]
    return texto, total_paginas

# ===========================================================================
# API
# ===========================================================================
@retry(retry=retry_if_exception_type(ResourceExhausted), wait=wait_exponential(multiplier=2, min=4, max=20), stop=stop_after_attempt(5), reraise=True)
def solicitar_gemini(api_key, modelo, prompt, schema=None):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(modelo)
    config_kwargs = {"max_output_tokens": 65536, "temperature": 0.0, "response_mime_type": "application/json"}
    if schema is not None:
        try:
            config_kwargs["response_schema"] = schema
        except Exception:
            pass
    config = genai.types.GenerationConfig(**config_kwargs)
    return model.generate_content(prompt, generation_config=config).text

@retry(retry=retry_if_exception_type(OpenAIRateLimitError), wait=wait_exponential(multiplier=2, min=4, max=20), stop=stop_after_attempt(5), reraise=True)
def solicitar_openai(api_key, modelo, prompt):
    client = OpenAI(api_key=api_key)
    response = client.chat.completions.create(
        model=modelo, messages=[{"role": "user", "content": prompt}],
        temperature=0.0, max_tokens=16384, response_format={"type": "json_object"}
    )
    return response.choices[0].message.content

def llamar_ia(api_key, proveedor, modelo, prompt, schema=None):
    if proveedor == "Google Gemini":
        return solicitar_gemini(api_key, modelo, prompt, schema=schema)
    else:
        return solicitar_openai(api_key, modelo, prompt)

# ===========================================================================
# JSON ROBUSTO
# ===========================================================================
def reparar_json_truncado(texto):
    in_string = False
    escape_next = False
    llaves = 0
    corchetes = 0
    last_safe_pos = 0
    for i, char in enumerate(texto):
        if escape_next:
            escape_next = False
            continue
        if in_string:
            if char == '\\':
                escape_next = True
            elif char == '"':
                in_string = False
                last_safe_pos = i + 1
            continue
        if char == '"':
            in_string = True
        elif char == '{':
            llaves += 1
            last_safe_pos = i + 1
        elif char == '}':
            llaves -= 1
            last_safe_pos = i + 1
        elif char == '[':
            corchetes += 1
            last_safe_pos = i + 1
        elif char == ']':
            corchetes -= 1
            last_safe_pos = i + 1
        elif char in (',', ':', ' ', '\n', '\r', '\t'):
            last_safe_pos = i + 1
    repair = texto[:last_safe_pos]
    if in_string:
        repair += '"'
    repair = repair.rstrip()
    if repair.endswith(','):
        repair = repair[:-1]
    repair += ']' * max(corchetes, 0)
    repair += '}' * max(llaves, 0)
    return repair

def parsear_json_robusto(respuesta):
    texto = respuesta.strip()
    if texto.startswith("```json"):
        texto = texto[7:]
    elif texto.startswith("```"):
        texto = texto[3:]
    if texto.endswith("```"):
        texto = texto[:-3]
    texto = texto.strip()
    try:
        return json.loads(texto)
    except json.JSONDecodeError:
        pass
    match = re.search(r'(\{[\s\S]*\})', texto)
    if match:
        try:
            return json.loads(match.group(1))
        except json.JSONDecodeError:
            pass
    json_start = texto.find('{')
    if json_start >= 0:
        json_body = texto[json_start:]
        try:
            return json.loads(reparar_json_truncado(json_body))
        except json.JSONDecodeError:
            pass
        for end_pos in range(len(json_body), max(len(json_body) - 8000, json_start), -1):
            if end_pos <= json_start:
                break
            if json_body[end_pos - 1] == '}':
                try:
                    return json.loads(reparar_json_truncado(json_body[:end_pos]))
                except json.JSONDecodeError:
                    continue
    try:
        return json.loads(re.sub(r'[\x00-\x1f]', '', texto))
    except json.JSONDecodeError:
        pass
    raise ValueError(f"JSON irrecuperable. Inicio: {texto[:500]}...")

# ===========================================================================
# SESSION STATE
# ===========================================================================
if 'fase1_resultado' not in st.session_state:
    st.session_state.fase1_resultado = None
if 'fase1_form_data' not in st.session_state:
    st.session_state.fase1_form_data = None

api_key_usuario = st.session_state.get("api_key_global", "")
proveedor_ia = st.session_state.get("proveedor_ia_global", "Google Gemini")
modelo_seleccionado = st.session_state.get("modelo_global", "gemini-2.5-flash")

with st.sidebar:
    st.markdown("##### ⚡ Ponderación RA")
    if not api_key_usuario:
        st.error("🔒 Configura tu API Key en la página de Inicio")
    else:
        st.success(f"✅ {proveedor_ia} · {modelo_seleccionado}")

st.markdown('<div class="main-header">Sistema de Ponderación Analítica por R.A.</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">Extracción Literal con CR + Contenidos · Formato Oficial MINERD</div>', unsafe_allow_html=True)

# ===========================================================================
# FORMULARIO
# ===========================================================================
with st.form("form_ponderacion", clear_on_submit=False):
    st.markdown('<div class="section-title">📄 1. Fuente Curricular (PDF)</div>', unsafe_allow_html=True)
    archivo_pdf = st.file_uploader("Cargue el documento PDF oficial del diseño curricular", type=["pdf"],
        help="RECOMENDACIÓN: Sube solo las páginas del módulo a trabajar.")

    st.markdown('<div class="section-title">🏛️ 2. Datos Institucionales</div>', unsafe_allow_html=True)
    col_inst1, col_inst2 = st.columns(2)
    with col_inst1:
        politecnico = st.text_input("Nombre del Politécnico", value="Politécnico Salesiano Arquides Calderón")
        docente = st.text_input("Nombre del Docente", value="Ing. Bernardo Antonio Hernández Batista")
    with col_inst2:
        ano_escolar = st.text_input("Año Escolar", placeholder="Ej: 2026-2027", value="2026-2027")
        regional_distrito = st.text_input("Regional / Distrito", value="Regional 06 - Distrito 06-02 (Moca)")

    st.markdown('<div class="section-title">📝 3. Parámetros Curriculares</div>', unsafe_allow_html=True)
    col_par1, col_par2 = st.columns(2)
    with col_par1:
        bachillerato = st.text_input("Bachillerato Técnico", placeholder="Ej: Logística y Transporte")
        modulo = st.text_input("Módulo Formativo", placeholder="Ej: MF_358_3: Impuesto al consumo y a vehículos de motor")
    with col_par2:
        unidad_competencia = st.text_input("Unidad de Competencia (UC)", placeholder="Ej: UC_358_3: Asistir en declaraciones...")
        semanas_totales = st.number_input("Semanas Totales del Módulo", min_value=1, max_value=50, value=38)

    st.markdown("""
    <div style="background:#FEF3C7;border-left:4px solid #F59E0B;padding:10px 14px;border-radius:6px;margin:10px 0;">
        <b>💡 Búsqueda Flexible:</b> Escribe el nombre del módulo de cualquier forma.<br>
        El sistema detecta la nomenclatura propia de tu documento y extrae <b>R.A. + CR + CE + Contenidos</b> respetándola.
    </div>
    """, unsafe_allow_html=True)

    st.markdown("<br>", unsafe_allow_html=True)
    submit_button = st.form_submit_button("🔍 Fase 1: Buscar Módulo y Extraer R.A. + CR + Contenidos")

# ===========================================================================
# FASE 1: EXTRACCIÓN CON CR Y CE DINÁMICOS
# ===========================================================================
if submit_button:
    if not api_key_usuario:
        st.error("🔒 Debes ingresar tu API Key.")
    elif not archivo_pdf or not modulo:
        st.warning("📝 Carga el PDF y escribe el nombre del Módulo.")
    elif not bachillerato or not unidad_competencia:
        st.warning("📝 Completa Bachillerato y UC.")
    else:
        with st.spinner(f'🧠 Fase 1: Extrayendo R.A. + CR + Contenidos de "{modulo}" con {modelo_seleccionado}...'):
            try:
                texto_curriculo, total_paginas = extraer_texto_pdf(archivo_pdf)

                with st.expander(f"📖 Preview ({total_paginas} págs, {len(texto_curriculo):,} chars)"):
                    st.text(texto_curriculo[:3000])
                    if len(texto_curriculo) > 3000:
                        st.caption(f"... y {len(texto_curriculo) - 3000:,} chars más")

                # ═══════════════════════════════════════════════════════════
                # PROMPT FASE 1 — EXTRACCIÓN ADAPTATIVA AL ESQUEMA DEL DOCUMENTO
                # ═══════════════════════════════════════════════════════════
                prompt_fase1 = f"""Eres un motor de EXTRACCIÓN LITERAL de documentos curriculares del MINERD.
Tu ÚNICA función es localizar información en el documento y copiarla EXACTAMENTE como aparece, sin modificar ni una palabra.

═══════════════════════════════════════════════════════════════════════════
TAREA PRINCIPAL:
═══════════════════════════════════════════════════════════════════════════
Buscar el MÓDULO FORMATIVO: "{modulo}"
Extraer TODOS sus Resultados de Aprendizaje (R.A.) con:
  - Sus Criterios de Realización (CR)
  - Sus Criterios de Evaluación (CE) si existen
  - Sus Contenidos (Conceptuales, Procedimentales, Actitudinales)

═══════════════════════════════════════════════════════════════════════════
ESTRATEGIA DE BÚSQUEDA FLEXIBLE (aplicar en orden):
═══════════════════════════════════════════════════════════════════════════
a) Coincidencia exacta del nombre
b) Por código MF_XXX_X si está presente
c) Por palabras clave del nombre del módulo
d) Por coincidencia parcial con las 2-3 palabras más distintivas
e) Buscar en sección "PLAN DE ESTUDIOS" o "MÓDULOS FORMATIVOS" del documento

═══════════════════════════════════════════════════════════════════════════
PASO 0 — DETECTA EL ESQUEMA PROPIO DE ESTE DOCUMENTO (OBLIGATORIO):
═══════════════════════════════════════════════════════════════════════════
Los diseños curriculares del MINERD NO siempre usan la misma nomenclatura ni
numeración. Antes de extraer nada, identifica cómo ESTE documento específico
nombra sus componentes. Ejemplos de variantes que puedes encontrar:

  • Los R.A. pueden llamarse: "R.A.", "RA1", "Elemento de Competencia (EC)",
    "EC8.1", "Objetivo de Aprendizaje", "Resultado Esperado", "Competencia
    Específica", etc.

  • Los Criterios de Realización (CR) pueden llamarse: "Criterio de
    Realización", "CR8.1.1", "Criterio de Desempeño", "Indicador de Logro",
    o pueden aparecer SIN etiqueta, solo como una lista (numerada 8.1.1,
    simple 1., con letras a), con viñetas "-", con romanos i., etc.)
    inmediatamente debajo de cada R.A.

  • Los Criterios de Evaluación (CE) pueden llamarse: "Criterio de
    Evaluación", "CE1.1", "Criterio de Valoración", o pueden NO existir como
    sección separada (en ese caso, derívalos de los CR más específicos).

REGLA DE ORO: USA LA NOMENCLATURA Y LOS CÓDIGOS TAL COMO APARECEN EN *ESTE*
DOCUMENTO. No fuerces el formato "CR8.1.1" ni "EC8.1" si el documento usa otro
esquema — copia los códigos reales que encuentres (letras, números, viñetas,
romanos, o ninguno). Si el documento no numera sus criterios, extrae el texto
de cada uno igual, simplemente sin código antepuesto.

Resume en 1-2 frases el esquema que detectaste en el campo ESQUEMA_DETECTADO
del JSON de salida (ej: "EC/CR con códigos numéricos tipo EC8.1/CR8.1.1",
"R.A. numerados RA1, RA2... con criterios en viñetas sin código",
"Resultado de Aprendizaje / Criterio de Evaluación con letras a), b), c)").

═══════════════════════════════════════════════════════════════════════════
PROCESO DE EXTRACCIÓN:
═══════════════════════════════════════════════════════════════════════════
1) R.A. = cada Resultado de Aprendizaje / Elemento de Competencia (bajo el
   nombre que use el documento). Extrae su código/ID EXACTO tal como aparece
   (o "R.A. N" en orden si el documento no los codifica).

2) CRITERIOS DE REALIZACIÓN (CR) — ¡CRÍTICO!:
   Los CR están normalmente DIRECTAMENTE DEBAJO de cada R.A. en el documento.
   DEBES extraer TODOS los CR de cada R.A., preservando:
   - El código/numeración propia del documento SI EXISTE (sea cual sea su
     formato real).
   - El texto literal completo del criterio.
   NUNCA dejes el campo CR vacío o con "NO ESPECIFICADO" si el documento
   tiene criterios para ese R.A. — extrae el texto aunque no tenga código.

3) CRITERIOS DE EVALUACIÓN (CE):
   Si existen CE explícitos (diferentes de los CR), extraerlos con su propia
   nomenclatura. Si NO existen CE explícitos separados de los CR, entonces
   deriva los CE a partir de los CR: cada CR específico funciona como un
   criterio evaluable.

4) CONTENIDOS — ¡CRÍTICO!:
   Busca los Contenidos del módulo. Pueden aparecer como:
   a) Sección explícita "Contenidos" con subsecciones (Conceptuales,
      Procedimentales, Actitudinales)
   b) Texto descriptivo dentro del módulo
   c) Tabla de contenidos asociada al módulo
   d) Si NO encuentras contenidos explícitos, INFIERE los contenidos
      conceptuales a partir de los CR: los temas, conceptos y procedimientos
      mencionados en los CR son los contenidos que el R.A. debe trabajar.
   El campo CONTENIDOS debe incluir TODO lo que el R.A. aborda: conceptos,
   procedimientos, normas, formularios, plataformas, etc. Separa cada
   contenido con {MARKER_NL}.

═══════════════════════════════════════════════════════════════════════════
REGLAS ESTRICTAS:
═══════════════════════════════════════════════════════════════════════════
1. EXTRACCIÓN PURA: Solo copia lo que aparece en el documento.
2. CERO ALUCINACIÓN: NO inventes CR ni CE que no estén en el documento.
3. Si no encuentras el módulo: MODULO_ENCONTRADO = false.
4. RESPETA LA NOMENCLATURA DEL DOCUMENTO: incluye el código de cada CR/CE tal
   como aparece en ESTE documento (sea CR8.1.1, 1.1, a), viñeta, o ninguno).
   No inventes ni fuerces un formato que el documento no usa.
5. CONTENIDOS COMPLETOS: Extrae TODOS los contenidos, no solo algunos.

CODIFICACIÓN OBLIGATORIA (evita errores JSON):
- Saltos de línea → {MARKER_NL}
- Comillas dobles → {MARKER_DQ}
- Tabulaciones → {MARKER_TAB}
- NUNCA saltos de línea literales dentro de valores JSON.

FORMATO DE SALIDA (JSON VÁLIDO):
{{
  "MODULO_ENCONTRADO": true,
  "MOTIVO_NO_ENCONTRADO": "",
  "NOMBRE_MODULO_HALLADO": "nombre literal encontrado",
  "CODIGO_MODULO": "MF_XXX_X o vacío",
  "ESQUEMA_DETECTADO": "breve descripción de cómo ESTE documento nombra y numera sus R.A., CR y CE",
  "CANTIDAD_RA": 0,
  "RESULTADOS_APRENDIZAJE": [
    {{
      "ID": "código/identificador EXACTO tal como aparece en el documento (ej: EC8.1, RA1, Resultado 1)",
      "TEXTO": "texto literal completo del R.A.",
      "CR": "Criterios de realización tal como aparecen en el documento, con o sin código propio, separados por {MARKER_NL}. Ejemplo CON códigos: 'CR8.1.1 Texto literal{MARKER_NL}CR8.1.2 Texto literal'. Ejemplo SIN códigos: 'Texto literal del primer criterio{MARKER_NL}Texto literal del segundo criterio'",
      "CE": "Criterios de evaluación literales separados por {MARKER_NL}, o derivados de los CR si no existen como sección separada",
      "CONTENIDOS": "Contenido conceptual 1{MARKER_NL}Contenido conceptual 2{MARKER_NL}Contenido procedimental 1{MARKER_NL}Contenido actitudinal 1"
    }}
  ]
}}

═══════════════════════════════════════════════════════════════════════════
TEXTO DEL DOCUMENTO A ANALIZAR:
═══════════════════════════════════════════════════════════════════════════
{texto_curriculo}
"""

                respuesta_fase1 = llamar_ia(api_key_usuario, proveedor_ia, modelo_seleccionado, prompt_fase1, schema=SCHEMA_EXTRACCION)
                datos_fase1 = parsear_json_robusto(respuesta_fase1)
                datos_fase1 = decodificar_marcadores(datos_fase1)

                if not datos_fase1.get("MODULO_ENCONTRADO", False):
                    motivo = datos_fase1.get("MOTIVO_NO_ENCONTRADO", "")
                    st.error(f'❌ Módulo "{modulo}" NO encontrado.')
                    if motivo:
                        st.warning(f"**IA reporta:** {motivo}")
                    st.info("💡 Revisa el preview del texto. Copia el nombre exacto del módulo como aparece en el PDF.")
                    st.stop()

                ras_extraidos = datos_fase1.get("RESULTADOS_APRENDIZAJE", [])
                if not ras_extraidos:
                    st.error("❌ Módulo encontrado pero sin R.A./EC.")
                    st.stop()

                st.session_state.fase1_resultado = datos_fase1
                st.session_state.fase1_form_data = {
                    "politecnico": politecnico, "docente": docente,
                    "ano_escolar": ano_escolar, "regional_distrito": regional_distrito,
                    "bachillerato": bachillerato, "modulo": modulo,
                    "unidad_competencia": unidad_competencia,
                    "semanas_totales": semanas_totales
                }

                codigo_mod = datos_fase1.get("CODIGO_MODULO", "")
                nombre_mod = datos_fase1.get("NOMBRE_MODULO_HALLADO", modulo)
                st.success(f"✅ **Fase 1:** {len(ras_extraidos)} R.A. con CR y Contenidos extraídos de **{codigo_mod + ': ' if codigo_mod else ''}{nombre_mod}**")

                esquema_detectado = datos_fase1.get("ESQUEMA_DETECTADO", "")
                if esquema_detectado:
                    st.info(f"🔎 **Esquema detectado en el documento:** {esquema_detectado}")

            except ValueError as ve:
                st.error(f"❌ {ve}")
            except ResourceExhausted:
                st.error("❌ Límite de API. Espera unos momentos.")
            except Exception as e:
                st.error(f"⚠️ Error Fase 1: {e}")
                if 'respuesta_fase1' in dir():
                    with st.expander("🔍 Respuesta cruda"):
                        st.text(respuesta_fase1[:3000])

# ===========================================================================
# VERIFICACIÓN + FASE 2
# ===========================================================================
if st.session_state.fase1_resultado is not None:
    datos_fase1 = st.session_state.fase1_resultado
    form_data = st.session_state.fase1_form_data
    ras = datos_fase1.get("RESULTADOS_APRENDIZAJE", [])

    st.markdown("---")
    st.markdown('<div class="section-title">📋 Verificación: R.A. + CR + Contenidos Extraídos</div>', unsafe_allow_html=True)

    codigo_mod = datos_fase1.get("CODIGO_MODULO", "")
    nombre_mod = datos_fase1.get("NOMBRE_MODULO_HALLADO", "")
    esquema_detectado = datos_fase1.get("ESQUEMA_DETECTADO", "")

    col_r1, col_r2, col_r3 = st.columns(3)
    with col_r1:
        st.markdown(f"**🏛️ Módulo:** {codigo_mod + ': ' if codigo_mod else ''}{nombre_mod}")
    with col_r2:
        st.markdown(f"**📊 R.A.:** {len(ras)}")
    with col_r3:
        st.markdown(f"**📅 Semanas:** {form_data['semanas_totales']}")

    if esquema_detectado:
        st.caption(f"🔎 Esquema detectado en este documento: {esquema_detectado}")

    with st.expander("🔍 VER DETALLE — R.A. + CR + Contenidos", expanded=True):
        for i, ra in enumerate(ras):
            ra_id = ra.get('ID', f'R.A. {i+1}')
            ra_texto = ra.get('TEXTO', 'Sin texto')
            ra_cr = ra.get('CR', 'NO ESPECIFICADO')
            ra_ce = ra.get('CE', 'NO ESPECIFICADO')
            ra_cont = ra.get('CONTENIDOS', 'NO ESPECIFICADO')

            st.markdown(f"""
            <div class="ra-card">
                <div class="ra-id">{ra_id}</div>
                <div class="ra-text">{ra_texto[:300]}{'...' if len(ra_texto) > 300 else ''}</div>
            </div>
            """, unsafe_allow_html=True)

            # Verificación adaptativa: no todos los documentos numeran sus CR,
            # así que la ausencia de código NO es un error — solo un CR vacío lo es.
            cr_vacio = (not ra_cr) or ra_cr.strip().upper() in ("", "NO ESPECIFICADO")
            has_cr_codes = bool(re.search(r'(CR\s?\d|CE\s?\d|EC\s?\d|RA\s?\d|\d+\.\d+(\.\d+)?|(^|\n)[a-zA-Z]\)|[IVXivx]+\.)', ra_cr or ""))
            col_d1, col_d2, col_d3 = st.columns(3)
            with col_d1:
                if cr_vacio:
                    label_cr = "**⚠️ C.R. (VACÍO — verificar):**"
                elif has_cr_codes:
                    label_cr = "**📌 C.R. (con códigos del documento):**"
                else:
                    label_cr = "**📌 C.R. (sin códigos — este documento no los numera):**"
                st.markdown(label_cr)
                st.text(ra_cr[:400] + ("..." if len(ra_cr) > 400 else ""))
            with col_d2:
                st.markdown("**📌 C.E.:**")
                st.text(ra_ce[:400] + ("..." if len(ra_ce) > 400 else ""))
            with col_d3:
                st.markdown("**📌 Contenidos:**")
                st.text(ra_cont[:400] + ("..." if len(ra_cont) > 400 else ""))

            if i < len(ras) - 1:
                st.divider()

    st.markdown("""
    <div style="background:#D1FAE5;border-left:4px solid #10B981;padding:10px 14px;border-radius:6px;margin:10px 0;">
        <b>✅ Verificación:</b> Confirma que cada RA tenga sus C.R. y Contenidos completos, respetando la
        nomenclatura propia de ESTE documento (con o sin códigos numéricos — ambos son válidos).<br>
        Si hay ⚠️ "VACÍO", significa que no se encontraron criterios para ese R.A. — verifica el PDF y reinicia si es necesario.
    </div>
    """, unsafe_allow_html=True)

    col_a1, col_a2 = st.columns(2)
    with col_a1:
        confirmar = st.button("✅ Fase 2: Confirmar y Generar Documento Completo", type="primary", use_container_width=True)
    with col_a2:
        reiniciar = st.button("🔄 Reiniciar", use_container_width=True)

    if reiniciar:
        st.session_state.fase1_resultado = None
        st.session_state.fase1_form_data = None
        st.rerun()

    # =======================================================================
    # FASE 2: PONDERACIÓN + WORD COMPLETO
    # =======================================================================
    if confirmar:
        with st.spinner(f'🧠 Fase 2: Generando ponderación completa con {modelo_seleccionado}...'):
            try:
                ras_para_prompt = codificar_marcadores_texto(json.dumps(ras, ensure_ascii=False))

                prompt_fase2 = f"""Actúa como Especialista Curricular ETP del MINERD.

TAREA: A partir de los R.A. con CR, CE y Contenidos extraídos literalmente (respetando la nomenclatura propia
de ESTE documento), genera la ponderación COMPLETA.

DATOS DE ENTRADA (VERIFICADOS — COPIAR EXACTAMENTE):
{ras_para_prompt}

MÓDULO: {form_data['modulo']}
UC: {form_data['unidad_competencia']}
SEMANAS TOTALES: {form_data['semanas_totales']}

REGLAS:
1. RESPETA R.A., CR y CE EXACTAMENTE como fueron extraídos. NO modifiques sus textos ni códigos.
2. Suma Valor (%) = 100%. Suma Semanas = {form_data['semanas_totales']}.
3. Para cada RA: Bloom, Fase, Instrumento, Indicadores, Dependencias, Desempeño L/EP/NA.
4. NO inventes R.A. adicionales.
5. CR Y CE: Copia EXACTAMENTE como aparecen en los datos de entrada, respetando la nomenclatura y códigos
   propios de ESTE documento (tengan o no numeración tipo CR8.1.1 — usa lo que realmente traiga el input).
6. CONTENIDOS: Copia EXACTAMENTE como aparecen en los datos de entrada.

CODIFICACIÓN: Saltos de línea → {MARKER_NL} · Comillas dobles → {MARKER_DQ} · NUNCA saltos de línea literales en valores.

FORMATO DE SALIDA (JSON):
{{
  "RESUMEN": "Resumen ejecutivo",
  "METODOLOGIA": "Metodología de ponderación detallada",
  "SECUENCIA": "Secuencia y orden de los R.A.",
  "TABLA_GENERAL": [
    {{"RA": "ID y texto", "BLOOM": "Nivel", "FASE": "Inicio/Desarrollo/Cierre", "VALOR": "Ej: 20%", "SEMANAS": "Ej: 7", "INSTRUMENTO": "Instrumento"}}
  ],
  "MATRICES": [
    {{
      "TEXTO": "ID y texto del RA (copiado exactamente)",
      "BLOOM": "Nivel de Bloom",
      "CONTENIDOS": "Contenidos (copiados EXACTAMENTE del input)",
      "CR": "Criterios de Realización (copiados EXACTAMENTE del input, con la nomenclatura que ese input traiga)",
      "CE": "Criterios de Evaluación (copiados EXACTAMENTE del input)",
      "RECURSOS": "Recursos didácticos sugeridos",
      "VALOR": "Ej: 20%",
      "SEMANAS": "Ej: 7",
      "INDICADORES": "3-4 indicadores observables separados por {MARKER_NL}",
      "DEPENDENCIAS": "RAs previos o Ninguna",
      "DESEMPENIO_L": "Descriptor Logrado",
      "DESEMPENIO_EP": "Descriptor En Proceso",
      "DESEMPENIO_NA": "Descriptor Necesita Apoyo"
    }}
  ],
  "PLAN_CONTINGENCIA": "Plan si un RA no se logra",
  "CONEXIONES_INTERCURRICULARES": "Conexiones con otras áreas",
  "GLOSARIO": [{{"TERMINO": "Término", "DEFINICION": "Definición"}}]
}}

TABLA_GENERAL y MATRICES: exactamente {len(ras)} elementos.
CR en MATRICES: COPIAR EXACTAMENTE del input, respetando la nomenclatura/códigos que ese input ya trae (sean cuales sean).
"""

                respuesta_fase2 = llamar_ia(api_key_usuario, proveedor_ia, modelo_seleccionado, prompt_fase2, schema=SCHEMA_PONDERACION)
                datos_fase2 = parsear_json_robusto(respuesta_fase2)
                datos_fase2 = decodificar_marcadores(datos_fase2)

                resumen = datos_fase2.get("RESUMEN", "")
                metodologia = datos_fase2.get("METODOLOGIA", "")
                secuencia = datos_fase2.get("SECUENCIA", "")
                tabla_gen = datos_fase2.get("TABLA_GENERAL", [])
                matrices = datos_fase2.get("MATRICES", [])
                plan_contingencia = datos_fase2.get("PLAN_CONTINGENCIA", "")
                conexiones = datos_fase2.get("CONEXIONES_INTERCURRICULARES", "")
                glosario = datos_fase2.get("GLOSARIO", [])

                if not matrices:
                    st.error("❌ Matriz vacía.")
                    st.stop()

                # ============================================================
                # WORD — FORMATO OFICIAL COMPLETO
                # ============================================================
                doc = Document()
                doc.styles['Normal'].font.name = 'Calibri'
                doc.styles['Normal'].font.size = Pt(11)
                for section in doc.sections:
                    section.left_margin = Inches(0.5)
                    section.right_margin = Inches(0.5)

                def shade_cell(cell, color):
                    shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
                    cell._tc.get_or_add_tcPr().append(shd)

                def clean(text):
                    return str(text).replace("**", "").replace("***", "")

                # --- ENCABEZADO ---
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run("MINISTERIO DE EDUCACIÓN DE LA REPÚBLICA DOMINICANA\n")
                r.bold = True
                r.font.size = Pt(12)

                p2 = doc.add_paragraph()
                p2.add_run(f"{form_data['regional_distrito']}\n\n").bold = True
                p2.add_run(f"{form_data['politecnico']}\n").bold = True
                p2.add_run('"Formando Honrados Ciudadanos y Buenos Cristianos"\n\n').italic = True
                rt = p2.add_run("PLANIFICACIÓN DIDÁCTICA - MÓDULO FORMATIVO\n")
                rt.bold = True
                rt.font.size = Pt(13)
                p2.add_run("Sistema de Ponderación por Resultados de Aprendizaje\n")
                p2.add_run("______________________________________________________________________")

                # --- II. DATOS GENERALES ---
                doc.add_heading('II. Datos Generales del Módulo', level=2)
                td = doc.add_table(rows=6, cols=2)
                td.style = 'Table Grid'
                for idx, (e, v) in enumerate([
                    ("Centro Educativo", form_data['politecnico']),
                    ("Docente", form_data['docente']),
                    ("Regional / Distrito", form_data['regional_distrito']),
                    ("Año Escolar", form_data['ano_escolar']),
                    ("Módulo Formativo", form_data['modulo']),
                    ("Duración Total", f"{form_data['semanas_totales']} Semanas")
                ]):
                    td.rows[idx].cells[0].text = e
                    td.rows[idx].cells[0].paragraphs[0].runs[0].bold = True
                    td.rows[idx].cells[1].text = str(v)

                # --- III. RESUMEN ---
                doc.add_heading('III. Resumen Ejecutivo', level=2)
                doc.add_paragraph(clean(resumen))

                # --- IV. METODOLOGÍA ---
                doc.add_heading('IV. Metodología de Ponderación', level=2)
                doc.add_paragraph(clean(metodologia))

                # --- V. SECUENCIA ---
                doc.add_heading('V. Secuencia y Orden de R.A.', level=2)
                doc.add_paragraph(clean(secuencia))

                # --- VI. TABLA GENERAL ---
                doc.add_heading('VI. Tabla General de Distribución', level=2)
                if tabla_gen:
                    tg = doc.add_table(rows=1, cols=6)
                    tg.style = 'Table Grid'
                    for i, t in enumerate(["R.A.", "Bloom", "Fase", "Valor (%)", "Semanas", "Instrumento"]):
                        tg.rows[0].cells[i].text = t
                        tg.rows[0].cells[i].paragraphs[0].runs[0].bold = True
                        shade_cell(tg.rows[0].cells[i], "E2E8F0")
                    for item in tabla_gen:
                        rc = tg.add_row().cells
                        rc[0].text = clean(item.get("RA", ""))
                        rc[1].text = clean(item.get("BLOOM", ""))
                        rc[2].text = clean(item.get("FASE", ""))
                        rc[3].text = clean(item.get("VALOR", ""))
                        rc[4].text = clean(item.get("SEMANAS", ""))
                        rc[5].text = clean(item.get("INSTRUMENTO", ""))
                    rc = tg.add_row().cells
                    rc[0].text = "TOTAL"
                    rc[3].text = "100%"
                    rc[4].text = str(form_data['semanas_totales'])
                    for i in range(6):
                        if rc[i].paragraphs[0].runs:
                            rc[i].paragraphs[0].runs[0].bold = True

                # --- VII. VALIDACIÓN ---
                doc.add_heading('VII. Validación de la Distribución', level=2)
                tv = doc.add_table(rows=3, cols=3)
                tv.style = 'Table Grid'
                for i, t in enumerate(["Concepto", "Valor Obtenido", "Estado"]):
                    tv.rows[0].cells[i].text = t
                    tv.rows[0].cells[i].paragraphs[0].runs[0].bold = True
                    shade_cell(tv.rows[0].cells[i], "E2E8F0")
                suma_pct = sum(float(re.sub(r'[^\d.]', '', item.get("VALOR", "0"))) for item in tabla_gen)
                suma_sem = sum(int(re.sub(r'[^\d]', '', item.get("SEMANAS", "0"))) for item in tabla_gen)
                tv.rows[1].cells[0].text = "Suma de Porcentajes"
                tv.rows[1].cells[1].text = f"{suma_pct:.0f}%"
                tv.rows[1].cells[2].text = "✅ Correcto" if abs(suma_pct - 100) < 1 else "❌ Incorrecto"
                tv.rows[2].cells[0].text = "Suma de Semanas"
                tv.rows[2].cells[1].text = str(suma_sem)
                tv.rows[2].cells[2].text = "✅ Correcto" if suma_sem == form_data['semanas_totales'] else "❌ Incorrecto"

                # --- VIII. MATRIZ DE DESARROLLO CURRICULAR ---
                doc.add_heading('VIII. Matriz de Desarrollo Curricular', level=2)

                for idx_m, item in enumerate(matrices):
                    try:
                        tm = doc.add_table(rows=2, cols=7)
                        tm.style = 'Table Grid'
                        for i, t in enumerate(["R.A. y Bloom", "Contenidos", "C.R.", "C.E.", "Recursos", "Valor", "Semanas"]):
                            tm.cell(0, i).text = t
                            tm.cell(0, i).paragraphs[0].runs[0].bold = True
                            shade_cell(tm.cell(0, i), "E2E8F0")

                        bloom = clean(item.get("BLOOM", ""))
                        ra_texto = clean(item.get("TEXTO", ""))
                        ra_bloom = f"{ra_texto} (Nivel de Bloom: {bloom})" if bloom else ra_texto

                        vals = [
                            ra_bloom,
                            clean(item.get("CONTENIDOS", "")),
                            clean(item.get("CR", "")),
                            clean(item.get("CE", "")),
                            clean(item.get("RECURSOS", "Bibliografía y medios oficiales")),
                            clean(item.get("VALOR", "")),
                            clean(item.get("SEMANAS", ""))
                        ]
                        for i, d in enumerate(vals):
                            tm.cell(1, i).text = d
                        doc.add_paragraph()

                        # Indicadores
                        indicadores = clean(item.get("INDICADORES", ""))
                        if indicadores:
                            pi = doc.add_paragraph()
                            pi.add_run("Indicadores de Logro Observables:\n").bold = True
                            pi.add_run(indicadores)

                        # Dependencias
                        deps = clean(item.get("DEPENDENCIAS", ""))
                        pd = doc.add_paragraph()
                        pd.add_run("Dependencias: ").bold = True
                        pd.add_run(deps if deps else "Ninguna (RA inicial)")

                        # Escala de Valoración
                        doc.add_paragraph()
                        pe = doc.add_paragraph()
                        pe.add_run("Escala de Valoración:").bold = True

                        te = doc.add_table(rows=4, cols=2)
                        te.style = 'Table Grid'
                        te.rows[0].cells[0].text = "Nivel"
                        te.rows[0].cells[1].text = "Descriptor de Desempeño"
                        for c in te.rows[0].cells:
                            c.paragraphs[0].runs[0].bold = True
                            shade_cell(c, "D1FAE5")
                        te.rows[1].cells[0].text = "Logrado (L)"
                        te.rows[1].cells[1].text = clean(item.get("DESEMPENIO_L", ""))
                        te.rows[2].cells[0].text = "En Proceso (EP)"
                        te.rows[2].cells[1].text = clean(item.get("DESEMPENIO_EP", ""))
                        te.rows[3].cells[0].text = "Necesita Apoyo (NA)"
                        te.rows[3].cells[1].text = clean(item.get("DESEMPENIO_NA", ""))

                        doc.add_paragraph("__________________________________________________")
                        doc.add_paragraph()
                    except Exception as e:
                        st.warning(f"⚠️ Error RA {idx_m+1}: {e}")

                # --- IX. PLAN CONTINGENCIA ---
                doc.add_heading('IX. Plan de Contingencia', level=2)
                doc.add_paragraph(clean(plan_contingencia))

                # --- X. CONEXIONES ---
                doc.add_heading('X. Conexiones Intercurriculares', level=2)
                doc.add_paragraph(clean(conexiones))

                # --- XI. GLOSARIO ---
                doc.add_heading('XI. Glosario Técnico', level=2)
                if glosario:
                    tgl = doc.add_table(rows=1, cols=2)
                    tgl.style = 'Table Grid'
                    tgl.rows[0].cells[0].text = "Término"
                    tgl.rows[0].cells[1].text = "Definición"
                    for c in tgl.rows[0].cells:
                        c.paragraphs[0].runs[0].bold = True
                        shade_cell(c, "E2E8F0")
                    for term in glosario:
                        rc = tgl.add_row().cells
                        rc[0].text = clean(term.get("TERMINO", ""))
                        rc[1].text = clean(term.get("DEFINICION", ""))

                # --- XII. CONTROL AVANCE ---
                doc.add_heading('XII. Tabla de Control de Avance por R.A.', level=2)
                p_av = doc.add_paragraph()
                p_av.add_run("El docente marca el progreso semanal por cada R.A. (L = Logrado, EP = En Proceso, NA = Necesita Apoyo).").italic = True

                num_sem = 10
                tca = doc.add_table(rows=len(tabla_gen) + 1, cols=1 + num_sem)
                tca.style = 'Table Grid'
                tca.rows[0].cells[0].text = "R.A."
                tca.rows[0].cells[0].paragraphs[0].runs[0].bold = True
                shade_cell(tca.rows[0].cells[0], "E2E8F0")
                for s in range(num_sem):
                    tca.rows[0].cells[s+1].text = f"Sem {s+1}"
                    tca.rows[0].cells[s+1].paragraphs[0].runs[0].bold = True
                    tca.rows[0].cells[s+1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    shade_cell(tca.rows[0].cells[s+1], "E2E8F0")
                for i, item in enumerate(tabla_gen):
                    ra_label = clean(item.get("RA", ""))
                    if len(ra_label) > 60:
                        ra_label = ra_label[:57] + "..."
                    tca.rows[i+1].cells[0].text = ra_label

                # --- XIII. OBSERVACIONES ---
                doc.add_heading('XIII. Observaciones del Coordinador', level=2)
                for _ in range(3):
                    doc.add_paragraph("____________________________________________________________")

                # --- FIRMAS ---
                doc.add_paragraph("\n\n")
                pf = doc.add_paragraph()
                pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pf.add_run("_________________________\t_________________________\t_________________________\n")
                pf.add_run("Director/a de Centro\tCoordinador/a ETP\tDocente\n").bold = True

                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                def sanear_nombre(texto):
                    if not texto: return "archivo"
                    t = unicodedata.normalize('NFKD', texto).encode('ascii', 'ignore').decode('ascii')
                    t = re.sub(r'[^A-Za-z0-9_\- ]', '', t).strip()
                    return re.sub(r'\s+', '_', t)[:60]

                st.success(f"✅ **Documento completo** — {len(matrices)} R.A. con CR y Contenidos · Formato MINERD")

                st.download_button(
                    label="📥 Descargar Documento de Ponderación (.docx)",
                    data=buffer,
                    file_name=f"Ponderacion_RA_{sanear_nombre(form_data['modulo'])}.docx",
                    mime="application/vnd.openxmlformats.officedocument.wordprocessingml.document",
                    type="primary"
                )

                st.session_state.fase1_resultado = None
                st.session_state.fase1_form_data = None

            except ResourceExhausted:
                st.error("❌ Límite de API.")
            except ValueError as ve:
                st.error(f"❌ Error JSON Fase 2: {ve}")
                if 'respuesta_fase2' in dir():
                    with st.expander("🔍 Respuesta cruda"):
                        st.text(respuesta_fase2[:3000])
            except Exception as e:
                st.error(f"⚠️ Error Fase 2: {e}")