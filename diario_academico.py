import streamlit as st
import google.generativeai as genai
from openai import OpenAI, RateLimitError as OpenAIRateLimitError
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from io import BytesIO
import PyPDF2
import json
import re
from google.api_core.exceptions import ResourceExhausted
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type

# --- ESTILOS CSS ---
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    html, body, [class*="css"] { font-family: 'Inter', sans-serif; background-color: #F8FAFC; color: #1E293B; }
    .main-header { font-size: 2.2rem; font-weight: 700; color: #0F172A; text-align: center; margin-bottom: 5px; line-height: 1.2; }
    .sub-header { text-align: center; color: #475569; font-size: 1.1rem; font-weight: 400; margin-bottom: 35px; }
    [data-testid="stForm"] { background-color: #FFFFFF; border: 1px solid #E2E8F0; border-radius: 8px; box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05); padding: 30px; }
    .section-title { color: #1D4ED8; font-weight: 600; font-size: 1.2rem; border-bottom: 2px solid #DBEAFE; padding-bottom: 8px; margin-top: 20px; margin-bottom: 18px; }
    div.stButton > button:first-child, div.stFormSubmitButton > button:first-child { background-color: #2563EB !important; color: #FFFFFF !important; border: none !important; border-radius: 6px !important; font-weight: 600 !important; padding: 10px 24px !important; width: 100%; }
    div.stButton > button:first-child:hover, div.stFormSubmitButton > button:first-child:hover { background-color: #1D4ED8 !important; }
</style>
""", unsafe_allow_html=True)

# --- FUNCIONES DE TEXTO Y JSON ROBUSTO ---
MARKER_NL = "<<NL>>"
MARKER_DQ = "<<DQ>>"
MARKER_TAB = "<<TAB>>"

def decodificar_marcadores(obj):
    if isinstance(obj, str):
        return obj.replace(MARKER_NL, "\n").replace(MARKER_DQ, '"').replace(MARKER_TAB, "\t")
    if isinstance(obj, dict):
        return {decodificar_marcadores(k): decodificar_marcadores(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [decodificar_marcadores(item) for item in obj]
    return obj

def reparar_json_truncado(texto):
    in_string = False
    escape_next = False
    llaves = corchetes = 0
    last_safe_pos = 0
    for i, char in enumerate(texto):
        if escape_next:
            escape_next = False
            continue
        if in_string:
            if char == "\\": escape_next = True
            elif char == '"':
                in_string = False
                last_safe_pos = i + 1
            continue
        if char == '"': in_string = True
        elif char == "{": llaves += 1; last_safe_pos = i + 1
        elif char == "}": llaves -= 1; last_safe_pos = i + 1
        elif char == "[": corchetes += 1; last_safe_pos = i + 1
        elif char == "]": corchetes -= 1; last_safe_pos = i + 1
        elif char in (",", ":", " ", "\n", "\r", "\t"): last_safe_pos = i + 1

    reparado = texto[:last_safe_pos]
    if in_string: reparado += '"'
    reparado = reparado.rstrip()
    if reparado.endswith(","): reparado = reparado[:-1]
    reparado += "]" * max(corchetes, 0) + "}" * max(llaves, 0)
    return reparado

def parsear_json_robusto(respuesta):
    if not respuesta or not respuesta.strip(): raise ValueError("La IA devolvió una respuesta vacía.")
    texto = respuesta.strip()
    if texto.startswith("```json"): texto = texto[7:]
    elif texto.startswith("```"): texto = texto[3:]
    if texto.endswith("```"): texto = texto[:-3]
    texto = texto.strip()
    try: return json.loads(texto, strict=False)
    except json.JSONDecodeError: pass
    match = re.search(r"(\{[\s\S]*\})", texto)
    if match:
        try: return json.loads(match.group(1), strict=False)
        except json.JSONDecodeError: pass
    json_start = texto.find("{")
    if json_start >= 0:
        cuerpo = texto[json_start:]
        try: return json.loads(reparar_json_truncado(cuerpo), strict=False)
        except json.JSONDecodeError: pass
    raise ValueError(f"JSON irrecuperable. Inicio: {texto[:400]}...")

# --- EXTRACTOR DE TEXTO (PDF Y WORD) ---
def extraer_texto_documento(archivo):
    if archivo.name.lower().endswith(".pdf"):
        pdf_reader = PyPDF2.PdfReader(archivo)
        return "".join([pagina.extract_text() or "" for pagina in pdf_reader.pages])
    elif archivo.name.lower().endswith(".docx"):
        doc = Document(archivo)
        texto = []
        for para in doc.paragraphs:
            texto.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    texto.append(cell.text)
        return "\n".join(texto)
    return ""

# --- FUNCIONES DE API CON REINTENTOS ---
@retry(retry=retry_if_exception_type(ResourceExhausted), wait=wait_exponential(multiplier=2, min=4, max=20), stop=stop_after_attempt(5), reraise=True)
def solicitar_gemini_con_reintento(api_key, modelo, prompt):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(modelo)
    respuesta = model.generate_content(
        prompt,
        generation_config=genai.types.GenerationConfig(
            max_output_tokens=8192, temperature=0.2, response_mime_type="application/json"
        )
    )
    return respuesta.text

@retry(retry=retry_if_exception_type(OpenAIRateLimitError), wait=wait_exponential(multiplier=2, min=4, max=20), stop=stop_after_attempt(5), reraise=True)
def solicitar_openai_con_reintento(api_key, modelo, prompt):
    client = OpenAI(api_key=api_key)
    response = client.chat.completions.create(
        model=modelo, messages=[{"role": "user", "content": prompt}],
        temperature=0.2, max_tokens=8192, response_format={"type": "json_object"}
    )
    return response.choices[0].message.content

# --- GENERADOR DE WORD (ESTRUCTURA PDF OFICIAL) ---
def generar_documento_oficial(datos, form):
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(9)

    # Ajuste de márgenes (estrechos)
    for section in doc.sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

    def shade_cell(cell, color):
        shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
        cell._tc.get_or_add_tcPr().append(shd)

    # ENCABEZADO
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_header.add_run("Ministerio de Educación de la República Dominicana\n").bold = True
    p_header.add_run("DATOS GENERALES").bold = True

    # TABLA 1: DATOS GENERALES
    t_gen = doc.add_table(rows=5, cols=4)
    t_gen.style = 'Table Grid'
    
    # Fila 0
    t_gen.cell(0,0).text = f"Nombre completo: {form['docente']}"
    t_gen.cell(0,0).merge(t_gen.cell(0,1))
    t_gen.cell(0,2).text = f"Cédula: {form['cedula']}"
    t_gen.cell(0,2).merge(t_gen.cell(0,3))
    
    # Fila 1
    t_gen.cell(1,0).text = f"Regional: {form['regional']}"
    t_gen.cell(1,1).text = f"Distrito: {form['distrito']}"
    t_gen.cell(1,2).text = f"Centro Educativo: {form['centro']}"
    t_gen.cell(1,2).merge(t_gen.cell(1,3))
    
    # Fila 2
    t_gen.cell(2,0).text = f"Nivel/Sub-Sistema: {form['nivel']}"
    t_gen.cell(2,1).text = f"Ciclo: {form['ciclo']}"
    t_gen.cell(2,2).text = f"Grado y Sección: {form['grado']}"
    t_gen.cell(2,3).text = f"Modalidad: {form['modalidad']}"
    
    # Fila 3
    t_gen.cell(3,0).text = f"Área: {form['area']}"
    t_gen.cell(3,1).text = f"Asignatura: {form['asignatura']}"
    t_gen.cell(3,2).text = f"Semana: {form['semana']}"
    t_gen.cell(3,3).text = f"Código: {form['codigo']}"
    
    # Fila 4
    t_gen.cell(4,0).text = f"Secuencia Didáctica: {form['secuencia']}"
    t_gen.cell(4,0).merge(t_gen.cell(4,1))
    t_gen.cell(4,2).text = f"Duración: {form['duracion']}\nFecha: {form['fecha']}"
    t_gen.cell(4,3).text = f"Actividad: {form['actividad']}"

    # Sombrear la tabla de datos generales
    for row in t_gen.rows:
        for cell in row.cells:
            shade_cell(cell, "E2E8F0")

    p_plan = doc.add_paragraph()
    p_plan.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_plan.add_run("\nPlanificación de Clase Diaria").bold = True

    # TABLA 2: COMPETENCIAS
    t_comp = doc.add_table(rows=4, cols=1)
    t_comp.style = 'Table Grid'
    
    t_comp.cell(0,0).text = "Competencias específicas"
    t_comp.cell(0,0).paragraphs[0].runs[0].bold = True
    shade_cell(t_comp.cell(0,0), "DBEAFE")
    t_comp.cell(1,0).text = "\n".join(datos.get("COMPETENCIAS_ESPECIFICAS", []))
    
    t_comp.cell(2,0).text = "Competencias Fundamentales"
    t_comp.cell(2,0).paragraphs[0].runs[0].bold = True
    shade_cell(t_comp.cell(2,0), "DBEAFE")
    # Generar cuadritos para las competencias fundamentales
    comp_fund = "  ".join([f"☑ {cf}" for cf in datos.get("COMPETENCIAS_FUNDAMENTALES", [])])
    t_comp.cell(3,0).text = comp_fund

    # TABLA 3: CONTENIDOS
    doc.add_paragraph()
    p_cont = doc.add_paragraph()
    p_cont.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_cont.add_run("CONTENIDOS").bold = True
    
    t_cont = doc.add_table(rows=2, cols=4)
    t_cont.style = 'Table Grid'
    headers_cont = ["Conceptos", "Procedimientos", "Actitudes y valores", "Indicadores de Logro"]
    for i, txt in enumerate(headers_cont):
        t_cont.cell(0,i).text = txt
        t_cont.cell(0,i).paragraphs[0].runs[0].bold = True
        shade_cell(t_cont.cell(0,i), "DBEAFE")
    
    cont_data = datos.get("CONTENIDOS", {})
    t_cont.cell(1,0).text = "\n".join(cont_data.get("CONCEPTOS", []))
    t_cont.cell(1,1).text = "\n".join(cont_data.get("PROCEDIMIENTOS", []))
    t_cont.cell(1,2).text = "\n".join(cont_data.get("ACTITUDES_VALORES", []))
    t_cont.cell(1,3).text = "\n".join(cont_data.get("INDICADORES_LOGRO", []))

    # TABLA 4: INTENCIÓN Y ESTRATEGIAS
    doc.add_paragraph()
    p_est = doc.add_paragraph()
    p_est.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_est.add_run("ESTRATEGIAS DE ENSEÑANZA-APRENDIZAJE").bold = True

    t_est = doc.add_table(rows=2, cols=2)
    t_est.style = 'Table Grid'
    t_est.cell(0,0).text = "Intención pedagógica del día"
    t_est.cell(0,0).paragraphs[0].runs[0].bold = True
    shade_cell(t_est.cell(0,0), "F1F5F9")
    t_est.cell(0,1).text = form['intencion']
    
    t_est.cell(1,0).text = "Estrategia / metodología"
    t_est.cell(1,0).paragraphs[0].runs[0].bold = True
    shade_cell(t_est.cell(1,0), "F1F5F9")
    t_est.cell(1,1).text = "\n".join(datos.get("ESTRATEGIAS", [form['estrategias']]))

    # TABLA 5: MOMENTOS (INICIO, DESARROLLO, CIERRE)
    doc.add_paragraph()
    t_mom = doc.add_table(rows=4, cols=3)
    t_mom.style = 'Table Grid'
    headers_mom = ["Momento / Tiempo", "Actividades", "Recursos"]
    for i, txt in enumerate(headers_mom):
        t_mom.cell(0,i).text = txt
        t_mom.cell(0,i).paragraphs[0].runs[0].bold = True
        shade_cell(t_mom.cell(0,i), "DBEAFE")

    momentos_data = datos.get("MOMENTOS", {})
    
    # Inicio
    inicio = momentos_data.get("INICIO", {})
    t_mom.cell(1,0).text = f"INICIO\nDe {inicio.get('TIEMPO', '10 minutos')}"
    t_mom.cell(1,0).paragraphs[0].runs[0].bold = True
    t_mom.cell(1,1).text = inicio.get("ACTIVIDADES", "")
    t_mom.cell(1,2).text = inicio.get("RECURSOS", "")
    
    # Desarrollo
    desarrollo = momentos_data.get("DESARROLLO", {})
    t_mom.cell(2,0).text = f"DESARROLLO\nDe {desarrollo.get('TIEMPO', '35 minutos')}"
    t_mom.cell(2,0).paragraphs[0].runs[0].bold = True
    t_mom.cell(2,1).text = desarrollo.get("ACTIVIDADES", "")
    t_mom.cell(2,2).text = desarrollo.get("RECURSOS", "")

    # Cierre
    cierre = momentos_data.get("CIERRE", {})
    t_mom.cell(3,0).text = f"CIERRE\nDe {cierre.get('TIEMPO', '5 minutos')}"
    t_mom.cell(3,0).paragraphs[0].runs[0].bold = True
    t_mom.cell(3,1).text = cierre.get("ACTIVIDADES", "")
    t_mom.cell(3,2).text = cierre.get("RECURSOS", "")

    # TABLA 6: EVALUACIÓN Y ADAPTACIONES
    doc.add_paragraph()
    t_ev = doc.add_table(rows=4, cols=3)
    t_ev.style = 'Table Grid'
    
    t_ev.cell(0,0).text = "Evidencias"
    t_ev.cell(0,0).paragraphs[0].runs[0].bold = True
    shade_cell(t_ev.cell(0,0), "DBEAFE")
    t_ev.cell(0,1).merge(t_ev.cell(0,2))
    t_ev.cell(0,1).text = "\n".join(datos.get("EVIDENCIAS", []))
    
    headers_ev = ["Tipo", "Técnicas", "Instrumentos"]
    for i, txt in enumerate(headers_ev):
        t_ev.cell(1,i).text = txt
        t_ev.cell(1,i).paragraphs[0].runs[0].bold = True
        shade_cell(t_ev.cell(1,i), "F1F5F9")
        
    t_ev.cell(2,0).text = "\n".join(datos.get("TIPO_EVALUACION", []))
    t_ev.cell(2,1).text = "\n".join(datos.get("TECNICAS", []))
    t_ev.cell(2,2).text = "\n".join(datos.get("INSTRUMENTOS", []))

    t_ev.cell(3,0).text = "ADAPTACIONES (Si aplica)"
    t_ev.cell(3,0).paragraphs[0].runs[0].bold = True
    shade_cell(t_ev.cell(3,0), "FEE2E2")
    t_ev.cell(3,1).merge(t_ev.cell(3,2))
    t_ev.cell(3,1).text = datos.get("ADAPTACIONES", "No aplica")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- CONFIGURACIÓN CENTRALIZADA ---
api_key_usuario = st.session_state.get("api_key_global", "")
proveedor_ia = st.session_state.get("proveedor_ia_global", "Google Gemini")
modelo_seleccionado = st.session_state.get("modelo_global", "gemini-2.5-flash")

with st.sidebar:
    st.markdown("##### ⚡ Plan Diario")
    if not api_key_usuario:
        st.error("🔒 Configura tu API Key en la página de Inicio")
    else:
        st.success(f"✅ {proveedor_ia} · {modelo_seleccionado}")

# --- ENCABEZADO ---
st.markdown('<div class="main-header">Esquema de Planificación Diaria Oficial</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">Formato Estructurado MINERD - Técnico Profesional</div>', unsafe_allow_html=True)

# --- FORMULARIO ---
with st.form("form_plandiario_oficial", clear_on_submit=False):
    
    st.markdown('<div class="section-title">🏫 1. Datos Generales e Institucionales</div>', unsafe_allow_html=True)
    
    col1, col2, col3 = st.columns(3)
    with col1:
        docente = st.text_input("Nombre completo del Docente")
        regional = st.text_input("Regional", value="06")
        nivel = st.text_input("Nivel/Sub-Sistema", value="Secundario")
        area = st.selectbox("Área", ["Matemáticas", "Lengua Española", "Ciencias Naturales", "Ciencias Sociales", "Idiomas", "Técnico Profesional"])
        secuencia = st.text_input("Secuencia Didáctica", placeholder="Ej: La matemática en la minería")
    with col2:
        cedula = st.text_input("Cédula")
        distrito = st.text_input("Distrito", value="06")
        ciclo = st.text_input("Ciclo", value="2do ciclo")
        asignatura = st.text_input("Asignatura", placeholder="Ej: Geometría")
        duracion = st.text_input("Duración", value="50 minutos")
    with col3:
        centro = st.text_input("Centro Educativo", placeholder="Ej: Politécnico Salesiano")
        grado = st.text_input("Grado y Sección", placeholder="Ej: 4to B")
        modalidad = st.text_input("Modalidad", value="Técnico Profesional")
        semana = st.text_input("Semana / Código", placeholder="Semana X / Código Y")
        codigo = st.text_input("Código de Asignatura")
        
    col4, col5 = st.columns(2)
    with col4:
        fecha = st.date_input("Fecha")
    with col5:
        actividad = st.text_input("Número de Actividad", placeholder="Ej: 2da")

    st.markdown('<div class="section-title">📄 2. Malla Curricular Oficial (Obligatoria)</div>', unsafe_allow_html=True)
    archivo_malla = st.file_uploader("Carga el documento de la Malla Curricular (PDF o Word)", type=["pdf", "docx"], help="La IA extraerá contenidos, competencias e indicadores desde este documento.")

    st.markdown('<div class="section-title">📝 3. Estrategia e Intención</div>', unsafe_allow_html=True)
    estrategias = st.text_input("Estrategia / Metodología", placeholder="Ej: Indagación dialógica, Estudio de Caso")
    tema = st.text_input("Tema a tratar", placeholder="Ej: Transformaciones geométricas")
    intencion = st.text_area("Intención pedagógica del día", height=70, placeholder="Ej: Utilizar el plano cartesiano para ubicar puntos mediante la herramienta...")

    st.markdown("<br>", unsafe_allow_html=True)
    submit_button = st.form_submit_button("⚙️ Generar Planificación Diaria Oficial")

# --- LÓGICA CORE ---
if submit_button:
    if not api_key_usuario:
        st.error("🔒 Debes ingresar tu API Key en la página de Inicio.")
    elif not archivo_malla:
        st.warning("⚠️ Debes cargar la Malla Curricular para extraer los elementos obligatorios.")
    elif not tema or not intencion:
        st.warning("📝 Por favor, completa el Tema y la Intención Pedagógica.")
    else:
        with st.spinner(f'🧠 Leyendo malla curricular y diseñando matriz oficial con {modelo_seleccionado}...'):
            respuesta_ia = None
            try:
                texto_malla = extraer_texto_documento(archivo_malla)
                if len(texto_malla) > 60000: texto_malla = texto_malla[:60000]

                if len(texto_malla.strip()) < 50:
                    st.error("❌ No se pudo extraer texto del archivo.")
                    st.stop()

                prompt_maestro = f"""Actúa como experto en planificación educativa del MINERD (República Dominicana).
Diseña la estructura metodológica detallada para una clase de {duracion} del área de {area}.

TEMA DE LA CLASE: {tema}
INTENCIÓN PEDAGÓGICA: {intencion}
ESTRATEGIAS BASE: {estrategias}

MALLA CURRICULAR OFICIAL (FUENTE ESTRICTA):
{texto_malla}

INSTRUCCIONES DE EXTRACCIÓN Y REDACCIÓN:
1. Extrae textualmente de la Malla las Competencias Específicas, Competencias Fundamentales, Conceptos, Procedimientos, Actitudes/Valores e Indicadores de Logro que correspondan al tema.
2. Diseña los Momentos de la clase (INICIO, DESARROLLO, CIERRE) redactando las 'Actividades' paso a paso (saludo, pase de lista, recuperación de saberes, desarrollo guiado, metacognición, etc.) y lista los 'Recursos' a utilizar.
3. Para la Evaluación, sugiere Evidencias, Tipos (Diagnóstica, Formativa), Técnicas e Instrumentos.

FORMATO DE SALIDA ESTRICTO (JSON NATIVO SIN MARKDOWN):
{{
  "COMPETENCIAS_ESPECIFICAS": ["Código - Descripción de competencia 1", "Código - Descripción 2"],
  "COMPETENCIAS_FUNDAMENTALES": ["Ética y Ciudadana", "Resolución de Problemas", "Científica y Tecnológica"],
  "CONTENIDOS": {{
    "CONCEPTOS": ["Concepto 1", "Concepto 2"],
    "PROCEDIMIENTOS": ["Procedimiento 1", "Procedimiento 2"],
    "ACTITUDES_VALORES": ["Actitud 1", "Valor 1"],
    "INDICADORES_LOGRO": ["IL-1 - Descripción"]
  }},
  "ESTRATEGIAS": ["{estrategias}"],
  "MOMENTOS": {{
    "INICIO": {{
      "TIEMPO": "8 minutos",
      "ACTIVIDADES": "Actividades de inicio paso a paso...",
      "RECURSOS": "Pizarra, proyector..."
    }},
    "DESARROLLO": {{
      "TIEMPO": "35 minutos",
      "ACTIVIDADES": "Actividades de desarrollo (trabajo grupal, explicación, etc)...",
      "RECURSOS": "Cuadernos, GeoGebra..."
    }},
    "CIERRE": {{
      "TIEMPO": "7 minutos",
      "ACTIVIDADES": "Preguntas de metacognición...",
      "RECURSOS": "Juego interactivo, cuaderno..."
    }}
  }},
  "EVIDENCIAS": ["Cuadernos", "Gráficos en plataforma"],
  "TIPO_EVALUACION": ["Formativa (Desarrollo y Cierre)", "Heteroevaluación"],
  "TECNICAS": ["Observación directa", "Análisis de procedimientos"],
  "INSTRUMENTOS": ["Lista de cotejo"],
  "ADAPTACIONES": "No aplica"
}}
"""
                if proveedor_ia == "Google Gemini":
                    respuesta_ia = solicitar_gemini_con_reintento(api_key_usuario, modelo_seleccionado, prompt_maestro)
                else:
                    respuesta_ia = solicitar_openai_con_reintento(api_key_usuario, modelo_seleccionado, prompt_maestro)

                datos = parsear_json_robusto(respuesta_ia)
                datos = decodificar_marcadores(datos)

                # Agrupar los datos del formulario
                datos_formulario = {
                    "docente": docente, "cedula": cedula, "regional": regional, "distrito": distrito,
                    "centro": centro, "nivel": nivel, "ciclo": ciclo, "grado": grado, 
                    "modalidad": modalidad, "area": area, "asignatura": asignatura, 
                    "semana": semana, "codigo": codigo, "secuencia": secuencia, 
                    "duracion": duracion, "fecha": fecha.strftime('%d/%m/%Y'), 
                    "actividad": actividad, "intencion": intencion, "estrategias": estrategias
                }

                buffer_docx = generar_documento_oficial(datos, datos_formulario)
                
                st.success("✅ ¡Planificación Diaria Oficial generada con éxito!")
                
                st.download_button(
                    label="📥 Descargar Planificación Oficial MINERD (.docx)",
                    data=buffer_docx,
                    file_name=f"Plan_Diario_Oficial_{area}_{fecha.strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary" 
                )
                
            except ResourceExhausted:
                st.error("❌ Límite de API Gemini alcanzado. Espera unos momentos.")
            except OpenAIRateLimitError:
                st.error("❌ Límite de API OpenAI alcanzado. Espera unos momentos.")
            except ValueError as ve:
                st.error(f"⚠️ Error de procesamiento de JSON: {ve}")
                if respuesta_ia:
                    with st.expander("🔍 Ver respuesta cruda de la IA (Depuración)"):
                        st.text(respuesta_ia[:3000])
            except Exception as e:
                st.error(f"⚠️ Error de procesamiento: {e}")