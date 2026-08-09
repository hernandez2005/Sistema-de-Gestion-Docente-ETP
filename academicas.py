import streamlit as st
import google.generativeai as genai
from openai import OpenAI, RateLimitError as OpenAIRateLimitError
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from io import BytesIO
import PyPDF2
import json
import re

# --- IMPORTACIONES PARA REINTENTOS DE API ---
from google.api_core.exceptions import ResourceExhausted
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type

# --- ESTILOS CSS ---
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    html, body, [class*="css"] { font-family: 'Inter', sans-serif; background-color: #F8FAFC; color: #1E293B; }
    .main-header { font-size: 2.2rem; font-weight: 700; color: #0F172A; text-align: center; margin-bottom: 5px; line-height: 1.2; }
    .sub-header { text-align: center; color: #475569; font-size: 1.1rem; font-weight: 400; margin-bottom: 35px; }
    [data-testid="stForm"] { background-color: #FFFFFF; border: 1px solid #E2E8F0; border-radius: 8px; box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05); padding: 30px; }
    .section-title { color: #1D4ED8; font-weight: 600; font-size: 1.2rem; border-bottom: 2px solid #DBEAFE; padding-bottom: 8px; margin-top: 25px; margin-bottom: 18px; }
    div.stButton > button:first-child, div.stFormSubmitButton > button:first-child { background-color: #2563EB !important; color: #FFFFFF !important; border: none !important; border-radius: 6px !important; font-weight: 600 !important; padding: 10px 24px !important; width: 100%; }
    div.stButton > button:first-child:hover, div.stFormSubmitButton > button:first-child:hover { background-color: #1D4ED8 !important; }
    .btn-extraer > button { background-color: #10B981 !important; color: white !important; }
    .btn-extraer > button:hover { background-color: #059669 !important; }
</style>
""", unsafe_allow_html=True)

# --- INICIALIZAR VARIABLES DE SESIÓN ---
if "elementos_extraidos" not in st.session_state:
    st.session_state.elementos_extraidos = {
        "comp_fundamentales": [], "comp_especificas": [], "indicadores": [],
        "conceptuales": [], "procedimentales": [], "actitudinales": []
    }

# --- FUNCIONES DE JSON ROBUSTO ---
MARKER_NL = "<<NL>>"
def reparar_json_truncado(texto):
    in_string = False; escape_next = False; llaves = corchetes = 0; last_safe_pos = 0
    for i, char in enumerate(texto):
        if escape_next: escape_next = False; continue
        if in_string:
            if char == "\\": escape_next = True
            elif char == '"': in_string = False; last_safe_pos = i + 1
            continue
        if char == '"': in_string = True
        elif char == "{": llaves += 1; last_safe_pos = i + 1
        elif char == "}": llaves -= 1; last_safe_pos = i + 1
        elif char == "[": corchetes += 1; last_safe_pos = i + 1
        elif char == "]": corchetes -= 1; last_safe_pos = i + 1
        elif char in (",", ":", " ", "\n", "\r", "\t"): last_safe_pos = i + 1
    reparado = texto[:last_safe_pos]
    if in_string: reparado += '"'
    reparado = reparado.rstrip().rstrip(",")
    reparado += "]" * max(corchetes, 0) + "}" * max(llaves, 0)
    return reparado

def parsear_json_robusto(respuesta):
    if not respuesta or not respuesta.strip(): raise ValueError("La IA devolvió una respuesta vacía.")
    texto = respuesta.strip().replace("```json", "").replace("```", "").strip()
    try: return json.loads(texto, strict=False)
    except: pass
    try: return json.loads(reparar_json_truncado(texto), strict=False)
    except: pass
    match = re.search(r"(\{[\s\S]*\})", texto)
    if match:
        try: return json.loads(reparar_json_truncado(match.group(1)), strict=False)
        except: pass
    raise ValueError("JSON irrecuperable tras reparación.")

def recortar_opcion(texto, limite=150):
    texto = re.sub(r"\s+", " ", texto).strip()
    return texto if len(texto) <= limite else texto[:limite-3] + "..."

# --- LECTOR DUAL (PDF Y DOCX) CON BÚSQUEDA INTELIGENTE ---
def obtener_seccion_grado(archivo, area, grado):
    texto_capturado = ""
    
    area_map = {
        "Lengua Española": ["lengua española", "español"],
        "Matemática": ["matemática", "matematicas"],
        "Ciencias Sociales": ["ciencias sociales", "sociales"],
        "Ciencias de la Naturaleza": ["ciencias de la naturaleza", "naturales", "biología"],
        "Inglés": ["lenguas extranjeras", "inglés", "ingles"],
        "Francés": ["lenguas extranjeras", "francés", "frances"],
        "Educación Física": ["educación física", "educacion fisica"],
        "Educación Artística": ["educación artística", "educacion artistica"],
        "Formación Integral Humana y Religiosa": ["formación integral", "formacion integral"]
    }
    
    grado_map = {
        "1ro de Secundaria": ["1er. grado", "primer grado", "1er grado", "primero"],
        "2do de Secundaria": ["2do. grado", "segundo grado", "2do grado", "segundo"],
        "3er de Secundaria": ["3er. grado", "tercer grado", "3er grado", "tercero"],
        "4to de Secundaria": ["4to. grado", "cuarto grado", "4to grado", "cuarto"],
        "5to de Secundaria": ["5to. grado", "quinto grado", "5to grado", "quinto"],
        "6to de Secundaria": ["6to. grado", "sexto grado", "6to grado", "sexto"]
    }

    kw_area = area_map.get(area, [area.lower()])
    kw_grado = grado_map.get(grado, [grado.lower()])

    if archivo.name.lower().endswith('.pdf'):
        pdf_reader = PyPDF2.PdfReader(archivo)
        capturando = False
        paginas_restantes = 0
        
        for pagina in pdf_reader.pages:
            texto_pag = pagina.extract_text() or ""
            texto_lower = texto_pag.lower()
            
            tiene_area = any(kw in texto_lower for kw in kw_area)
            tiene_grado = any(kw in texto_lower for kw in kw_grado)
            
            if tiene_area and tiene_grado and ("competencia" in texto_lower or "contenido" in texto_lower):
                capturando = True
                paginas_restantes = 8 
                
            if capturando and paginas_restantes > 0:
                texto_capturado += texto_pag + "\n\n"
                paginas_restantes -= 1
                
    elif archivo.name.lower().endswith('.docx'):
        doc = Document(archivo)
        texto_completo = "\n".join([p.text for p in doc.paragraphs])
        texto_lower = texto_completo.lower()
        
        idx_area = -1
        for kw in kw_area:
            idx = texto_lower.find(kw)
            if idx != -1:
                idx_area = idx
                break
        if idx_area == -1: idx_area = 0
            
        idx_grado = -1
        for kw in kw_grado:
            idx = texto_lower.find(kw, idx_area)
            if idx != -1:
                idx_grado = idx
                break
        
        if idx_grado != -1:
            texto_capturado = texto_completo[idx_grado:idx_grado + 20000]
        else:
            texto_capturado = texto_completo[:20000]

    return texto_capturado[:40000] 

# --- ESCÁNER DE RESPALDO (REGEX) POR SI FALLA LA IA ---
def extraer_elementos_malla_regex(seccion):
    elementos = {
        "comp_fundamentales": [], "comp_especificas": [], "indicadores": [],
        "conceptuales": [], "procedimentales": [], "actitudinales": []
    }
    sec_norm = re.sub(r'[ \t]+', ' ', seccion)
    
    m_comp = re.search(r'Competencias?\s+Espec[íi]ficas?\s+del\s+Grado\s*(.*?)(?=Contenidos|Conceptos|Indicadores|Ejes)', sec_norm, re.DOTALL | re.I)
    if m_comp:
        bloque = m_comp.group(1).strip()
        patron_comp = r'(Comunicativa|Pensamiento Lógico, Creativo y Crítico|Resolución de Problemas|Ética y Ciudadana|Científica y Tecnológica|Ambiental y de la Salud|Desarrollo Personal y Espiritual)\s*(.*?)(?=(?:Comunicativa|Pensamiento Lógico, Creativo y Crítico|Resolución de Problemas|Ética y Ciudadana|Científica y Tecnológica|Ambiental y de la Salud|Desarrollo Personal y Espiritual)|$)'
        for match in re.finditer(patron_comp, bloque, re.DOTALL | re.I):
            nombre = match.group(1).strip()
            desc = match.group(2).strip(' -•\t\n')
            if nombre and nombre not in elementos["comp_fundamentales"]: elementos["comp_fundamentales"].append(nombre)
            if desc and len(desc) > 10: elementos["comp_especificas"].append(f"{nombre}: {desc}")

    m1 = re.search(r'(?:Conceptos|Contenidos Conceptuales)\s*(.*?)(?=Procedimientos|Contenidos Procedimentales|Actitudes|Indicadores|Competencias|$)', sec_norm, re.DOTALL | re.I)
    if m1:
        elementos["conceptuales"] = [l.strip(" -•\t") for l in m1.group(1).split('\n') if len(l.strip(" -•\t")) > 3][:20]
        
    m2 = re.search(r'(?:Procedimientos|Contenidos Procedimentales)\s*(.*?)(?=Actitudes|Valores|Indicadores|Competencias|$)', sec_norm, re.DOTALL | re.I)
    if m2:
        elementos["procedimentales"] = [l.strip(" -•\t") for l in m2.group(1).split('\n') if len(l.strip(" -•\t")) > 3][:20]
        
    m3 = re.search(r'(?:Actitudes y [Vv]alores|Contenidos Actitudinales)\s*(.*?)(?=Indicadores|Competencias|Malla|Ejes|$)', sec_norm, re.DOTALL | re.I)
    if m3:
        elementos["actitudinales"] = [l.strip(" -•\t") for l in m3.group(1).split('\n') if len(l.strip(" -•\t")) > 3][:20]
        
    m_ind = re.search(r'Indicadores?\s+de\s+Logro\s*(.*?)(?=Área\s+de|Malla Curricular|Primer Ciclo|Segundo Ciclo|Ejes Transversales|$)', sec_norm, re.DOTALL | re.I)
    if m_ind:
        elementos["indicadores"] = [l.strip(" -•\t") for l in m_ind.group(1).split('\n') if len(l.strip(" -•\t")) > 10][:25]
        
    return elementos

# --- FUNCIONES DE API CON REINTENTOS ---
@retry(retry=retry_if_exception_type(ResourceExhausted), wait=wait_exponential(multiplier=2, min=4, max=20), stop=stop_after_attempt(5), reraise=True)
def solicitar_gemini(api_key, modelo, prompt):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(modelo)
    respuesta = model.generate_content(
        prompt,
        generation_config=genai.types.GenerationConfig(max_output_tokens=8192, temperature=0.1, response_mime_type="application/json")
    )
    return respuesta.text

@retry(retry=retry_if_exception_type(OpenAIRateLimitError), wait=wait_exponential(multiplier=2, min=4, max=20), stop=stop_after_attempt(5), reraise=True)
def solicitar_openai(api_key, modelo, prompt):
    client = OpenAI(api_key=api_key)
    response = client.chat.completions.create(
        model=modelo, messages=[{"role": "user", "content": prompt}],
        temperature=0.1, max_tokens=8192, response_format={"type": "json_object"}
    )
    return response.choices[0].message.content

# --- FUNCIÓN PARA GENERAR WORD ---
def generar_word_situacion(datos_json, form_data):
    doc = Document()
    doc.styles['Normal'].font.name = 'Arial'
    doc.styles['Normal'].font.size = Pt(10)

    for section in doc.sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)

    def shade_cell(cell, color):
        shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
        cell._tc.get_or_add_tcPr().append(shd)

    p_tit = doc.add_paragraph()
    p_tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_tit = p_tit.add_run("PLANIFICACIÓN POR SITUACIÓN DE APRENDIZAJE\n")
    r_tit.bold = True
    r_tit.font.size = Pt(14)

    doc.add_heading("I. DATOS GENERALES", level=2)
    t1 = doc.add_table(rows=5, cols=2)
    t1.style = 'Table Grid'
    d1 = [
        (f"Centro educativo: {form_data['centro']}", f"Docente: {form_data['docente']}"),
        (f"Área Académica: {form_data['area']}", f"Nivel: Secundario"),
        (f"Grado: {form_data['grado']}", f"Tiempo estimado: {form_data['duracion']}"),
        (f"Fecha de inicio: {form_data['fecha']}", f"Título de la Unidad: {form_data['tema']}"),
        (f"Situación de aprendizaje No.: 1", "")
    ]
    for i, (col1, col2) in enumerate(d1):
        t1.cell(i, 0).text = col1
        t1.cell(i, 1).text = col2

    doc.add_heading("II. SITUACIÓN DE APRENDIZAJE", level=2)
    sit = datos_json.get("situacion_aprendizaje", {})
    p_sa = doc.add_paragraph()
    p_sa.add_run("Contexto:\n").bold = True
    p_sa.add_run(str(sit.get("contexto", "")) + "\n\n")
    p_sa.add_run("Situación o problemática:\n").bold = True
    p_sa.add_run(str(sit.get("situacion_problematica", "")) + "\n\n")
    p_sa.add_run("Desafío:\n").bold = True
    p_sa.add_run(str(sit.get("desafio", "")) + "\n\n")
    p_sa.add_run("Producto final:\n").bold = True
    p_sa.add_run(str(sit.get("producto_final", "")))

    doc.add_heading("III. COMPETENCIAS FUNDAMENTALES", level=2)
    t_cf = doc.add_table(rows=1, cols=2)
    t_cf.style = 'Table Grid'
    hdr_cf = t_cf.rows[0].cells
    hdr_cf[0].text = "Competencia fundamental"; hdr_cf[1].text = "¿Cómo se desarrolla en la situación?"
    for c in hdr_cf: 
        c.paragraphs[0].runs[0].bold = True
        shade_cell(c, "E2E8F0")
    for cf in datos_json.get("competencias_fundamentales_desarrollo", []):
        row = t_cf.add_row().cells
        row[0].text = str(cf.get("competencia", ""))
        row[1].text = str(cf.get("como_se_desarrolla", ""))

    doc.add_heading("IV. COMPETENCIAS ESPECÍFICAS", level=2)
    doc.add_paragraph("\n".join([f"• {c}" for c in form_data['comp_especificas']]))

    doc.add_heading("V. CONTENIDOS", level=2)
    t_cont = doc.add_table(rows=1, cols=3)
    t_cont.style = 'Table Grid'
    hdr_ct = t_cont.rows[0].cells
    hdr_ct[0].text = "Conceptuales"; hdr_ct[1].text = "Procedimentales"; hdr_ct[2].text = "Actitudinales"
    for c in hdr_ct: 
        c.paragraphs[0].runs[0].bold = True
        shade_cell(c, "DBEAFE")
    row_ct = t_cont.add_row().cells
    row_ct[0].text = "\n".join([f"- {c}" for c in form_data['conceptuales']])
    row_ct[1].text = "\n".join([f"- {c}" for c in form_data['procedimentales']])
    row_ct[2].text = "\n".join([f"- {c}" for c in form_data['actitudinales']])

    doc.add_heading("VI. INDICADORES DE LOGRO", level=2)
    t_ind = doc.add_table(rows=1, cols=2)
    t_ind.style = 'Table Grid'
    hdr_ind = t_ind.rows[0].cells
    hdr_ind[0].text = "Indicador de logro"; hdr_ind[1].text = "Evidencia que permitirá comprobarlo"
    for c in hdr_ind: 
        c.paragraphs[0].runs[0].bold = True
        shade_cell(c, "E2E8F0")
    for ie in datos_json.get("indicadores_evidencias", []):
        row = t_ind.add_row().cells
        row[0].text = str(ie.get("indicador", ""))
        row[1].text = str(ie.get("evidencia", ""))

    doc.add_heading("VII. SECUENCIA DE ACTIVIDADES", level=2)
    sec = datos_json.get("secuencia_actividades", {})
    fases_map = {
        "Fase 1. Inicio / Exploración": sec.get("fase_1_inicio", {}),
        "Fase 2. Desarrollo / Construcción": sec.get("fase_2_desarrollo", {}),
        "Fase 3. Aplicación / Producción": sec.get("fase_3_aplicacion", {}),
        "Fase 4. Socialización / Cierre": sec.get("fase_4_cierre", {})
    }
    
    for titulo_fase, datos_fase in fases_map.items():
        doc.add_heading(titulo_fase, level=3)
        p = doc.add_paragraph()
        p.add_run("Propósito: ").bold = True
        p.add_run(str(datos_fase.get("proposito", "")) + "\n")
        
        p.add_run("Actividades del docente:\n").bold = True
        for ad in datos_fase.get("act_docente", []): p.add_run(f"• {ad}\n")
        
        p.add_run("Actividades de los estudiantes:\n").bold = True
        for ae in datos_fase.get("act_estudiantes", []): p.add_run(f"• {ae}\n")
        
        if "producto_parcial" in datos_fase:
            p.add_run(f"Producto parcial/final: ").bold = True
            p.add_run(f"{datos_fase.get('producto_parcial', '')}\n")
            
        p.add_run("Recursos: ").bold = True
        p.add_run(f"{datos_fase.get('recursos', '')}\n")
        
        if "evidencias" in datos_fase:
            p.add_run("Evidencias: ").bold = True
            p.add_run(f"{datos_fase.get('evidencias', '')}\n")
            
        p.add_run("Tiempo: ").bold = True
        p.add_run(f"{datos_fase.get('tiempo', '')}")
        doc.add_paragraph("_" * 60)

    doc.add_heading("VIII. ESTRATEGIAS DE ENSEÑANZA Y APRENDIZAJE", level=2)
    doc.add_paragraph(", ".join(datos_json.get("estrategias", [])))

    doc.add_heading("IX. RECURSOS", level=2)
    rec = datos_json.get("recursos_detallados", {})
    p_rec = doc.add_paragraph()
    p_rec.add_run("Humanos: ").bold = True
    p_rec.add_run(", ".join(rec.get("humanos", [])) + "\n")
    p_rec.add_run("Tecnológicos: ").bold = True
    p_rec.add_run(", ".join(rec.get("tecnologicos", [])) + "\n")
    p_rec.add_run("Materiales: ").bold = True
    p_rec.add_run(", ".join(rec.get("materiales", [])))

    doc.add_heading("X. EVALUACIÓN", level=2)
    ev = datos_json.get("evaluacion", {})
    p_ev = doc.add_paragraph()
    p_ev.add_run("Diagnóstica: ").bold = True
    p_ev.add_run(str(ev.get("diagnostica", "")) + "\n")
    p_ev.add_run("Formativa: ").bold = True
    p_ev.add_run(str(ev.get("formativa", "")) + "\n")
    p_ev.add_run("Sumativa: ").bold = True
    p_ev.add_run(str(ev.get("sumativa", "")) + "\n")
    p_ev.add_run("Técnicas: ").bold = True
    p_ev.add_run(", ".join(ev.get("tecnicas", [])) + "\n")
    p_ev.add_run("Instrumentos: ").bold = True
    p_ev.add_run(", ".join(ev.get("instrumentos", [])))

    doc.add_heading("XI. MATRIZ DE EVALUACIÓN", level=2)
    t_mat = doc.add_table(rows=1, cols=5)
    t_mat.style = 'Table Grid'
    hdr_mat = t_mat.rows[0].cells
    for i, t in enumerate(["Competencia / Indicador", "Evidencia", "Técnica", "Instrumento", "Criterios"]):
        hdr_mat[i].text = t
        hdr_mat[i].paragraphs[0].runs[0].bold = True
        shade_cell(hdr_mat[i], "E2E8F0")
    for m in datos_json.get("matriz_evaluacion", []):
        row = t_mat.add_row().cells
        row[0].text = str(m.get("competencia_indicador", ""))
        row[1].text = str(m.get("evidencia", ""))
        row[2].text = str(m.get("tecnica", ""))
        row[3].text = str(m.get("instrumento", ""))
        row[4].text = str(m.get("criterios", ""))

    doc.add_heading("XII. EVIDENCIAS DE APRENDIZAJE", level=2)
    doc.add_paragraph(", ".join(datos_json.get("evidencias_aprendizaje", [])))

    doc.add_heading("XIII. PRODUCTO FINAL", level=2)
    doc.add_paragraph(str(datos_json.get("situacion_aprendizaje", {}).get("producto_final", "")))

    doc.add_heading("XIV. METACOGNICIÓN", level=2)
    doc.add_paragraph("\n".join([f"• {m}" for m in datos_json.get("metacognicion", [])]))

    doc.add_heading("XV. ATENCIÓN A LA DIVERSIDAD", level=2)
    doc.add_paragraph(str(datos_json.get("atencion_diversidad", "")))

    doc.add_heading("XVI. VINCULACIÓN CON EL CONTEXTO", level=2)
    vinc = datos_json.get("vinculacion", {})
    p_vinc = doc.add_paragraph()
    p_vinc.add_run("Relación con la familia: ").bold = True
    p_vinc.add_run(str(vinc.get("familia", "")) + "\n")
    p_vinc.add_run("Relación con la comunidad: ").bold = True
    p_vinc.add_run(str(vinc.get("comunidad", "")) + "\n")
    p_vinc.add_run("Aplicación en situaciones reales: ").bold = True
    p_vinc.add_run(str(vinc.get("aplicacion_real", "")))

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- CONFIGURACIÓN CENTRALIZADA ---
api_key_usuario = st.session_state.get("api_key_global", "")
proveedor_ia = st.session_state.get("proveedor_ia_global", "Google Gemini")
modelo_seleccionado = st.session_state.get("modelo_global", "gemini-2.5-flash")

with st.sidebar:
    st.markdown("##### ⚡ Planificador Académico")
    if not api_key_usuario:
        st.error("🔒 Configura tu API Key en Inicio")
    else:
        st.success(f"✅ {proveedor_ia} · {modelo_seleccionado}")

# --- ENCABEZADO ---
st.markdown('<div class="main-header">Plan de Unidad por Situación de Aprendizaje</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">Diseño alineado a la Plantilla Oficial MINERD (16 Criterios)</div>', unsafe_allow_html=True)

# ═════════════════════════════════════════════════════════════════════════════
# PASO 1: CARGA Y EXTRACCIÓN (SECCIÓN INDEPENDIENTE)
# ═════════════════════════════════════════════════════════════════════════════
st.markdown('<div class="section-title">📄 1. Carga de Malla Oficial y Extracción Automática</div>', unsafe_allow_html=True)
archivo_doc = st.file_uploader("Sube el PDF o Word (.docx) de la Adecuación Curricular", type=["pdf", "docx"])

col_sel1, col_sel2 = st.columns(2)
with col_sel1:
    areas = ["Matemática", "Lengua Española", "Ciencias Sociales", "Ciencias de la Naturaleza", "Inglés", "Francés", "Educación Física", "Educación Artística", "Formación Integral Humana y Religiosa"]
    area_sel = st.selectbox("Área / Asignatura", areas)
with col_sel2:
    grados = ["1ro de Secundaria", "2do de Secundaria", "3er de Secundaria", "4to de Secundaria", "5to de Secundaria", "6to de Secundaria"]
    grado_sel = st.selectbox("Grado", grados)

st.markdown('<div class="btn-extraer">', unsafe_allow_html=True)
btn_extraer = st.button("🔍 Extraer Componentes Curriculares del Documento", use_container_width=True)
st.markdown('</div>', unsafe_allow_html=True)

if btn_extraer:
    if not api_key_usuario:
        st.error("🔒 Configura tu API Key en la barra lateral primero.")
    elif not archivo_doc:
        st.warning("⚠️ Debes subir el PDF o Word curricular antes de extraer.")
    else:
        with st.spinner(f"🧠 Escaneando el documento y extrayendo los componentes de {area_sel} - {grado_sel}..."):
            try:
                # 1. Extracción estratégica de páginas
                seccion_grado = obtener_seccion_grado(archivo_doc, area_sel, grado_sel)
                
                if seccion_grado.strip():
                    # 2. Prompt estricto de clasificación con límites
                    prompt_extraccion = f"""Actúa como Analista Curricular experto del MINERD.
                    He extraído el texto del currículo de {area_sel} para {grado_sel}.
                    Tu objetivo es estructurar los componentes en un JSON válido.
                    
                    TEXTO CURRICULAR:
                    {seccion_grado}
                    
                    REGLAS INFALIBLES:
                    1. Extrae las Competencias Fundamentales mencionadas explícitamente.
                    2. Extrae las Competencias Específicas del grado.
                    3. Extrae un máximo de 15 Contenidos Conceptuales, 15 Procedimentales y 15 Actitudinales (para evitar saturación de memoria).
                    4. Extrae un máximo de 15 Indicadores de Logro.
                    5. PROHIBIDO usar comillas dobles (") dentro de los textos extraídos. Usa comillas simples (').
                    6. NO inventes contenido externo. 
                    
                    FORMATO JSON OBLIGATORIO:
                    {{
                      "comp_fundamentales": ["..."],
                      "comp_especificas": ["..."],
                      "conceptuales": ["..."],
                      "procedimentales": ["..."],
                      "actitudinales": ["..."],
                      "indicadores": ["..."]
                    }}
                    """
                    
                    try:
                        # Intento con IA
                        if proveedor_ia == "Google Gemini":
                            resp_ext = solicitar_gemini(api_key_usuario, modelo_seleccionado, prompt_extraccion)
                        else:
                            resp_ext = solicitar_openai(api_key_usuario, modelo_seleccionado, prompt_extraccion)
                            
                        datos_ext = parsear_json_robusto(resp_ext)
                        
                    except Exception as e_json:
                        # Fallback a Escáner Regex
                        st.warning("⚠️ El texto era demasiado complejo para la IA. Activando escáner de respaldo (Regex)...")
                        datos_ext = extraer_elementos_malla_regex(seccion_grado)
                    
                    if datos_ext and any(datos_ext.values()):
                        st.session_state.elementos_extraidos = datos_ext
                        st.success("✅ ¡Extracción completada! Selecciona los componentes en el formulario de abajo.")
                    else:
                        st.warning("⚠️ No se encontraron datos en la sección extraída. Verifica que el documento sea correcto.")
                else:
                    st.warning("⚠️ No se encontró la sección exacta de la materia y grado en el documento.")
            except Exception as e:
                st.error(f"Error procesando la extracción: {e}")

# --- OBTENER DATOS DE LA SESIÓN ---
elementos = st.session_state.elementos_extraidos

# ═════════════════════════════════════════════════════════════════════════════
# PASO 2: FORMULARIO PRINCIPAL (GENERACIÓN)
# ═════════════════════════════════════════════════════════════════════════════
with st.form("form_academico"):
    
    st.markdown('<div class="section-title">🏫 2. Datos Generales de la Clase</div>', unsafe_allow_html=True)
    col1, c2, c3 = st.columns(3)
    with col1:
        centro = st.text_input("Centro Educativo", value="Politécnico Salesiano Arquides Calderón")
        docente = st.text_input("Docente", value="Ing. Bernardo Antonio Hernández Batista")
    with c2:
        fecha = st.date_input("Fecha de Inicio")
        duracion = st.text_input("Tiempo Estimado", placeholder="Ej: 3 semanas")
    with c3:
        tema = st.text_input("Título de la Unidad", placeholder="Ej: La Noticia y el Reportaje")

    st.markdown('<div class="section-title">📝 3. Selección de Competencias e Indicadores</div>', unsafe_allow_html=True)
    st.markdown("*(Selecciona de las listas los elementos extraídos)*")
    
    c_f1, c_f2 = st.columns(2)
    with c_f1:
        if elementos["comp_fundamentales"]:
            sel_comp_fund = st.multiselect("Competencias Fundamentales", elementos["comp_fundamentales"])
        else:
            sel_comp_fund = st.text_area("Competencias Fundamentales (Manual)", height=50).split("\n")
            
    with c_f2:
        if elementos["comp_especificas"]:
            sel_comp_esp = st.multiselect("Competencias Específicas", elementos["comp_especificas"], format_func=recortar_opcion)
        else:
            sel_comp_esp = st.text_area("Competencias Específicas (Manual)", height=50).split("\n")

    if elementos["indicadores"]:
        sel_indicadores = st.multiselect("Indicadores de Logro", elementos["indicadores"], format_func=recortar_opcion)
    else:
        sel_indicadores = st.text_area("Indicadores de Logro (Manual)", height=50).split("\n")

    st.markdown('<div class="section-title">📚 4. Selección de Contenidos</div>', unsafe_allow_html=True)
    c_c1, c_c2, c_c3 = st.columns(3)
    with c_c1:
        if elementos["conceptuales"]:
            sel_conceptuales = st.multiselect("Conceptuales", elementos["conceptuales"])
        else:
            sel_conceptuales = st.text_area("Conceptuales (Manual)", height=50).split("\n")
    with c_c2:
        if elementos["procedimentales"]:
            sel_procedimentales = st.multiselect("Procedimentales", elementos["procedimentales"])
        else:
            sel_procedimentales = st.text_area("Procedimentales (Manual)", height=50).split("\n")
    with c_c3:
        if elementos["actitudinales"]:
            sel_actitudinales = st.multiselect("Actitudinales", elementos["actitudinales"])
        else:
            sel_actitudinales = st.text_area("Actitudinales (Manual)", height=50).split("\n")

    st.markdown('<div class="section-title">🌍 5. Contexto y Problema del Entorno</div>', unsafe_allow_html=True)
    contexto_situacion = st.text_area(
        "Describe la realidad, necesidad o problema del entorno de los estudiantes para anclar el aprendizaje:", 
        height=80, 
        placeholder="Ej: Los estudiantes de la comunidad han notado un incremento en la contaminación del río cercano. Necesitan investigar y crear una campaña..."
    )

    submit_button = st.form_submit_button("⚙️ Generar Planificación Oficial MINERD (.docx)")

# --- LÓGICA DE GENERACIÓN ---
if submit_button:
    def limpiar_lista(lst):
        return [item.strip() for item in lst if item and item.strip()]

    l_cf = limpiar_lista(sel_comp_fund)
    l_ce = limpiar_lista(sel_comp_esp)
    l_ind = limpiar_lista(sel_indicadores)
    l_con = limpiar_lista(sel_conceptuales)
    l_pro = limpiar_lista(sel_procedimentales)
    l_act = limpiar_lista(sel_actitudinales)

    if not api_key_usuario:
        st.error("🔒 Configura tu API Key en la barra lateral.")
    elif not l_cf or not l_ce or not l_ind or not contexto_situacion:
        st.warning("⚠️ Debes seleccionar Competencias, Indicadores y redactar el Contexto de la situación.")
    else:
        with st.spinner(f'🧠 Diseñando matriz curricular con {modelo_seleccionado}...'):
            try:
                prompt_maestro = f"""Actúa como Asesor Curricular Nivel Máster del MINERD.
Tu tarea es generar el cuerpo de una "Planificación por Situación de Aprendizaje" respetando de forma EXACTA los 16 renglones oficiales del formato.

INSUMOS DEL DOCENTE:
- Área: {area_sel} | Grado: {grado_sel}
- Tema: {tema}
- Contexto de la realidad: {contexto_situacion}
- Competencias Fundamentales seleccionadas: {', '.join(l_cf)}
- Competencias Específicas: {', '.join(l_ce)}
- Indicadores de Logro: {', '.join(l_ind)}
- Contenidos Conceptuales: {', '.join(l_con)}
- Contenidos Procedimentales: {', '.join(l_pro)}
- Contenidos Actitudinales: {', '.join(l_act)}

REGLAS DE DISEÑO:
1. Basado en el "Contexto de la realidad", redacta la "Situación de Aprendizaje" completa.
2. Para cada Indicador de Logro, propón 1 evidencia observable concreta.
3. Desarrolla las 4 fases de la Secuencia Didáctica (Inicio, Desarrollo, Aplicación, Cierre) detallando actividades docentes y estudiantes.
4. Propón Estrategias, Recursos, tipos de Evaluación y construye la Matriz de Evaluación.
5. NO uses saltos de línea literales (\\n) dentro de los strings del JSON.

FORMATO DE SALIDA ESTRICTO (JSON NATIVO):
{{
  "situacion_aprendizaje": {{
    "contexto": "Redacción ampliada...",
    "situacion_problematica": "Problema...",
    "desafio": "El reto...",
    "producto_final": "Producto..."
  }},
  "competencias_fundamentales_desarrollo": [
    {{"competencia": "Nombre", "como_se_desarrolla": "..."}}
  ],
  "indicadores_evidencias": [
    {{"indicador": "Texto", "evidencia": "..."}}
  ],
  "secuencia_actividades": {{
    "fase_1_inicio": {{"proposito": "...", "act_docente": ["..."], "act_estudiantes": ["..."], "recursos": "...", "evidencias": "...", "tiempo": "..."}},
    "fase_2_desarrollo": {{"proposito": "...", "act_docente": ["..."], "act_estudiantes": ["..."], "recursos": "...", "evidencias": "...", "tiempo": "..."}},
    "fase_3_aplicacion": {{"proposito": "...", "act_docente": ["..."], "act_estudiantes": ["..."], "producto_parcial": "...", "recursos": "...", "tiempo": "..."}},
    "fase_4_cierre": {{"proposito": "...", "act_docente": ["..."], "act_estudiantes": ["..."], "evidencias": "...", "tiempo": "..."}}
  }},
  "estrategias": ["Estrategia 1", "Estrategia 2"],
  "recursos_detallados": {{"humanos": ["..."], "tecnologicos": ["..."], "materiales": ["..."]}},
  "evaluacion": {{
    "diagnostica": "...", "formativa": "...", "sumativa": "...", "tecnicas": ["..."], "instrumentos": ["..."]
  }},
  "matriz_evaluacion": [
    {{"competencia_indicador": "...", "evidencia": "...", "tecnica": "...", "instrumento": "...", "criterios": "..."}}
  ],
  "evidencias_aprendizaje": ["Evidencia 1", "Evidencia 2"],
  "metacognicion": ["Pregunta 1", "Pregunta 2"],
  "atencion_diversidad": "...",
  "vinculacion": {{"familia": "...", "comunidad": "...", "aplicacion_real": "..."}}
}}
"""
                if proveedor_ia == "Google Gemini":
                    respuesta_ia = solicitar_gemini(api_key_usuario, modelo_seleccionado, prompt_maestro)
                else:
                    respuesta_ia = solicitar_openai(api_key_usuario, modelo_seleccionado, prompt_maestro)

                datos_json = parsear_json_robusto(respuesta_ia)

                datos_form = {
                    "centro": centro, "docente": docente, "fecha": fecha.strftime('%d/%m/%Y'),
                    "area": area_sel, "grado": grado_sel, "duracion": duracion, "tema": tema,
                    "comp_especificas": l_ce, "conceptuales": l_con, "procedimentales": l_pro, "actitudinales": l_act
                }

                buffer_docx = generar_word_situacion(datos_json, datos_form)
                
                st.success("✅ ¡Planificación estructurada y generada exitosamente!")
                st.download_button(
                    label="📥 Descargar Planificación Oficial MINERD (.docx)",
                    data=buffer_docx,
                    file_name=f"Planificacion_Academica_{area_sel}_{grado_sel}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )
                
            except Exception as e:
                st.error(f"⚠️ Error generando la planificación: {e}")