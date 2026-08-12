"""
alerta.py — Sistema de Alerta Temprana y Reforzamiento
Conectado a la Nube (Supabase) | UI/UX Premium | Reportes Automáticos
"""

import streamlit as st
import google.generativeai as genai
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from io import BytesIO
import json
import pandas as pd
import datetime
import re
from google.api_core.exceptions import ResourceExhausted
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type
from db import obtener_df, ejecutar_consulta  # <--- CONEXIÓN A LA NUBE

# ==========================================================================
# SECCIÓN 1 — Base de Datos (Nube)
# ==========================================================================
def init_db():
    """Inicializa la tabla de alertas en Supabase si no existe."""
    try:
        ejecutar_consulta('''
            CREATE TABLE IF NOT EXISTS alertas (
                id SERIAL PRIMARY KEY,
                fecha TEXT,
                docente TEXT,
                asignatura TEXT,
                seccion TEXT,
                competencia TEXT,
                total_estudiantes INTEGER,
                estado TEXT
            )
        ''')
    except Exception as e:
        print(f"Error inicializando tabla de alertas: {e}")

init_db()

def obtener_alertas_df():
    """Obtiene el historial de alertas desde Supabase."""
    try:
        return obtener_df('''
            SELECT id, fecha AS "Fecha", docente AS "Docente", asignatura AS "Módulo/Asignatura", 
                   seccion AS "Sección", competencia AS "Competencia/RA", 
                   total_estudiantes AS "Total Estudiantes", estado AS "Estado" 
            FROM alertas ORDER BY id DESC
        ''')
    except Exception:
        return pd.DataFrame()

# ==========================================================================
# SECCIÓN 2 — Estilos y UI Premium
# ==========================================================================
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700;800&display=swap');
    html, body, [class*="css"] { font-family: 'Inter', sans-serif; background-color: #F8FAFC; color: #1E293B; }
    
    .main-header {
        background: linear-gradient(135deg, #0F172A 0%, #1E3A8A 100%);
        padding: 2rem;
        border-radius: 12px;
        color: white;
        margin-bottom: 2rem;
        box-shadow: 0 10px 25px -5px rgba(0, 0, 0, 0.2);
        text-align: center;
    }
    .main-title { font-size: 2.2rem; font-weight: 800; margin-bottom: 0.2rem; letter-spacing: -0.02em; }
    .main-subtitle { font-size: 1.1rem; font-weight: 400; opacity: 0.85; }
    
    .section-title { color: #1D4ED8; font-weight: 700; font-size: 1.3rem; border-bottom: 2px solid #E2E8F0; padding-bottom: 8px; margin-top: 20px; margin-bottom: 18px; }
    
    div.stButton > button:first-child, div.stFormSubmitButton > button:first-child { background-color: #2563EB !important; color: #FFFFFF !important; border: none !important; border-radius: 8px !important; font-weight: 600 !important; padding: 10px 24px !important; width: 100%; margin-top: 15px;}
    div.stButton > button:first-child:hover, div.stFormSubmitButton > button:first-child:hover { background-color: #1D4ED8 !important; box-shadow: 0 4px 12px rgba(29, 78, 216, 0.2); transform: translateY(-1px); }
</style>
""", unsafe_allow_html=True)

# ==========================================================================
# SECCIÓN 3 — Funciones de JSON Robusto
# ==========================================================================
def reparar_json_truncado(texto):
    in_string = False; escape_next = False; llaves = corchetes = 0; last_safe_pos = 0
    for i, char in enumerate(texto):
        if escape_next: escape_next = False; continue
        if in_string:
            if char == "\\": escape_next = True
            elif char == '"': in_string = False; last_safe_pos = i + 1
            continue
        if char == '"': in_string = True
        elif char == "{": llaves += 1; last_safe_pos = i + 1
        elif char == "}": llaves -= 1; last_safe_pos = i + 1
        elif char == "[": corchetes += 1; last_safe_pos = i + 1
        elif char == "]": corchetes -= 1; last_safe_pos = i + 1
        elif char in (",", ":", " ", "\n", "\r", "\t"): last_safe_pos = i + 1
    reparado = texto[:last_safe_pos]
    if in_string: reparado += '"'
    reparado = reparado.rstrip().rstrip(",")
    reparado += "]" * max(corchetes, 0) + "}" * max(llaves, 0)
    return reparado

def parsear_json_robusto(respuesta):
    if not respuesta or not respuesta.strip(): raise ValueError("La IA devolvió una respuesta vacía.")
    texto = respuesta.strip().replace("```json", "").replace("```", "").strip()
    try: return json.loads(texto, strict=False)
    except: pass
    try: return json.loads(reparar_json_truncado(texto), strict=False)
    except: pass
    match = re.search(r"(\{[\s\S]*\})", texto)
    if match:
        try: return json.loads(reparar_json_truncado(match.group(1)), strict=False)
        except: pass
    raise ValueError("JSON irrecuperable tras reparación.")

# --- LLAMADAS A API ---
@retry(retry=retry_if_exception_type(ResourceExhausted), wait=wait_exponential(multiplier=2, min=4, max=20), stop=stop_after_attempt(5), reraise=True)
def solicitar_gemini_json(api_key, modelo, prompt):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(modelo)
    respuesta = model.generate_content(prompt, generation_config=genai.types.GenerationConfig(max_output_tokens=8192, temperature=0.2, response_mime_type="application/json"))
    return respuesta.text

def solicitar_openai_json(api_key, modelo, prompt):
    client = OpenAI(api_key=api_key)
    response = client.chat.completions.create(model=modelo, messages=[{"role": "user", "content": prompt}], temperature=0.2, max_tokens=8192, response_format={"type": "json_object"})
    return response.choices[0].message.content

# ==========================================================================
# SECCIÓN 4 — UI Principal
# ==========================================================================
api_key_usuario = st.session_state.get("api_key_global", "")
proveedor_ia = st.session_state.get("proveedor_ia_global", "Google Gemini")
modelo_seleccionado = st.session_state.get("modelo_global", "gemini-1.5-flash")

with st.sidebar:
    st.markdown("##### ⚡ Alerta Temprana v2.0")
    if not api_key_usuario: st.error("🔒 Configura tu API Key en Inicio")
    else: st.success(f"✅ {proveedor_ia} · {modelo_seleccionado}")

st.markdown(f"""
<div class="main-header">
    <div class="main-title">Sistema de Alerta Temprana y Reforzamiento</div>
    <div class="main-subtitle">Diagnóstico, plan de recuperación y almacenamiento permanente en la Nube</div>
</div>
""", unsafe_allow_html=True)

tab1, tab2 = st.tabs(["🚨 Generar Alerta Temprana", "📋 Historial Institucional"])

# ----------------- TAB 1: REDACTAR PLAN -----------------
with tab1:
    with st.form("form_alerta", clear_on_submit=False):
        
        with st.container(border=True):
            st.markdown('##### 🏫 1. Datos del Contexto e Identificación')
            col1, col2 = st.columns(2)
            with col1:
                docente = st.text_input("Nombre del Docente", value="Ing. Bernardo Hernández")
                asignatura = st.text_input("Módulo / Asignatura", placeholder="Ej: Sistemas Operativos / Redes LAN")
            with col2:
                politecnico = st.text_input("Centro Educativo", value="Politécnico Arquides Calderón")
                seccion = st.text_input("Sección / Grado", placeholder="Ej: 6to de Informática")
            
        with st.container(border=True):
            st.markdown('##### ⚠️ 2. Registro de Estudiantes y Brechas Detectadas')
            competencia_evaluada = st.text_input("Resultado de Aprendizaje (R.A.) o Competencia Evaluada", placeholder="Ej: R.A.1 Configurar los parámetros de red local...")
            
            estudiantes_apoyo = st.text_area(
                "Estudiantes Identificados 'En Proceso' o 'Necesitan Apoyo' (Nombres y dificultad principal)", 
                height=120, 
                placeholder="Ej:\n1. Carlos Pérez - Dificultad en el direccionamiento IP estático.\n2. María Gómez - Confusión en la identificación de topologías físicas.\n3. Juan Martínez - No puede diferenciar entre switch y router."
            )
            
        with st.container(border=True):
            st.markdown('##### 👥 3. Perfil del Grupo y Condiciones de Recuperación')
            col_p1, col_p2 = st.columns(2)
            with col_p1:
                caracteristicas_grupo = st.text_area(
                    "Características del grupo y NEAE", 
                    height=80,
                    placeholder="Ej: Grupo visual-kinestésico. 2 estudiantes con dislexia, 1 con TDAH. Buena disposición colaborativa."
                )
            with col_p2:
                tiempo_disponible = st.text_input(
                    "Tiempo disponible para recuperación", 
                    placeholder="Ej: 3 semanas / 6 sesiones de tutoría"
                )
        
        submit_button = st.form_submit_button("⚙️ Generar Plan de Recuperación Integral y Guardar en la Nube")

    # --- LÓGICA CORE ---
    if submit_button:
        if not api_key_usuario:
            st.error("🔒 Debes ingresar tu API Key en la página de Inicio (barra lateral).")
        elif not asignatura or not estudiantes_apoyo or not competencia_evaluada:
            st.warning("📝 Por favor, completa la asignatura, la competencia y el listado de estudiantes.")
        else:
            with st.spinner(f'🧠 Diseñando plan de recuperación multinivel con {modelo_seleccionado}...'):
                try:
                    prompt_maestro = f"""Actúa como un Coordinador Pedagógico de Alto Nivel, Especialista en Educación Técnico Profesional (ETP) del MINERD y Experto en Pedagogía Diferenciada e Intervención Psicopedagógica.

CONTEXTO:
- Competencia / R.A. Evaluado: {competencia_evaluada}
- Estudiantes con brechas detectadas: 
{estudiantes_apoyo}
- Características del grupo y NEAE: {caracteristicas_grupo}
- Tiempo disponible para recuperación: {tiempo_disponible}

OBJETIVO:
Diseñar un Sistema de Alerta Temprana INTEGRAL que incluya: diagnóstico con niveles de severidad, plan de recuperación multinivel con estrategias diferenciadas (A y B), recursos específicos, separación de responsabilidades (docente/estudiante/familia), indicadores de progreso intermedios, plan de contingencia y guía de comunicación familiar.

REGLAS ESTRICTAS:
1. CLASIFICACIÓN DE SEVERIDAD: Para cada estudiante, determina el nivel de severidad de su brecha:
   - CRÍTICO (🔴): Brecha fundamental que impide avanzar en el módulo. Requiere intervención inmediata.
   - MODERADO (🟡): Brecha significativa que dificulta el desempeño, pero no lo bloquea completamente. Requiere tutoría focalizada.
   - LEVE (🟢): Brecha menor o laguna de conocimiento puntual. Se resuelve con práctica guiada.
2. ESTRATEGIA DIFERENCIADA: Para cada estudiante proporciona DOS estrategias.
3. RECURSOS ESPECÍFICOS: Para cada estudiante sugiere al menos 2 recursos concretos.
4. ADAPTACIONES NEAE: Si hay NEAE, propone adaptaciones específicas.
5. SEPARACIÓN DE ROLES: Indica qué debe hacer Docente/Estudiante/Familia.
6. INDICADORES DE PROGRESO: Define 3-4 indicadores intermedios observables.
7. PLAN DE CONTINGENCIA: Acción escalate si no hay mejora.
8. COMUNICACIÓN FAMILIAR: Borrador de comunicado.

FORMATO DE SALIDA ESTRICTO (JSON NATIVO OBLIGATORIO):
NO uses saltos de línea literales (\\n) dentro de los textos. Usa texto plano.
{{
  "DIAGNOSTICO_GENERAL": "Análisis panorámico de las brechas detectadas...",
  "NIVEL_ALERTA_GLOBAL": "ROJO / AMARILLO / VERDE",
  "PLAN_ACCION_ESTUDIANTES": [
    {{
      "ESTUDIANTE": "Nombre del estudiante",
      "BRECHA_DETECTADA": "Resumen técnico de la dificultad",
      "SEVERIDAD": "CRÍTICO / MODERADO / LEVE",
      "TIEMPO_ESTIMADO": "Ej: 2 semanas",
      "ESTRATEGIA_A": "Intervención principal...",
      "ESTRATEGIA_B": "Intervención alternativa...",
      "RECURSOS": "Recursos...",
      "ADAPTACION_NEAE": "Adaptaciones...",
      "ACCION_DOCENTE": "Acción docente...",
      "ACCION_ESTUDIANTE": "Acción estudiante...",
      "ACCION_FAMILIA": "Acción familia..."
    }}
  ],
  "ACTIVIDAD_RECUPERACION_GRUPAL": {{
    "TITULO": "Título",
    "DESCRIPCION": "Descripción...",
    "PASOS": ["Paso 1...", "Paso 2..."],
    "RECURSOS": "Materiales...",
    "TIEMPO_ESTIMADO": "Duración"
  }},
  "INDICADORES_PROGRESO": [
    "Indicador 1...",
    "Indicador 2..."
  ],
  "CRITERIO_REVALUACION": "Cómo se comprobará la mejora...",
  "PLAN_CONTINGENCIA": "Acción a tomar si no hay mejora...",
  "COMUNICADO_FAMILIA": "Borrador de comunicado respetuoso..."
}}
"""
                    if proveedor_ia == "Google Gemini":
                        respuesta_ia = solicitar_gemini_json(api_key_usuario, modelo_seleccionado, prompt_maestro)
                    else:
                        respuesta_ia = solicitar_openai_json(api_key_usuario, modelo_seleccionado, prompt_maestro)

                    datos = parsear_json_robusto(respuesta_ia)

                    diagnostico = datos.get("DIAGNOSTICO_GENERAL", "")
                    nivel_alerta = datos.get("NIVEL_ALERTA_GLOBAL", "")
                    plan_estudiantes = datos.get("PLAN_ACCION_ESTUDIANTES", [])
                    actividad_recu = datos.get("ACTIVIDAD_RECUPERACION_GRUPAL", {})
                    indicadores = datos.get("INDICADORES_PROGRESO", [])
                    criterio_reval = datos.get("CRITERIO_REVALUACION", "")
                    plan_contingencia = datos.get("PLAN_CONTINGENCIA", "")
                    comunicado_familia = datos.get("COMUNICADO_FAMILIA", "")

                    # --- GUARDAR EN BASE DE DATOS SUPABASE ---
                    estado_db = "Indefinida"
                    if "ROJO" in str(nivel_alerta).upper(): estado_db = "Roja"
                    elif "AMARILLO" in str(nivel_alerta).upper(): estado_db = "Amarilla"
                    elif "VERDE" in str(nivel_alerta).upper(): estado_db = "Verde"
                    
                    try:
                        ejecutar_consulta('''
                            INSERT INTO alertas (fecha, docente, asignatura, seccion, competencia, total_estudiantes, estado)
                            VALUES (%s, %s, %s, %s, %s, %s, %s)
                        ''', (datetime.date.today().strftime('%Y-%m-%d'), docente, asignatura, seccion, competencia_evaluada, len(plan_estudiantes), estado_db))
                    except Exception as e:
                        st.error(f"Error al sincronizar con Supabase: {e}")

                    # --- CONSTRUCCIÓN DEL DOCUMENTO WORD ---
                    doc = Document()
                    doc.styles['Normal'].font.name = 'Calibri'
                    doc.styles['Normal'].font.size = Pt(11)

                    sections = doc.sections
                    for section in sections:
                        section.left_margin = Inches(0.75)
                        section.right_margin = Inches(0.75)

                    def shade_cell(cell, color):
                        shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
                        cell._tc.get_or_add_tcPr().append(shd)

                    p_encabezado = doc.add_paragraph()
                    p_encabezado.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_encabezado.add_run(f"{politecnico}\n").bold = True
                    p_encabezado.add_run("Sistema de Alerta Temprana y Plan de Recuperación Integral (ETP)\n").bold = True
                    
                    doc.add_paragraph(f"Docente: {docente} | Módulo: {asignatura} | Sección: {seccion}")
                    doc.add_paragraph(f"Competencia / R.A. Analizado: {competencia_evaluada}")
                    doc.add_paragraph(f"Tiempo Disponible: {tiempo_disponible}")
                    doc.add_paragraph("_" * 70)

                    doc.add_heading("📊 1. Diagnóstico de Alerta Temprana", level=1)
                    
                    mapa_alerta = {
                        "ROJO": ("🔴 ALERTA ROJA — Intervención URGENTE", "FECACA"),
                        "AMARILLO": ("🟡 ALERTA AMARILLA — Requiere atención focalizada", "FEF3C7"),
                        "VERDE": ("🟢 ALERTA VERDE — Reforzamiento leve", "D1FAE5"),
                    }
                    
                    texto_alerta, color_alerta = ("⚪ Sin clasificar", "F1F5F9")
                    for k, v in mapa_alerta.items():
                        if k in str(nivel_alerta).upper():
                            texto_alerta, color_alerta = v
                            break
                    
                    t_alerta = doc.add_table(rows=1, cols=1)
                    t_alerta.style = 'Table Grid'
                    cell_a = t_alerta.rows[0].cells[0]
                    shade_cell(cell_a, color_alerta)
                    p_a = cell_a.paragraphs[0]
                    p_a.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_a = p_a.add_run(texto_alerta)
                    run_a.bold = True
                    run_a.font.size = Pt(12)
                    
                    doc.add_paragraph()
                    p_diag = doc.add_paragraph(diagnostico)
                    p_diag.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    doc.add_paragraph()

                    doc.add_heading("🎯 2. Matriz de Intervención Personalizada", level=1)
                    
                    tabla_resumen = doc.add_table(rows=1, cols=5)
                    tabla_resumen.style = 'Table Grid'
                    
                    hdr_resumen = tabla_resumen.rows[0].cells
                    headers_res = ["Estudiante", "Severidad", "Brecha Detectada", "Tiempo Est.", "Estrategia Principal"]
                    for i, h_text in enumerate(headers_res):
                        p = hdr_resumen[i].paragraphs[0]
                        run = p.add_run(h_text)
                        run.bold = True
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        shade_cell(hdr_resumen[i], "E2E8F0")

                    mapa_severidad_color = {"CRÍTICO": "FEE2E2", "MODERADO": "FEF3C7", "LEVE": "D1FAE5"}
                    
                    for est in plan_estudiantes:
                        row_cells = tabla_resumen.add_row().cells
                        row_cells[0].text = str(est.get("ESTUDIANTE", ""))
                        row_cells[0].paragraphs[0].runs[0].bold = True
                        
                        sev = str(est.get("SEVERIDAD", ""))
                        row_cells[1].text = sev
                        
                        color_sev = "FFFFFF"
                        for k, v in mapa_severidad_color.items():
                            if k in sev.upper():
                                color_sev = v
                                break
                        shade_cell(row_cells[1], color_sev)
                        
                        row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        row_cells[2].text = str(est.get("BRECHA_DETECTADA", ""))
                        row_cells[3].text = str(est.get("TIEMPO_ESTIMADO", ""))
                        row_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        row_cells[4].text = str(est.get("ESTRATEGIA_A", ""))

                    doc.add_paragraph()

                    doc.add_heading("📋 3. Desglose Detallado de Intervención por Estudiante", level=1)

                    for idx, est in enumerate(plan_estudiantes):
                        nombre = str(est.get("ESTUDIANTE", f"Estudiante {idx+1}"))
                        sev = str(est.get("SEVERIDAD", ""))
                        
                        icono_sev = "⚪"
                        if "CRÍTICO" in sev.upper(): icono_sev = "🔴"
                        elif "MODERADO" in sev.upper(): icono_sev = "🟡"
                        elif "LEVE" in sev.upper(): icono_sev = "🟢"
                        
                        doc.add_heading(f"{icono_sev} {nombre} — Severidad: {sev}", level=3)
                        
                        t_det = doc.add_table(rows=7, cols=2)
                        t_det.style = 'Table Grid'
                        
                        campos = [
                            ("Brecha Detectada", str(est.get("BRECHA_DETECTADA", ""))),
                            ("Tiempo Estimado de Recuperación", str(est.get("TIEMPO_ESTIMADO", ""))),
                            ("Estrategia A (Principal)", str(est.get("ESTRATEGIA_A", ""))),
                            ("Estrategia B (Alternativa)", str(est.get("ESTRATEGIA_B", ""))),
                            ("Recursos Específicos", str(est.get("RECURSOS", ""))),
                            ("Adaptaciones NEAE", str(est.get("ADAPTACION_NEAE", ""))),
                            ("Acción del Docente", str(est.get("ACCION_DOCENTE", ""))),
                        ]
                        
                        for r, (label, valor) in enumerate(campos):
                            t_det.cell(r, 0).text = label
                            t_det.cell(r, 0).paragraphs[0].runs[0].bold = True
                            shade_cell(t_det.cell(r, 0), "F1F5F9")
                            t_det.cell(r, 1).text = valor
                            t_det.cell(r, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        
                        doc.add_paragraph()
                        p_roles = doc.add_paragraph()
                        p_roles.add_run("👤 Acción del Estudiante: ").bold = True
                        p_roles.add_run(str(est.get("ACCION_ESTUDIANTE", "")))
                        
                        p_fam = doc.add_paragraph()
                        p_fam.add_run("🏠 Acción de la Familia: ").bold = True
                        p_fam.add_run(str(est.get("ACCION_FAMILIA", "")))
                        doc.add_paragraph("_" * 50)
                        doc.add_paragraph()

                    doc.add_heading(f"⚙️ 4. Actividad de Recuperación: {actividad_recu.get('TITULO', 'Nivelación')}", level=1)
                    p_act = doc.add_paragraph(str(actividad_recu.get("DESCRIPCION", "")))
                    p_act.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

                    doc.add_heading("Pasos Metodológicos de la Tutoría:", level=3)
                    for paso in actividad_recu.get("PASOS", []):
                        p_paso = doc.add_paragraph(str(paso), style='List Bullet')
                        p_paso.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    
                    doc.add_paragraph()
                    p_recursos_act = doc.add_paragraph()
                    p_recursos_act.add_run("📦 Recursos de la Actividad: ").bold = True
                    p_recursos_act.add_run(str(actividad_recu.get("RECURSOS", "")))
                    
                    p_tiempo_act = doc.add_paragraph()
                    p_tiempo_act.add_run("⏱️ Tiempo Estimado: ").bold = True
                    p_tiempo_act.add_run(str(actividad_recu.get("TIEMPO_ESTIMADO", "")))
                    doc.add_paragraph()

                    doc.add_heading("📈 5. Indicadores de Progreso Intermedio", level=1)
                    doc.add_paragraph("Utilice estos indicadores para monitorear la mejora del estudiante ANTES de la reevaluación final. Marque con ✓ al observar evidencia del indicador.")
                    
                    if indicadores:
                        t_ind = doc.add_table(rows=len(indicadores) + 1, cols=4)
                        t_ind.style = 'Table Grid'
                        
                        hdr_ind = t_ind.rows[0].cells
                        for i, txt in enumerate(["No.", "Indicador de Progreso", "¿Evidencia?", "Fecha"]):
                            hdr_ind[i].text = txt
                            hdr_ind[i].paragraphs[0].runs[0].bold = True
                            hdr_ind[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            shade_cell(hdr_ind[i], "E2E8F0")
                        
                        for i, ind in enumerate(indicadores):
                            t_ind.cell(i+1, 0).text = str(i+1)
                            t_ind.cell(i+1, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            t_ind.cell(i+1, 1).text = str(ind)
                            t_ind.cell(i+1, 2).text = "☐ Sí  ☐ No"
                            t_ind.cell(i+1, 3).text = "___/___/____"

                    doc.add_paragraph()
                    doc.add_heading("✅ 6. Criterio de Cierre y Reevaluación", level=1)
                    p_reval = doc.add_paragraph(criterio_reval)
                    p_reval.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    doc.add_paragraph()

                    doc.add_heading("🆘 7. Plan de Contingencia", level=1)
                    p_cont = doc.add_paragraph(plan_contingencia)
                    p_cont.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    doc.add_paragraph()

                    doc.add_heading("💌 8. Guía de Comunicación con la Familia", level=1)
                    doc.add_paragraph("Borrador sugerido para comunicar la situación del estudiante a su familia. Personalice según considere oportuno:")
                    p_com = doc.add_paragraph()
                    p_com.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p_com.add_run(comunicado_familia)
                    doc.add_paragraph()

                    doc.add_heading("📝 9. Registro de Seguimiento Semanal", level=1)
                    doc.add_paragraph("Complete este registro cada semana para documentar la evolución de cada estudiante.")
                    
                    if plan_estudiantes:
                        n_semanas = 4
                        t_seg = doc.add_table(rows=len(plan_estudiantes) + 1, cols=n_semanas + 2)
                        t_seg.style = 'Table Grid'
                        
                        hdr_seg = t_seg.rows[0].cells
                        hdr_seg[0].text = "Estudiante"
                        hdr_seg[0].paragraphs[0].runs[0].bold = True
                        shade_cell(hdr_seg[0], "E2E8F0")
                        hdr_seg[1].text = "Severidad"
                        hdr_seg[1].paragraphs[0].runs[0].bold = True
                        shade_cell(hdr_seg[1], "E2E8F0")
                        for s in range(n_semanas):
                            hdr_seg[s+2].text = f"Semana {s+1}"
                            hdr_seg[s+2].paragraphs[0].runs[0].bold = True
                            hdr_seg[s+2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                            shade_cell(hdr_seg[s+2], "F8FAFC")
                        
                        for i, est in enumerate(plan_estudiantes):
                            t_seg.cell(i+1, 0).text = str(est.get("ESTUDIANTE", ""))
                            t_seg.cell(i+1, 0).paragraphs[0].runs[0].bold = True
                            sev_str = str(est.get("SEVERIDAD", ""))
                            sev_icono = "⚪"
                            if "CRÍTICO" in sev_str.upper(): sev_icono = "🔴"
                            elif "MODERADO" in sev_str.upper(): sev_icono = "🟡"
                            elif "LEVE" in sev_str.upper(): sev_icono = "🟢"
                            
                            t_seg.cell(i+1, 1).text = f"{sev_icono} {sev_str}"
                            t_seg.cell(i+1, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    doc.add_paragraph()
                    p_leyenda = doc.add_paragraph()
                    run_ley = p_leyenda.add_run("Leyenda para cada celda: L = Logrado | EP = En Proceso | NA = Necesita Apoyo | NP = No Presentó")
                    run_ley.italic = True
                    run_ley.font.size = Pt(9)

                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)
                    
                    criticos = sum(1 for e in plan_estudiantes if "CRÍTICO" in str(e.get("SEVERIDAD","")).upper())
                    moderados = sum(1 for e in plan_estudiantes if "MODERADO" in str(e.get("SEVERIDAD","")).upper())
                    leves = sum(1 for e in plan_estudiantes if "LEVE" in str(e.get("SEVERIDAD","")).upper())
                    
                    st.success("✅ ¡Plan de Recuperación Integral generado y guardado en la nube con éxito!")
                    
                    with st.container(border=True):
                        st.markdown("##### 📊 Resumen de la Alerta Generada")
                        col_r1, col_r2, col_r3, col_r4 = st.columns(4)
                        with col_r1: st.metric("🔴 Críticos", criticos)
                        with col_r2: st.metric("🟡 Moderados", moderados)
                        with col_r3: st.metric("🟢 Leves", leves)
                        with col_r4: st.metric("📋 Total", len(plan_estudiantes))
                    
                    st.download_button(
                        label="📥 Descargar Plan de Recuperación Integral (.docx)",
                        data=buffer,
                        file_name=f"Plan_Recuperacion_{asignatura[:15].replace(' ','_')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary" 
                    )
                    
                except ResourceExhausted:
                    st.error("❌ Se alcanzó el límite de API.")
                except Exception as e:
                    st.error(f"⚠️ Error de procesamiento: {e}")

# ----------------- TAB 2: HISTORIAL DE ALERTAS -----------------
with tab2:
    st.markdown('<div class="section-title">🗄️ Archivo Institucional de Alertas (Nube)</div>', unsafe_allow_html=True)
    
    df_alertas = obtener_alertas_df()
    
    if df_alertas.empty:
        st.info("Aún no hay alertas de recuperación registradas en la base de datos de Supabase.")
    else:
        # Formateo visual del estado para la tabla
        df_mostrar = df_alertas.copy()
        def format_estado(val):
            val_str = str(val).strip().title()
            if val_str == 'Roja': return '🔴 Roja'
            if val_str == 'Amarilla': return '🟡 Amarilla'
            if val_str == 'Verde': return '🟢 Verde'
            return val
            
        if 'Estado' in df_mostrar.columns:
            df_mostrar['Estado'] = df_mostrar['Estado'].apply(format_estado)
            
        st.dataframe(df_mostrar.drop(columns=['id']), use_container_width=True, hide_index=True)
        
        st.markdown("---")
        with st.expander("🗑️ Gestionar Historial (Eliminar Registro)"):
            st.warning("⚠️ Precaución: Eliminar un registro borrará la constancia de la alerta en la base de datos principal y en tu Dashboard.")
            opciones_del = {
                f"ID {row['id']}: {row['Fecha']} - {row['Docente']} ({row['Módulo/Asignatura']})": row['id']
                for _, row in df_alertas.iterrows()
            }
            seleccion_del = st.selectbox("Selecciona la alerta a eliminar", list(opciones_del.keys()), index=None, placeholder="Elige...")
            if seleccion_del is not None and st.button("Eliminar permanentemente"):
                id_a_eliminar = opciones_del[seleccion_del]
                try:
                    ejecutar_consulta("DELETE FROM alertas WHERE id = %s", (id_a_eliminar,))
                    st.success("Alerta eliminada de la base de datos institucional en la nube.")
                    st.rerun()
                except Exception as e:
                    st.error(f"Error al eliminar registro: {e}")