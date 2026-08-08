import streamlit as st
import google.generativeai as genai
from openai import OpenAI, RateLimitError as OpenAIRateLimitError
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement
from io import BytesIO
import PyPDF2
import json
import re
import unicodedata
from google.api_core.exceptions import ResourceExhausted
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type

# --- ESTILOS CSS ---
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    html, body, [class*="css"] { font-family: 'Inter', sans-serif; background-color: #F8FAFC; color: #1E293B; }
    .main-header { font-size: 2.2rem; font-weight: 700; color: #0F172A; text-align: center; margin-bottom: 5px; line-height: 1.2; }
    .sub-header { text-align: center; color: #475569; font-size: 1.1rem; font-weight: 400; margin-bottom: 35px; }
    [data-testid="stForm"] { background-color: #FFFFFF; border: 1px solid #E2E8F0; border-radius: 8px; box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05); padding: 30px; }
    .section-title { color: #1D4ED8; font-weight: 600; font-size: 1.2rem; border-bottom: 2px solid #DBEAFE; padding-bottom: 8px; margin-top: 20px; margin-bottom: 18px; }
    div.stButton > button:first-child, div.stFormSubmitButton > button:first-child { background-color: #2563EB !important; color: #FFFFFF !important; border: none !important; border-radius: 6px !important; font-weight: 600 !important; padding: 10px 24px !important; width: 100%; }
    div.stButton > button:first-child:hover, div.stFormSubmitButton > button:first-child:hover { background-color: #1D4ED8 !important; }
</style>
""", unsafe_allow_html=True)

# ===========================================================================
# UTILIDADES DE TEXTO Y JSON ROBUSTO (Blindaje contra auditorías/truncados)
# ===========================================================================
MARKER_NL = "<<NL>>"
MARKER_DQ = "<<DQ>>"
MARKER_TAB = "<<TAB>>"

def decodificar_marcadores(obj):
    if isinstance(obj, str):
        return obj.replace(MARKER_NL, "\n").replace(MARKER_DQ, '"').replace(MARKER_TAB, "\t")
    if isinstance(obj, dict):
        return {decodificar_marcadores(k): decodificar_marcadores(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [decodificar_marcadores(item) for item in obj]
    return obj

def reparar_json_truncado(texto):
    in_string = False
    escape_next = False
    llaves = corchetes = 0
    last_safe_pos = 0

    for i, char in enumerate(texto):
        if escape_next:
            escape_next = False
            continue
        if in_string:
            if char == "\\":
                escape_next = True
            elif char == '"':
                in_string = False
                last_safe_pos = i + 1
            continue
        if char == '"':
            in_string = True
        elif char == "{":
            llaves += 1
            last_safe_pos = i + 1
        elif char == "}":
            llaves -= 1
            last_safe_pos = i + 1
        elif char == "[":
            corchetes += 1
            last_safe_pos = i + 1
        elif char == "]":
            corchetes -= 1
            last_safe_pos = i + 1
        elif char in (",", ":", " ", "\n", "\r", "\t"):
            last_safe_pos = i + 1

    reparado = texto[:last_safe_pos]
    if in_string:
        reparado += '"'
    reparado = reparado.rstrip()
    if reparado.endswith(","):
        reparado = reparado[:-1]
    reparado += "]" * max(corchetes, 0)
    reparado += "}" * max(llaves, 0)
    return reparado

def parsear_json_robusto(respuesta):
    if not respuesta or not respuesta.strip():
        raise ValueError("La IA devolvió una respuesta vacía.")

    texto = respuesta.strip()
    if texto.startswith("```json"):
        texto = texto[7:]
    elif texto.startswith("```"):
        texto = texto[3:]
    if texto.endswith("```"):
        texto = texto[:-3]
    texto = texto.strip()

    try:
        return json.loads(texto, strict=False)
    except json.JSONDecodeError:
        pass

    match = re.search(r"(\{[\s\S]*\})", texto)
    if match:
        try:
            return json.loads(match.group(1), strict=False)
        except json.JSONDecodeError:
            pass

    json_start = texto.find("{")
    if json_start >= 0:
        cuerpo = texto[json_start:]
        try:
            return json.loads(reparar_json_truncado(cuerpo), strict=False)
        except json.JSONDecodeError:
            pass
        for fin in range(len(cuerpo), max(len(cuerpo) - 8000, json_start), -1):
            if fin <= json_start:
                break
            if cuerpo[fin - 1] == "}":
                try:
                    return json.loads(reparar_json_truncado(cuerpo[:fin]), strict=False)
                except json.JSONDecodeError:
                    continue

    raise ValueError(f"JSON irrecuperable. Inicio de la respuesta: {texto[:400]}...")

# --- FUNCIONES DE API CON REINTENTOS ---
@retry(retry=retry_if_exception_type(ResourceExhausted), wait=wait_exponential(multiplier=2, min=4, max=20), stop=stop_after_attempt(5), reraise=True)
def solicitar_gemini_con_reintento(api_key, modelo, prompt):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(modelo)
    respuesta = model.generate_content(
        prompt,
        generation_config=genai.types.GenerationConfig(
            max_output_tokens=8192, 
            temperature=0.1,
            response_mime_type="application/json" 
        ) 
    )
    return respuesta.text

@retry(retry=retry_if_exception_type(OpenAIRateLimitError), wait=wait_exponential(multiplier=2, min=4, max=20), stop=stop_after_attempt(5), reraise=True)
def solicitar_openai_con_reintento(api_key, modelo, prompt):
    client = OpenAI(api_key=api_key)
    response = client.chat.completions.create(
        model=modelo,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1,
        max_tokens=8192,
        response_format={"type": "json_object"} 
    )
    return response.choices[0].message.content

# --- FUNCIÓN MODULAR PARA GENERAR WORD ---
def generar_documento_word(datos_json, form_data):
    doc = Document()
    
    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)

    def shade_cell(cell, color):
        shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
        cell._tc.get_or_add_tcPr().append(shd)

    def fijar_anchos_columna(tabla, anchos_pulgadas):
        tabla.autofit = False
        for row in tabla.rows:
            for idx, ancho in enumerate(anchos_pulgadas):
                if idx < len(row.cells):
                    row.cells[idx].width = Inches(ancho)

    # [Encabezado]
    p_encabezado = doc.add_paragraph()
    p_encabezado.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    r_poli = p_encabezado.add_run(f"Nombre del Politécnico: {form_data['politecnico']}\n")
    r_poli.bold = True
    r_poli.font.size = Pt(12)
    
    if form_data['eslogan']:
        r_eslogan = p_encabezado.add_run(f"Eslogan: {form_data['eslogan']}\n")
        r_eslogan.italic = True
        r_eslogan.font.size = Pt(10)
        
    r_titulo = p_encabezado.add_run("Matriz de Planificación por Resultados de Aprendizaje")
    r_titulo.bold = True
    r_titulo.font.size = Pt(12)
    
    doc.add_paragraph() 
    
    # [Datos Generales]
    datos_gen = datos_json.get("DATOS_GENERALES", {})
    
    def add_datos_linea(label1, val1, label2, val2, label3=None, val3=None):
        p = doc.add_paragraph()
        p.add_run(label1).bold = True
        p.add_run(f" {val1} | ")
        p.add_run(label2).bold = True
        p.add_run(f" {val2}")
        if label3:
            p.add_run(" | ")
            p.add_run(label3).bold = True
            p.add_run(f" {val3}")

    add_datos_linea("Familia Profesional:", datos_gen.get("FAMILIA", "N/E"), "Denominación:", datos_gen.get("DENOMINACION", "N/E"))
    add_datos_linea("Módulo:", form_data['modulo'], "Sesión:", "Única", "Nivel:", datos_gen.get("NIVEL", "N/E"))
    add_datos_linea("Docente:", form_data['docente'], "Código:", datos_gen.get("CODIGO_MODULO", "N/E"), "Horas:", datos_gen.get("HORAS", "N/E"))
    
    fechas_split = form_data['fechas'].split("-")
    f_inicio = fechas_split[0].replace("Inicio:", "").strip() if len(fechas_split) > 0 else form_data['fechas']
    f_final = fechas_split[1].replace("Final:", "").strip() if len(fechas_split) > 1 else ""
    add_datos_linea("Fecha de Inicio:", f_inicio, "Fecha de Final:", f_final)

    p_uc = doc.add_paragraph()
    p_uc.add_run("Asociada a la Unidad de Competencia (UC): ").bold = True
    p_uc.add_run(form_data['uc_input']) 
    
    p_ra = doc.add_paragraph()
    p_ra.add_run("Resultado de Aprendizaje (RA): ").bold = True
    p_ra.add_run(form_data['ra'])
    
    doc.add_paragraph()
    
    # [Tabla Principal]
    tabla_matriz = doc.add_table(rows=1, cols=6)
    tabla_matriz.style = 'Table Grid'
    
    encabezados = ["Elementos de Capacidad", "Nivel", "Fechas", "Actividades de Enseñanza-Aprendizaje", "Instrumento de Evaluación", "Contenidos"]
    
    hdr_cells = tabla_matriz.rows[0].cells
    for i, nombre in enumerate(encabezados):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(nombre)
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(hdr_cells[i], "DBEAFE") 
        
    anchos = [1.5, 0.5, 1.0, 2.5, 1.5, 1.5]
    
    for fila in datos_json.get("TABLA_MATRIZ", []):
        row_cells = tabla_matriz.add_row().cells
        row_cells[0].text = str(fila.get("EC", ""))
        row_cells[1].text = str(fila.get("NIVEL", ""))
        row_cells[2].text = str(fila.get("FECHAS", ""))
        row_cells[3].text = str(fila.get("ACTIVIDAD", ""))
        row_cells[4].text = str(fila.get("INSTRUMENTO", ""))
        row_cells[5].text = str(fila.get("CONTENIDOS", ""))
        
    fijar_anchos_columna(tabla_matriz, anchos)

    doc.add_paragraph("\n\n")
    
    # [Firmas]
    t_firmas = doc.add_table(rows=2, cols=2)
    t_firmas.cell(0,0).text = "__________________________"
    t_firmas.cell(0,1).text = "___________________________"
    t_firmas.cell(1,0).text = "Docente"
    t_firmas.cell(1,1).text = "Coordinador Módulos Formativos ETP"
    
    for row in t_firmas.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- CONFIGURACIÓN CENTRALIZADA ---
api_key_usuario = st.session_state.get("api_key_global", "")
proveedor_ia = st.session_state.get("proveedor_ia_global", "Google Gemini")
modelo_seleccionado = st.session_state.get("modelo_global", "gemini-2.5-flash")

with st.sidebar:
    st.markdown("##### ⚡ Planificación Modular (RA)")
    if not api_key_usuario:
        st.error("🔒 Configura tu API Key en la página de Inicio")
    else:
        st.success(f"✅ {proveedor_ia} · {modelo_seleccionado}")

# --- ENCABEZADO ---
st.markdown('<div class="main-header">Matriz de Planificación por R.A.</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">Compilación Curricular ETP - Alineada a Auditoría MINERD</div>', unsafe_allow_html=True)

# --- FORMULARIO ---
with st.form("form_planificacion", clear_on_submit=False):
    
    st.markdown('<div class="section-title">📄 1. Fuente de Conocimiento Curricular</div>', unsafe_allow_html=True)
    archivo_pdf = st.file_uploader("Cargue el documento PDF oficial", type=["pdf"], help="RECOMENDACIÓN: Sube solo las páginas del módulo a trabajar.")
    
    st.markdown('<div class="section-title">🏛️ 2. Arquitectura Institucional</div>', unsafe_allow_html=True)
    col_inst1, col_inst2 = st.columns(2)
    with col_inst1:
        politecnico = st.text_input("Nombre del Politécnico", placeholder="Ingrese el nombre del centro educativo")
        docente = st.text_input("Nombre del Docente", placeholder="Ingrese su nombre completo")
    with col_inst2:
        eslogan = st.text_input("Eslogan del Politécnico", placeholder="Ingrese el eslogan institucional")
        coordinador = st.text_input("Coordinador Módulos Formativos", placeholder="Nombre del coordinador")
        
    st.markdown('<div class="section-title">📝 3. Parámetros de Operación</div>', unsafe_allow_html=True)
    
    col1, col2 = st.columns(2)
    with col1:
        modulo = st.text_input("Módulo Formativo", placeholder="Ej: MF_358_3 Impuestos al Consumo...")
        fechas = st.text_input("Fechas estimadas", placeholder="Ej: Inicio: 12/11/2026 - Final: 18/12/2026")
    with col2:
        cantidad_ec = st.number_input("Cantidad EXACTA de Elementos de Capacidad (EC)", min_value=1, value=3)
        cantidad_actividades = st.number_input("Cantidad EXACTA de Actividades a diseñar", min_value=1, value=6)
        
    uc_input = st.text_area("🔗 Unidad de Competencia (UC)", height=80, placeholder="Pega aquí la Unidad de Competencia asociada...")
    ra = st.text_area("🎯 Resultado de Aprendizaje (RA)", height=80, placeholder="Pega aquí el RA completo a planificar...")
    
    st.markdown("<br>", unsafe_allow_html=True)
    submit_button = st.form_submit_button("⚙️ Iniciar Compilación de Matriz Oficial")

# --- LÓGICA CORE ---
if submit_button:
    if not api_key_usuario:
        st.error("🔒 Debes ingresar tu API Key en la página de Inicio (barra lateral).")
    elif not archivo_pdf or not modulo or not ra or not uc_input or not politecnico or not docente:
        st.warning("📝 Parámetros Incompletos: Carga el PDF y completa todos los datos institucionales y curriculares.")
    else:
        with st.spinner(f'🧠 Ejecutando análisis curricular experto con {modelo_seleccionado}...'):
            respuesta_ia = None
            try:
                # 1. Extracción del texto del PDF
                pdf_reader = PyPDF2.PdfReader(archivo_pdf)
                texto_curriculo = "".join([pagina.extract_text() or "" for pagina in pdf_reader.pages])
                if len(texto_curriculo) > 60000: 
                    texto_curriculo = texto_curriculo[:60000]

                # 2. PROMPT MAESTRO ALINEADO A AUDITORÍA
                prompt_maestro = f"""Actúa como un Especialista Curricular de Alto Nivel de la ETP (MINERD).
He extraído el texto del Diseño Curricular oficial. Debes buscar el Módulo Formativo (MF) y el Resultado de Aprendizaje (RA), extraer los Criterios y Contenidos, y diseñar la matriz.

INSUMOS:
- Módulo Formativo: {modulo}
- Fechas estimadas: {fechas}
- Cantidad de EC a crear: {cantidad_ec}
- Cantidad de Actividades a diseñar: {cantidad_actividades}

REGLAS ESTRICTAS DE DISEÑO (PARA APROBAR AUDITORÍA):
1. EXTRACCIÓN AUTÓNOMA: Localiza en el PDF la Familia Profesional, Denominación, Nivel y Horas totales.
2. NIVELES DE DESEMPEÑO: Las actividades DEBEN estar distribuidas en los niveles 1 (Conocimiento), 2 (Aplicación) y 3 (Dominio/Autonomía).
3. INSTRUMENTOS VARIADOS: NO repitas el mismo instrumento en todas las actividades. Varía entre rúbricas, listas de cotejo, cuestionarios, informes, simulaciones, etc.
4. COHERENCIA: Los contenidos de cada fila deben estar directamente alineados con la actividad y el EC.
5. TEXTO PLANO: No utilices formato Markdown. Si necesitas salto de línea en un texto, usa {MARKER_NL}.

FORMATO DE SALIDA ESTRICTO (JSON NATIVO):
{{
  "DATOS_GENERALES": {{
    "FAMILIA": "[Extracción del PDF]",
    "DENOMINACION": "[Extracción del PDF]",
    "NIVEL": "[Extracción del PDF]",
    "CODIGO_MODULO": "[Extracción del PDF]",
    "HORAS": "[Extracción del PDF]"
  }},
  "TABLA_MATRIZ": [
    {{
      "EC": "[Redacción clara del Elemento de Capacidad]",
      "NIVEL": "[1, 2 o 3]",
      "FECHAS": "{fechas}",
      "ACTIVIDAD": "[Descripción técnica de la actividad]",
      "INSTRUMENTO": "[Instrumento variado y pertinente]",
      "CONTENIDOS": "[Contenidos específicos alineados]"
    }}
  ]
}}
El arreglo "TABLA_MATRIZ" debe contener exactamente {cantidad_actividades} objetos.

DOCUMENTO CURRICULAR A ANALIZAR:
{texto_curriculo}
"""
                # 3. Petición a la IA
                if proveedor_ia == "Google Gemini":
                    respuesta_ia = solicitar_gemini_con_reintento(api_key_usuario, modelo_seleccionado, prompt_maestro)
                else:
                    respuesta_ia = solicitar_openai_con_reintento(api_key_usuario, modelo_seleccionado, prompt_maestro)

                # 4. PARSEO JSON ROBUSTO
                datos_json = parsear_json_robusto(respuesta_ia)
                datos_json = decodificar_marcadores(datos_json)
                
                if not datos_json.get("TABLA_MATRIZ", []):
                    st.error("❌ El JSON se procesó pero la tabla de actividades está vacía.")
                    st.stop()

                # 5. PREPARAR DATOS DEL FORMULARIO
                datos_formulario = {
                    "politecnico": politecnico,
                    "eslogan": eslogan,
                    "modulo": modulo,
                    "docente": docente,
                    "fechas": fechas,
                    "uc_input": uc_input,
                    "ra": ra
                }

                # 6. GENERACIÓN DEL ARCHIVO WORD
                buffer_docx = generar_documento_word(datos_json, datos_formulario)
                
                st.success(f"✅ ¡Matriz Curricular generada con éxito! ({len(datos_json.get('TABLA_MATRIZ', []))} actividades)")
                
                st.download_button(
                    label="📥 Descargar Matriz Oficial de Planificación (.docx)",
                    data=buffer_docx,
                    file_name=f"Matriz_Planificacion_RA.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary" 
                )
                
            except ResourceExhausted:
                st.error("❌ Límite de API Gemini alcanzado. Espera unos momentos.")
            except OpenAIRateLimitError:
                st.error("❌ Límite de API OpenAI alcanzado. Espera unos momentos.")
            except ValueError as ve:
                st.error(f"⚠️ Error de procesamiento de JSON: {ve}")
                if respuesta_ia:
                    with st.expander("🔍 Ver respuesta cruda de la IA (Depuración)"):
                        st.text(respuesta_ia[:3000])
            except Exception as e:
                st.error(f"⚠️ Error Inesperado: {e}")