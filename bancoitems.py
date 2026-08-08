import streamlit as st
import google.generativeai as genai
from openai import OpenAI, RateLimitError as OpenAIRateLimitError
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from io import BytesIO
import PyPDF2
import json
import re
from google.api_core.exceptions import ResourceExhausted
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type

# --- ESTILOS CSS ---
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    html, body, [class*="css"] { font-family: 'Inter', sans-serif; background-color: #F8FAFC; color: #1E293B; }
    .main-header { font-size: 2.2rem; font-weight: 700; color: #0F172A; text-align: center; margin-bottom: 5px; line-height: 1.2; }
    .sub-header { text-align: center; color: #475569; font-size: 1.1rem; font-weight: 400; margin-bottom: 35px; }
    [data-testid="stForm"] { background-color: #FFFFFF; border: 1px solid #E2E8F0; border-radius: 8px; box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05); padding: 30px; }
    .section-title { color: #1D4ED8; font-weight: 600; font-size: 1.2rem; border-bottom: 2px solid #DBEAFE; padding-bottom: 8px; margin-top: 20px; margin-bottom: 18px; }
    .items-group { color: #475569; font-weight: 600; font-size: 0.95rem; margin-top: 14px; margin-bottom: 8px; }
    div.stButton > button:first-child, div.stFormSubmitButton > button:first-child { background-color: #2563EB !important; color: #FFFFFF !important; border: none !important; border-radius: 6px !important; font-weight: 600 !important; padding: 10px 24px !important; width: 100%; }
    div.stButton > button:first-child:hover, div.stFormSubmitButton > button:first-child:hover { background-color: #1D4ED8 !important; }
</style>
""", unsafe_allow_html=True)

# ===========================================================================
# MARCADORES SEGUROS (evitan que saltos de línea/comillas literales del
# modelo rompan el JSON — causa raíz del error "Unterminated string")
# ===========================================================================
MARKER_NL = "<<NL>>"
MARKER_DQ = "<<DQ>>"
MARKER_TAB = "<<TAB>>"

def decodificar_marcadores(obj):
    if isinstance(obj, str):
        return obj.replace(MARKER_NL, '\n').replace(MARKER_DQ, '"').replace(MARKER_TAB, '\t')
    elif isinstance(obj, dict):
        return {decodificar_marcadores(k): decodificar_marcadores(v) for k, v in obj.items()}
    elif isinstance(obj, list):
        return [decodificar_marcadores(item) for item in obj]
    return obj

# ===========================================================================
# JSON ROBUSTO — repara JSON truncado (por corte de tokens) o con pequeños
# errores de formato, en vez de fallar directo con json.loads
# ===========================================================================
def reparar_json_truncado(texto):
    in_string = False
    escape_next = False
    llaves = 0
    corchetes = 0
    last_safe_pos = 0
    for i, char in enumerate(texto):
        if escape_next:
            escape_next = False
            continue
        if in_string:
            if char == '\\':
                escape_next = True
            elif char == '"':
                in_string = False
                last_safe_pos = i + 1
            continue
        if char == '"':
            in_string = True
        elif char == '{':
            llaves += 1
            last_safe_pos = i + 1
        elif char == '}':
            llaves -= 1
            last_safe_pos = i + 1
        elif char == '[':
            corchetes += 1
            last_safe_pos = i + 1
        elif char == ']':
            corchetes -= 1
            last_safe_pos = i + 1
        elif char in (',', ':', ' ', '\n', '\r', '\t'):
            last_safe_pos = i + 1
    repair = texto[:last_safe_pos]
    if in_string:
        repair += '"'
    repair = repair.rstrip()
    if repair.endswith(','):
        repair = repair[:-1]
    repair += ']' * max(corchetes, 0)
    repair += '}' * max(llaves, 0)
    return repair

def parsear_json_robusto(respuesta):
    texto = respuesta.strip()
    if texto.startswith("```json"):
        texto = texto[7:]
    elif texto.startswith("```"):
        texto = texto[3:]
    if texto.endswith("```"):
        texto = texto[:-3]
    texto = texto.strip()
    try:
        return json.loads(texto)
    except json.JSONDecodeError:
        pass
    match = re.search(r'(\{[\s\S]*\})', texto)
    if match:
        try:
            return json.loads(match.group(1))
        except json.JSONDecodeError:
            pass
    json_start = texto.find('{')
    if json_start >= 0:
        json_body = texto[json_start:]
        try:
            return json.loads(reparar_json_truncado(json_body))
        except json.JSONDecodeError:
            pass
        for end_pos in range(len(json_body), max(len(json_body) - 8000, json_start), -1):
            if end_pos <= json_start:
                break
            if json_body[end_pos - 1] == '}':
                try:
                    return json.loads(reparar_json_truncado(json_body[:end_pos]))
                except json.JSONDecodeError:
                    continue
    try:
        return json.loads(re.sub(r'[\x00-\x1f]', '', texto))
    except json.JSONDecodeError:
        pass
    raise ValueError(f"JSON irrecuperable. Inicio: {texto[:500]}...")

# --- LLAMADAS A API CON JSON FORZADO Y LÍMITES DE TOKENS ADECUADOS ---
@retry(retry=retry_if_exception_type(ResourceExhausted), wait=wait_exponential(multiplier=2, min=4, max=20), stop=stop_after_attempt(5), reraise=True)
def solicitar_gemini_json(api_key, modelo, prompt):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(modelo)
    respuesta = model.generate_content(
        prompt,
        generation_config=genai.types.GenerationConfig(
            max_output_tokens=65536,
            temperature=0.0,
            response_mime_type="application/json"
        )
    )
    return respuesta.text

@retry(retry=retry_if_exception_type(OpenAIRateLimitError), wait=wait_exponential(multiplier=2, min=4, max=20), stop=stop_after_attempt(5), reraise=True)
def solicitar_openai_json(api_key, modelo, prompt):
    client = OpenAI(api_key=api_key)
    response = client.chat.completions.create(
        model=modelo,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.0,
        max_tokens=16384,
        response_format={"type": "json_object"}
    )
    return response.choices[0].message.content

# --- CONFIGURACIÓN CENTRALIZADA (desde main.py) ---
api_key_usuario = st.session_state.get("api_key_global", "")
proveedor_ia = st.session_state.get("proveedor_ia_global", "Google Gemini")
modelo_seleccionado = st.session_state.get("modelo_global", "gemini-2.5-flash")

with st.sidebar:
    st.markdown("##### ⚡ Banco de Ítems")
    if not api_key_usuario:
        st.error("🔒 Configura tu API Key en la página de Inicio")
    else:
        st.success(f"✅ {proveedor_ia} · {modelo_seleccionado}")

# --- ENCABEZADO ---
st.markdown('<div class="main-header">Generador de Banco de Ítems y Pruebas Teóricas</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">Evaluación escrita diversificada con anclaje estricto al contenido curricular ETP</div>', unsafe_allow_html=True)

# --- FORMULARIO ---
with st.form("form_banco_items", clear_on_submit=False):

    st.markdown('<div class="section-title">📄 1. Fuente Curricular (PDF)</div>', unsafe_allow_html=True)
    archivo_pdf = st.file_uploader("Cargue el documento PDF oficial del módulo", type=["pdf"], help="Sube las páginas específicas del contenido.")

    st.markdown('<div class="section-title">🏛️ 2. Datos Institucionales y de la Prueba</div>', unsafe_allow_html=True)
    col_inst1, col_inst2 = st.columns(2)
    with col_inst1:
        politecnico = st.text_input("Centro Educativo", value="Politécnico Salesiano Arquides Calderón")
        docente = st.text_input("Nombre del Docente", value="Ing. Bernardo Antonio Hernández Batista")
        asignatura = st.text_input("Módulo / Asignatura", placeholder="Ej: Ofimática o Redes LAN")
    with col_inst2:
        titulo_prueba = st.text_input("Título de la Evaluación", value="Prueba Teórica de Validación de Competencias")
        valor_total = st.number_input("Puntuación Total de la Prueba", min_value=10, max_value=100, value=100)
        fecha = st.date_input("Fecha de Aplicación")

    # ── 9 TIPOS DE ÍTEMS ORGANIZADOS EN 3 GRUPOS ──
    st.markdown('<div class="section-title">📝 3. Estructura de Ítems a Generar (9 Tipos)</div>', unsafe_allow_html=True)

    st.markdown('<p class="items-group">🟢 Ítems Objetivos (Selección)</p>', unsafe_allow_html=True)
    col_q1, col_q2, col_q3 = st.columns(3)
    with col_q1:
        cant_mc = st.number_input("O1. Opción Múltiple", min_value=0, max_value=20, value=4)
    with col_q2:
        cant_ci = st.number_input("O2. Correcto / Incorrecto (C/I)", min_value=0, max_value=20, value=4)
    with col_q3:
        cant_fill = st.number_input("O3. Completar Espacios", min_value=0, max_value=15, value=3)

    st.markdown('<p class="items-group">🟡 Ítems de Relación y Clasificación</p>', unsafe_allow_html=True)
    col_q4, col_q5, col_q6 = st.columns(3)
    with col_q4:
        cant_match = st.number_input("R1. Apareamiento", min_value=0, max_value=10, value=3)
    with col_q5:
        cant_order = st.number_input("R2. Ordenamiento / Secuencia", min_value=0, max_value=10, value=2)
    with col_q6:
        cant_classify = st.number_input("R3. Clasificación / Categorización", min_value=0, max_value=10, value=2)

    st.markdown('<p class="items-group">🔴 Ítems de Desarrollo y Aplicación</p>', unsafe_allow_html=True)
    col_q7, col_q8, col_q9 = st.columns(3)
    with col_q7:
        cant_short = st.number_input("D1. Respuesta Corta", min_value=0, max_value=10, value=2)
    with col_q8:
        cant_case = st.number_input("D2. Caso Práctico", min_value=0, max_value=5, value=1)
    with col_q9:
        cant_tech = st.number_input("D3. Análisis Técnico", min_value=0, max_value=10, value=2)

    st.markdown("<br>", unsafe_allow_html=True)
    submit_button = st.form_submit_button("⚙️ Generar Banco de Ítems Diversificado (Word)")

# --- LÓGICA CORE ---
if submit_button:
    if not api_key_usuario:
        st.error("🔒 Debes ingresar tu API Key en la página de Inicio (barra lateral).")
    elif not archivo_pdf or not asignatura:
        st.warning("📝 Por favor, carga el PDF curricular y define la asignatura.")
    else:
        total_items = cant_mc + cant_ci + cant_fill + cant_match + cant_order + cant_classify + cant_short + cant_case + cant_tech
        if total_items == 0:
            st.warning("📝 Selecciona al menos 1 ítem de cualquier tipo para generar la prueba.")
        else:
            with st.spinner(f'🧠 Leyendo documento y generando {total_items} ítems diversificados con {modelo_seleccionado}...'):
                respuesta_ia = None
                try:
                    # 1. Extracción del PDF
                    pdf_reader = PyPDF2.PdfReader(archivo_pdf)
                    texto_curriculo = "".join([pagina.extract_text() or "" for pagina in pdf_reader.pages])
                    if not texto_curriculo.strip():
                        raise ValueError("No se pudo extraer texto del PDF cargado.")
                    if len(texto_curriculo) > 80000:
                        texto_curriculo = texto_curriculo[:80000]

                    # 2. PROMPT MAESTRO CON 9 TIPOS DE ÍTEMS
                    prompt_maestro = f"""Actúa como un Evaluador Educativo Máster y Especialista en la Educación Técnico Profesional (ETP). 

OBJETIVO CRÍTICO:
Diseña una Prueba Teórica y Técnica diversificada con 9 tipos de ítems, basándote EXCLUSIVAMENTE Y DE FORMA LITERAL en el texto del currículo proporcionado abajo.

REGLAS ABSOLUTAS (CERO ALUCINACIONES Y CERO CONOCIMIENTO EXTERNO):
1. PROHIBIDO INVENTAR: No utilices conocimiento general ni externo. Si un concepto no está en el PDF, no lo uses.
2. EXTRACCIÓN DE CONTENIDO: Extrae conceptos, definiciones, normativas y procedimientos del documento y conviértelos en ítems.

ESPECIFICACIÓN DE CADA TIPO DE ÍTEM:

3. OPCIÓN MÚLTIPLE ({cant_mc} preguntas): Cada pregunta deriva de un párrafo real del texto. 4 opciones (A, B, C, D) y letra correcta.

4. CORRECTO / INCORRECTO ({cant_ci} ítems): Afirmaciones extraídas del documento (algunas verdaderas, otras alteradas sutilmente para que sean falsas).

5. COMPLETAR ESPACIOS ({cant_fill} ítems): Enunciados técnicos con 1 o 2 espacios en blanco marcados como "______". Las respuestas deben ser términos exactos del documento.

6. APAREAMIENTO ({cant_match} ítems): Relaciona conceptos y definiciones reales extraídas textualmente del PDF.

7. ORDENAMIENTO / SECUENCIA ({cant_order} ítems): Procesos, procedimientos o pasos técnicos del documento que deben ser ordenados lógicamente. Presenta los pasos desordenados y proporciona el orden correcto.

8. CLASIFICACIÓN / CATEGORIZACIÓN ({cant_classify} ítems): Elementos o conceptos del documento que deben clasificarse en 2 o 3 categorías extraídas del texto. Proporciona los elementos y las categorías, con la clasificación correcta.

9. RESPUESTA CORTA ({cant_short} preguntas): Preguntas que requieren una respuesta breve de 1 a 3 oraciones, basadas en definiciones o procedimientos del documento. Incluye la respuesta esperada.

10. CASO PRÁCTICO ({cant_case} casos): Un escenario técnico realista derivado del contexto del módulo, seguido de 2-3 preguntas guiadas. Incluye respuestas guía.

11. ANÁLISIS TÉCNICO ({cant_tech} preguntas): Plantea un análisis basado exclusivamente en los procesos técnicos reales del documento que requieran desarrollo extenso.

═══════════════════════════════════════════════════════════════════════════
CODIFICACIÓN OBLIGATORIA DE TEXTO — CRÍTICO PARA UN JSON VÁLIDO:
═══════════════════════════════════════════════════════════════════════════
Dentro de CUALQUIER valor de texto del JSON (preguntas, enunciados, descripciones,
respuestas, pasos, etc.), NUNCA uses saltos de línea, tabulaciones o comillas
dobles literales. En su lugar, usa estos marcadores:
- Salto de línea → {MARKER_NL}
- Comilla doble  → {MARKER_DQ}
- Tabulación     → {MARKER_TAB}
Esto aplica incluso dentro de descripciones largas de casos prácticos o
respuestas de varias oraciones: todo debe ir en una sola línea de texto plano
usando {MARKER_NL} donde normalmente pondrías un salto de línea.

FORMATO DE SALIDA ESTRICTO (JSON NATIVO OBLIGATORIO, UNA SOLA LÍNEA POR VALOR):
Devuelve un objeto JSON válido con la siguiente estructura exacta:
{{
  "MULTIPLE_CHOICE": [
    {{
      "PREGUNTA": "Texto basado estrictamente en el PDF...",
      "OPCIONES": ["A) Opción 1", "B) Opción 2", "C) Opción 3", "D) Opción 4"],
      "RESPUESTA_CORRECTA": "A"
    }}
  ],
  "CORRECT_INCORRECT": [
    {{
      "ENUNCIADO": "Afirmación técnica extraída del texto...",
      "ES_CORRECTO": true
    }}
  ],
  "FILL_IN_BLANK": [
    {{
      "ENUNCIADO": "Texto con ______ para completar y otro ______ si aplica.",
      "RESPUESTAS": ["término 1", "término 2"]
    }}
  ],
  "MATCHING": [
    {{
      "PREMISA": "Concepto real del PDF...",
      "RESPUESTA": "Definición real del PDF..."
    }}
  ],
  "ORDERING": [
    {{
      "ENUNCIADO": "Ordene los siguientes pasos según el procedimiento descrito:",
      "PASOS": ["Paso desordenado C", "Paso desordenado A", "Paso desordenado B"],
      "ORDEN_CORRECTO": ["Paso A", "Paso B", "Paso C"]
    }}
  ],
  "CLASSIFICATION": [
    {{
      "ENUNCIADO": "Clasifique los siguientes elementos según su categoría:",
      "CATEGORIAS": ["Categoría 1", "Categoría 2"],
      "ELEMENTOS": ["Elemento A", "Elemento B", "Elemento C", "Elemento D"],
      "CLASIFICACION_CORRECTA": {{
        "Categoría 1": ["Elemento A", "Elemento C"],
        "Categoría 2": ["Elemento B", "Elemento D"]
      }}
    }}
  ],
  "SHORT_ANSWER": [
    {{
      "PREGUNTA": "Pregunta que requiere respuesta breve...",
      "RESPUESTA_ESPERADA": "Respuesta esperada de 1 a 3 oraciones."
    }}
  ],
  "CASE_STUDY": [
    {{
      "TITULO": "Título del caso práctico",
      "DESCRIPCION": "Escenario técnico realista derivado del módulo...",
      "PREGUNTAS": ["Pregunta 1 del caso...", "Pregunta 2 del caso..."],
      "RESPUESTAS_GUIA": ["Respuesta guía 1...", "Respuesta guía 2..."]
    }}
  ],
  "TECHNICAL_ANALYSIS": [
    "Pregunta de análisis basada en los procesos del documento..."
  ]
}}

DOCUMENTO CURRICULAR OFICIAL A ANALIZAR (BASE ÚNICA Y OBLIGATORIA):
{texto_curriculo}
"""
                    # 3. Petición a la IA con JSON garantizado
                    if proveedor_ia == "Google Gemini":
                        respuesta_ia = solicitar_gemini_json(api_key_usuario, modelo_seleccionado, prompt_maestro)
                    else:
                        respuesta_ia = solicitar_openai_json(api_key_usuario, modelo_seleccionado, prompt_maestro)

                    # 4. Parseo robusto (con reparación de truncado) + decodificación de marcadores
                    datos = parsear_json_robusto(respuesta_ia)
                    datos = decodificar_marcadores(datos)

                    mc_list = datos.get("MULTIPLE_CHOICE", [])
                    ci_list = datos.get("CORRECT_INCORRECT", [])
                    fill_list = datos.get("FILL_IN_BLANK", [])
                    match_list = datos.get("MATCHING", [])
                    order_list = datos.get("ORDERING", [])
                    classify_list = datos.get("CLASSIFICATION", [])
                    short_list = datos.get("SHORT_ANSWER", [])
                    case_list = datos.get("CASE_STUDY", [])
                    tech_list = datos.get("TECHNICAL_ANALYSIS", [])

                    if not any([mc_list, ci_list, fill_list, match_list, order_list, classify_list, short_list, case_list, tech_list]):
                        st.error("❌ La IA no devolvió ítems. Intenta reducir la cantidad solicitada o revisa el PDF.")
                        st.stop()

                    # --- CONSTRUCCIÓN DEL DOCUMENTO WORD ---
                    doc = Document()
                    doc.styles['Normal'].font.name = 'Calibri'
                    doc.styles['Normal'].font.size = Pt(11)

                    sections = doc.sections
                    for section in sections:
                        section.left_margin = Inches(0.75)
                        section.right_margin = Inches(0.75)

                    def shade_cell(cell, color):
                        shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
                        cell._tc.get_or_add_tcPr().append(shd)

                    # Encabezado Institucional
                    p_encabezado = doc.add_paragraph()
                    p_encabezado.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_inst = p_encabezado.add_run(f"{politecnico}\n")
                    run_inst.bold = True
                    run_inst.font.size = Pt(13)

                    run_tit = p_encabezado.add_run(f"{titulo_prueba.upper()}\n")
                    run_tit.bold = True
                    run_tit.font.size = Pt(12)

                    doc.add_paragraph(f"Asignatura / Módulo: {asignatura} | Docente: {docente} | Fecha: {fecha.strftime('%d/%m/%Y')}")
                    doc.add_paragraph(f"Valor Total: {valor_total} Puntos | Calificación Obtenida: _________")
                    doc.add_paragraph("Nombre del Estudiante: _________________________________________________ | Sección: _______")
                    doc.add_paragraph("_" * 70)

                    p_inst = doc.add_paragraph()
                    p_inst.add_run("Instrucciones Generales: ").bold = True
                    p_inst.add_run("Responda cada sección basándose estrictamente en los contenidos del material de estudio oficial del módulo. Lea cuidadosamente cada tipo de ítem antes de responder.")

                    seccion_num = 0

                    # ═══════════════════════════════════════════════
                    # SECCIÓN I: OPCIÓN MÚLTIPLE
                    # ═══════════════════════════════════════════════
                    if mc_list:
                        seccion_num += 1
                        doc.add_heading(f"I. Selección Múltiple", level=2)
                        doc.add_paragraph("Instrucción: Seleccione la opción correcta para cada enunciado.")
                        for idx, item in enumerate(mc_list):
                            p_q = doc.add_paragraph()
                            p_q.add_run(f"{seccion_num}.{idx+1} ").bold = True
                            p_q.add_run(str(item.get("PREGUNTA", "")))
                            opciones = item.get("OPCIONES", [])
                            for opc in opciones:
                                p_opc = doc.add_paragraph(str(opc))
                                p_opc.paragraph_format.left_indent = Inches(0.25)
                            doc.add_paragraph()

                    # ═══════════════════════════════════════════════
                    # SECCIÓN II: CORRECTO / INCORRECTO
                    # ═══════════════════════════════════════════════
                    if ci_list:
                        seccion_num += 1
                        doc.add_heading("II. Criterio de Correcto (C) e Incorrecto (I)", level=2)
                        doc.add_paragraph("Instrucción: Escriba 'C' si la afirmación es correcta o 'I' si es incorrecta en el espacio indicado.")
                        for idx, item in enumerate(ci_list):
                            p_ci = doc.add_paragraph()
                            p_ci.add_run("_____ ").bold = True
                            p_ci.add_run(f"{seccion_num}.{idx+1}. {str(item.get('ENUNCIADO', ''))}")
                        doc.add_paragraph()

                    # ═══════════════════════════════════════════════
                    # SECCIÓN III: COMPLETAR ESPACIOS
                    # ═══════════════════════════════════════════════
                    if fill_list:
                        seccion_num += 1
                        doc.add_heading("III. Completar Espacios", level=2)
                        doc.add_paragraph("Instrucción: Complete los espacios en blanco con el término o concepto correcto según el contenido del módulo.")
                        for idx, item in enumerate(fill_list):
                            p_fill = doc.add_paragraph()
                            p_fill.add_run(f"{seccion_num}.{idx+1}. ").bold = True
                            p_fill.add_run(str(item.get("ENUNCIADO", "")))
                        doc.add_paragraph()

                    # ═══════════════════════════════════════════════
                    # SECCIÓN IV: ORDENAMIENTO / SECUENCIA
                    # ═══════════════════════════════════════════════
                    if order_list:
                        seccion_num += 1
                        doc.add_heading("IV. Ordenamiento / Secuencia", level=2)
                        doc.add_paragraph("Instrucción: Ordene lógicamente los siguientes pasos o elementos escribiendo los números 1, 2, 3... según el orden correcto del procedimiento.")
                        for idx, item in enumerate(order_list):
                            p_ord = doc.add_paragraph()
                            p_ord.add_run(f"{seccion_num}.{idx+1}. ").bold = True
                            p_ord.add_run(str(item.get("ENUNCIADO", "")))
                            pasos = item.get("PASOS", [])
                            for p_idx, paso in enumerate(pasos):
                                p_paso = doc.add_paragraph()
                                p_paso.add_run(f"_____ ").bold = True
                                p_paso.add_run(f"{chr(97 + p_idx)}) {str(paso)}")
                                p_paso.paragraph_format.left_indent = Inches(0.25)
                            doc.add_paragraph()

                    # ═══════════════════════════════════════════════
                    # SECCIÓN V: APAREAMIENTO
                    # ═══════════════════════════════════════════════
                    if match_list:
                        seccion_num += 1
                        doc.add_heading("V. Apareamiento", level=2)
                        doc.add_paragraph("Instrucción: Relacione los conceptos de la columna izquierda con la definición correspondiente de la derecha.")
                        tabla_match = doc.add_table(rows=1, cols=2)
                        tabla_match.style = 'Table Grid'
                        hdr = tabla_match.rows[0].cells
                        hdr[0].text = "Premisas / Conceptos"
                        hdr[1].text = "Términos / Definiciones"
                        hdr[0].paragraphs[0].runs[0].bold = True
                        hdr[1].paragraphs[0].runs[0].bold = True
                        shade_cell(hdr[0], "E2E8F0")
                        shade_cell(hdr[1], "E2E8F0")
                        for item in match_list:
                            row = tabla_match.add_row().cells
                            row[0].text = str(item.get("PREMISA", ""))
                            row[1].text = str(item.get("RESPUESTA", ""))
                        doc.add_paragraph()

                    # ═══════════════════════════════════════════════
                    # SECCIÓN VI: CLASIFICACIÓN / CATEGORIZACIÓN
                    # ═══════════════════════════════════════════════
                    if classify_list:
                        seccion_num += 1
                        doc.add_heading("VI. Clasificación / Categorización", level=2)
                        doc.add_paragraph("Instrucción: Clasifique cada elemento en la categoría que corresponda según el contenido del módulo.")
                        for idx, item in enumerate(classify_list):
                            p_cl = doc.add_paragraph()
                            p_cl.add_run(f"{seccion_num}.{idx+1}. ").bold = True
                            p_cl.add_run(str(item.get("ENUNCIADO", "")))

                            categorias = item.get("CATEGORIAS", [])
                            elementos = item.get("ELEMENTOS", [])

                            tabla_cl = doc.add_table(rows=1, cols=len(categorias) + 1)
                            tabla_cl.style = 'Table Grid'
                            hdr_cl = tabla_cl.rows[0].cells
                            hdr_cl[0].text = "Elemento"
                            hdr_cl[0].paragraphs[0].runs[0].bold = True
                            shade_cell(hdr_cl[0], "E2E8F0")
                            for c_idx, cat in enumerate(categorias):
                                hdr_cl[c_idx + 1].text = str(cat)
                                hdr_cl[c_idx + 1].paragraphs[0].runs[0].bold = True
                                shade_cell(hdr_cl[c_idx + 1], "E2E8F0")

                            for elem in elementos:
                                row_cl = tabla_cl.add_row().cells
                                row_cl[0].text = str(elem)
                                row_cl[0].paragraphs[0].runs[0].bold = True
                            doc.add_paragraph()

                    # ═══════════════════════════════════════════════
                    # SECCIÓN VII: RESPUESTA CORTA
                    # ═══════════════════════════════════════════════
                    if short_list:
                        seccion_num += 1
                        doc.add_heading("VII. Respuesta Corta", level=2)
                        doc.add_paragraph("Instrucción: Responda brevemente cada pregunta en 1 a 3 oraciones, basándose en el contenido del módulo.")
                        for idx, item in enumerate(short_list):
                            p_sa = doc.add_paragraph()
                            p_sa.add_run(f"{seccion_num}.{idx+1}. ").bold = True
                            p_sa.add_run(str(item.get("PREGUNTA", "")))
                            doc.add_paragraph("_" * 60)
                            doc.add_paragraph()

                    # ═══════════════════════════════════════════════
                    # SECCIÓN VIII: CASO PRÁCTICO
                    # ═══════════════════════════════════════════════
                    if case_list:
                        seccion_num += 1
                        doc.add_heading("VIII. Caso Práctico", level=2)
                        doc.add_paragraph("Instrucción: Lea detenidamente cada caso y responda las preguntas planteadas con base en los procesos técnicos del módulo.")
                        for idx, item in enumerate(case_list):
                            p_case_title = doc.add_paragraph()
                            p_case_title.add_run(f"{seccion_num}.{idx+1}. {str(item.get('TITULO', ''))}").bold = True

                            p_case_desc = doc.add_paragraph()
                            p_case_desc.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                            p_case_desc.add_run("Contexto: ").bold = True
                            p_case_desc.add_run(str(item.get("DESCRIPCION", "")))

                            preguntas_caso = item.get("PREGUNTAS", [])
                            for q_idx, preg in enumerate(preguntas_caso):
                                p_preg = doc.add_paragraph()
                                p_preg.add_run(f"  a{q_idx+1}) ").bold = True
                                p_preg.add_run(str(preg))
                                doc.add_paragraph("_" * 60)
                            doc.add_paragraph()

                    # ═══════════════════════════════════════════════
                    # SECCIÓN IX: ANÁLISIS TÉCNICO
                    # ═══════════════════════════════════════════════
                    if tech_list:
                        seccion_num += 1
                        doc.add_heading("IX. Análisis Técnico y Resolución de Casos", level=2)
                        doc.add_paragraph("Instrucción: Analice detalladamente cada planteamiento y desarrolle su respuesta fundamentada en los procesos del módulo.")
                        for idx, item in enumerate(tech_list):
                            p_t = doc.add_paragraph()
                            p_t.add_run(f"{seccion_num}.{idx+1} ").bold = True
                            p_t.add_run(str(item))
                            doc.add_paragraph("\n" + "_" * 75 + "\n" + "_" * 75 + "\n")

                    # ═══════════════════════════════════════════════
                    # SOLUCIONARIO OFICIAL (PÁGINA FINAL)
                    # ═══════════════════════════════════════════════
                    doc.add_page_break()
                    p_sol = doc.add_paragraph()
                    p_sol.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run_sol_title = p_sol.add_run("🔑 SOLUCIONARIO OFICIAL (EXCLUSIVO PARA EL DOCENTE)\n")
                    run_sol_title.bold = True
                    run_sol_title.font.size = Pt(13)
                    p_sol.add_run("Claves de corrección basadas estrictamente en el contenido cargado.\n")
                    doc.add_paragraph("_" * 70)

                    sol_seccion = 0

                    if mc_list:
                        sol_seccion += 1
                        doc.add_heading(f"Sección I: Selección Múltiple", level=3)
                        for idx, item in enumerate(mc_list):
                            p_ans = doc.add_paragraph()
                            p_ans.add_run(f"Pregunta {idx+1}: ").bold = True
                            p_ans.add_run(f"→ {item.get('RESPUESTA_CORRECTA', 'N/A')}")

                    if ci_list:
                        sol_seccion += 1
                        doc.add_heading(f"Sección II: Correcto (C) e Incorrecto (I)", level=3)
                        for idx, item in enumerate(ci_list):
                            p_ans_ci = doc.add_paragraph()
                            estado = "Correcto (C)" if item.get('ES_CORRECTO') else "Incorrecto (I)"
                            p_ans_ci.add_run(f"Afirmación {idx+1}: ").bold = True
                            p_ans_ci.add_run(f"→ {estado}")

                    if fill_list:
                        sol_seccion += 1
                        doc.add_heading(f"Sección III: Completar Espacios", level=3)
                        for idx, item in enumerate(fill_list):
                            respuestas = item.get("RESPUESTAS", [])
                            texto_resp = " / ".join([str(r) for r in respuestas])
                            p_ans_fill = doc.add_paragraph()
                            p_ans_fill.add_run(f"Ítem {idx+1}: ").bold = True
                            p_ans_fill.add_run(f"→ {texto_resp}")

                    if order_list:
                        sol_seccion += 1
                        doc.add_heading(f"Sección IV: Ordenamiento / Secuencia", level=3)
                        for idx, item in enumerate(order_list):
                            orden = item.get("ORDEN_CORRECTO", [])
                            p_ans_ord = doc.add_paragraph()
                            p_ans_ord.add_run(f"Ítem {idx+1} — Orden correcto: ").bold = True
                            for o_idx, paso in enumerate(orden):
                                p_ans_ord.add_run(f"\n  {o_idx+1}. {str(paso)}")

                    if match_list:
                        sol_seccion += 1
                        doc.add_heading(f"Sección V: Apareamiento", level=3)
                        for idx, item in enumerate(match_list):
                            p_ans2 = doc.add_paragraph()
                            p_ans2.add_run(f"Ítem {idx+1}: ").bold = True
                            p_ans2.add_run(f"{item.get('PREMISA')}  ←→  {item.get('RESPUESTA')}")

                    if classify_list:
                        sol_seccion += 1
                        doc.add_heading(f"Sección VI: Clasificación / Categorización", level=3)
                        for idx, item in enumerate(classify_list):
                            p_ans_cl = doc.add_paragraph()
                            p_ans_cl.add_run(f"Ítem {idx+1} — Clasificación correcta:").bold = True
                            clasif = item.get("CLASIFICACION_CORRECTA", {})
                            for cat, elems in clasif.items():
                                p_cat = doc.add_paragraph()
                                p_cat.paragraph_format.left_indent = Inches(0.25)
                                p_cat.add_run(f"• {str(cat)}: ").bold = True
                                if isinstance(elems, list):
                                    p_cat.add_run(", ".join([str(e) for e in elems]))
                                else:
                                    p_cat.add_run(str(elems))

                    if short_list:
                        sol_seccion += 1
                        doc.add_heading(f"Sección VII: Respuesta Corta", level=3)
                        for idx, item in enumerate(short_list):
                            p_ans_sa = doc.add_paragraph()
                            p_ans_sa.add_run(f"Pregunta {idx+1}: ").bold = True
                            p_ans_sa.add_run(f"→ {str(item.get('RESPUESTA_ESPERADA', ''))}")

                    if case_list:
                        sol_seccion += 1
                        doc.add_heading(f"Sección VIII: Caso Práctico", level=3)
                        for idx, item in enumerate(case_list):
                            p_ans_case = doc.add_paragraph()
                            p_ans_case.add_run(f"Caso {idx+1}: {str(item.get('TITULO', ''))}").bold = True
                            respuestas = item.get("RESPUESTAS_GUIA", [])
                            for r_idx, resp in enumerate(respuestas):
                                p_r = doc.add_paragraph()
                                p_r.paragraph_format.left_indent = Inches(0.25)
                                p_r.add_run(f"a{r_idx+1}) ").bold = True
                                p_r.add_run(f"→ {str(resp)}")

                    if tech_list:
                        sol_seccion += 1
                        doc.add_heading(f"Sección IX: Análisis Técnico", level=3)
                        doc.add_paragraph("Las respuestas deben evaluarse con rúbrica analítica según la profundidad del análisis, la correcta aplicación de los procesos técnicos del módulo y la fundamentación en el contenido curricular.")

                    buffer = BytesIO()
                    doc.save(buffer)
                    buffer.seek(0)

                    # Resumen de ítems generados
                    resumen = []
                    if mc_list: resumen.append(f"Opción Múltiple: {len(mc_list)}")
                    if ci_list: resumen.append(f"C/I: {len(ci_list)}")
                    if fill_list: resumen.append(f"Completar: {len(fill_list)}")
                    if match_list: resumen.append(f"Apareamiento: {len(match_list)}")
                    if order_list: resumen.append(f"Ordenamiento: {len(order_list)}")
                    if classify_list: resumen.append(f"Clasificación: {len(classify_list)}")
                    if short_list: resumen.append(f"Resp. Corta: {len(short_list)}")
                    if case_list: resumen.append(f"Caso Práctico: {len(case_list)}")
                    if tech_list: resumen.append(f"Análisis: {len(tech_list)}")

                    st.success(f"✅ ¡Prueba diversificada generada con éxito! → {' | '.join(resumen)}")

                    st.download_button(
                        label="📥 Descargar Prueba Teórica Diversificada (.docx)",
                        data=buffer,
                        file_name=f"Prueba_Diversificada_{asignatura[:10]}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        type="primary"
                    )

                except ResourceExhausted:
                    st.error("❌ Se alcanzó el límite de API. Espera unos momentos e inténtalo de nuevo.")
                except ValueError as ve:
                    st.error(f"⚠️ Error de procesamiento: {ve}")
                    if respuesta_ia:
                        with st.expander("🔍 Ver respuesta cruda de la IA (para diagnóstico)"):
                            st.text(respuesta_ia[:3000])
                    st.info("💡 Sugerencia: si esto ocurre con frecuencia, reduce la cantidad total de ítems solicitados en una sola generación.")
                except Exception as e:
                    st.error(f"⚠️ Error de procesamiento: {e}")
                    if respuesta_ia:
                        with st.expander("🔍 Ver respuesta cruda de la IA (para diagnóstico)"):
                            st.text(respuesta_ia[:3000])