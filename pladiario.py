import streamlit as st
import google.generativeai as genai
from openai import OpenAI, RateLimitError as OpenAIRateLimitError
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml, OxmlElement
from io import BytesIO
import json
import re
import unicodedata
from google.api_core.exceptions import ResourceExhausted
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type

# --- ESTILOS CSS ---
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    html, body, [class*="css"] { font-family: 'Inter', sans-serif; background-color: #F8FAFC; color: #1E293B; }
    .main-header { font-size: 2.2rem; font-weight: 700; color: #0F172A; text-align: center; margin-bottom: 5px; line-height: 1.2; }
    .sub-header { text-align: center; color: #475569; font-size: 1.1rem; font-weight: 400; margin-bottom: 35px; }
    [data-testid="stForm"] { background-color: #FFFFFF; border: 1px solid #E2E8F0; border-radius: 8px; box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05); padding: 30px; }
    .section-title { color: #1D4ED8; font-weight: 600; font-size: 1.2rem; border-bottom: 2px solid #DBEAFE; padding-bottom: 8px; margin-top: 20px; margin-bottom: 18px; }
    div.stButton > button:first-child, div.stFormSubmitButton > button:first-child { background-color: #2563EB !important; color: #FFFFFF !important; border: none !important; border-radius: 6px !important; font-weight: 600 !important; padding: 10px 24px !important; width: 100%; }
    div.stButton > button:first-child:hover, div.stFormSubmitButton > button:first-child:hover { background-color: #1D4ED8 !important; }
</style>
""", unsafe_allow_html=True)

# ===========================================================================
# UTILIDADES DE TEXTO Y JSON ROBUSTO
# ===========================================================================
MARKER_NL = "<<NL>>"
MARKER_DQ = "<<DQ>>"
MARKER_TAB = "<<TAB>>"

def decodificar_marcadores(obj):
    if isinstance(obj, str):
        return obj.replace(MARKER_NL, "\n").replace(MARKER_DQ, '"').replace(MARKER_TAB, "\t")
    if isinstance(obj, dict):
        return {decodificar_marcadores(k): decodificar_marcadores(v) for k, v in obj.items()}
    if isinstance(obj, list):
        return [decodificar_marcadores(item) for item in obj]
    return obj

def reparar_json_truncado(texto):
    in_string = False
    escape_next = False
    llaves = corchetes = 0
    last_safe_pos = 0

    for i, char in enumerate(texto):
        if escape_next:
            escape_next = False
            continue
        if in_string:
            if char == "\\":
                escape_next = True
            elif char == '"':
                in_string = False
                last_safe_pos = i + 1
            continue
        if char == '"':
            in_string = True
        elif char == "{":
            llaves += 1
            last_safe_pos = i + 1
        elif char == "}":
            llaves -= 1
            last_safe_pos = i + 1
        elif char == "[":
            corchetes += 1
            last_safe_pos = i + 1
        elif char == "]":
            corchetes -= 1
            last_safe_pos = i + 1
        elif char in (",", ":", " ", "\n", "\r", "\t"):
            last_safe_pos = i + 1

    reparado = texto[:last_safe_pos]
    if in_string:
        reparado += '"'
    reparado = reparado.rstrip()
    if reparado.endswith(","):
        reparado = reparado[:-1]
    reparado += "]" * max(corchetes, 0)
    reparado += "}" * max(llaves, 0)
    return reparado

def parsear_json_robusto(respuesta):
    if not respuesta or not respuesta.strip():
        raise ValueError("La IA devolvió una respuesta vacía.")

    texto = respuesta.strip()
    if texto.startswith("```json"):
        texto = texto[7:]
    elif texto.startswith("```"):
        texto = texto[3:]
    if texto.endswith("```"):
        texto = texto[:-3]
    texto = texto.strip()

    try:
        return json.loads(texto, strict=False)
    except json.JSONDecodeError:
        pass

    match = re.search(r"(\{[\s\S]*\})", texto)
    if match:
        try:
            return json.loads(match.group(1), strict=False)
        except json.JSONDecodeError:
            pass

    json_start = texto.find("{")
    if json_start >= 0:
        cuerpo = texto[json_start:]
        try:
            return json.loads(reparar_json_truncado(cuerpo), strict=False)
        except json.JSONDecodeError:
            pass
        for fin in range(len(cuerpo), max(len(cuerpo) - 8000, json_start), -1):
            if fin <= json_start:
                break
            if cuerpo[fin - 1] == "}":
                try:
                    return json.loads(reparar_json_truncado(cuerpo[:fin]), strict=False)
                except json.JSONDecodeError:
                    continue

    raise ValueError(f"JSON irrecuperable. Inicio de la respuesta: {texto[:400]}...")

# --- FUNCIONES DE API CON REINTENTOS ---
@retry(retry=retry_if_exception_type(ResourceExhausted), wait=wait_exponential(multiplier=2, min=4, max=20), stop=stop_after_attempt(5), reraise=True)
def solicitar_gemini_con_reintento(api_key, modelo, prompt):
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(modelo)
    respuesta = model.generate_content(
        prompt,
        generation_config=genai.types.GenerationConfig(
            max_output_tokens=8192, 
            temperature=0.2,
            response_mime_type="application/json"
        )
    )
    return respuesta.text

@retry(retry=retry_if_exception_type(OpenAIRateLimitError), wait=wait_exponential(multiplier=2, min=4, max=20), stop=stop_after_attempt(5), reraise=True)
def solicitar_openai_con_reintento(api_key, modelo, prompt):
    client = OpenAI(api_key=api_key)
    response = client.chat.completions.create(
        model=modelo,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2,
        max_tokens=8192,
        response_format={"type": "json_object"}
    )
    return response.choices[0].message.content

# ===========================================================================
# UTILIDADES WORD
# ===========================================================================
def shade_cell(cell, color):
    shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
    cell._tc.get_or_add_tcPr().append(shd)

def fijar_anchos_columna(tabla, anchos_pulgadas):
    tabla.autofit = False
    for row in tabla.rows:
        for idx, ancho in enumerate(anchos_pulgadas):
            if idx < len(row.cells):
                row.cells[idx].width = Inches(ancho)

# --- FUNCIÓN MODULAR PARA GENERAR WORD (ESQUEMA EXACTO DEL PDF) ---
def generar_documento_plandiario(datos, form_data):
    doc = Document()
    doc.styles['Normal'].font.name = 'Calibri'
    doc.styles['Normal'].font.size = Pt(10)

    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)

    # Encabezado
    p_titulo = doc.add_paragraph()
    p_titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_tit1 = p_titulo.add_run("MINISTERIO DE EDUCACIÓN DE LA REPÚBLICA DOMINICANA\n")
    run_tit1.bold = True
    run_tit1.font.size = Pt(14)
    run_tit2 = p_titulo.add_run("PLANIFICACIÓN DE CLASE DIARIA\n")
    run_tit2.bold = True
    run_tit2.font.size = Pt(12)
    run_tit3 = p_titulo.add_run("Modalidad Técnico Profesional (ETP)")
    run_tit3.italic = True
    run_tit3.font.size = Pt(11)

    def add_table_header(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.color.rgb = RGBColor(37, 99, 235)

    # --- DATOS GENERALES ---
    add_table_header("DATOS GENERALES")
    t1 = doc.add_table(rows=7, cols=4)
    t1.style = 'Table Grid'
    
    t1.cell(0,0).text = "Nombre completo"
    t1.cell(0,1).text = form_data['docente']
    t1.cell(0,2).text = "Cédula"
    t1.cell(0,3).text = form_data['cedula']
    
    t1.cell(1,0).text = "Regional"
    t1.cell(1,1).text = form_data['regional']
    t1.cell(1,2).text = "Distrito"
    t1.cell(1,3).text = form_data['distrito']
    
    t1.cell(2,0).text = "Centro Educativo"
    t1.cell(2,1).text = form_data['centro']
    t1.cell(2,2).text = "Código del Centro"
    t1.cell(2,3).text = form_data['codigo_centro']
    
    t1.cell(3,0).text = "Nivel / Subsistema"
    t1.cell(3,1).text = "Secundaria"
    t1.cell(3,2).text = "Ciclo"
    t1.cell(3,3).text = "Segundo"
    
    t1.cell(4,0).text = "Grado y Sección"
    t1.cell(4,1).text = form_data['grado']
    t1.cell(4,2).text = "Modalidad"
    t1.cell(4,3).text = "Técnico Profesional (ETP)"
    
    t1.cell(5,0).text = "Área / Asignatura/ Módulo"
    t1.cell(5,1).text = form_data['modulo']
    t1.cell(5,2).text = "Fecha"
    t1.cell(5,3).text = form_data['fecha'].strftime('%d/%m/%Y')
    
    t1.cell(6,0).text = "Duración"
    t1.cell(6,1).text = "50 minutos"
    t1.cell(6,2).text = "Tanda"
    t1.cell(6,3).text = form_data['tanda']

    for row in t1.rows:
        for c_idx, cell in enumerate(row.cells):
            if cell.paragraphs[0].runs:
                if c_idx % 2 == 0:
                    cell.paragraphs[0].runs[0].bold = True
                    shade_cell(cell, "F1F5F9")
    fijar_anchos_columna(t1, [1.5, 2.5, 1.5, 2.0])
    doc.add_paragraph()

    # --- MATRIZ DE PLANIFICACIÓN DIARIA ---
    add_table_header("Matriz de Planificación Diaria o Por Actividad - Bachillerato Técnico")
    t2 = doc.add_table(rows=7, cols=2)
    t2.style = 'Table Grid'
    
    t2.cell(0,0).text = "Características del grupo de estudiantes"
    t2.cell(0,1).text = form_data['caracteristicas']
    
    t2.cell(1,0).text = "Módulo Formativo (MF)"
    t2.cell(1,1).text = form_data['modulo']
    
    t2.cell(2,0).text = "Resultado de Aprendizaje (RA)"
    t2.cell(2,1).text = form_data['ra']
    
    t2.cell(3,0).text = "Criterio de Evaluación (CE)"
    t2.cell(3,1).text = form_data['ce']
    
    t2.cell(4,0).text = "Elemento de Capacidad (EC)"
    t2.cell(4,1).text = form_data['ec']
    
    t2.cell(5,0).text = "Tipo / Tiempo Estimado / Estrategias / Valor"
    t2.cell(5,1).text = f"Tipo: {form_data['tipo_actividad']}\nTiempo: {form_data['tiempo_estimado']}\nEstrategias: {form_data['estrategias']}\nValor: {form_data['valor']}"

    comp = datos.get("COMPONENTES", {})
    t2.cell(6,0).text = "Componentes Curriculares (Contenidos)"
    t2.cell(6,1).text = (
        f"Conceptuales:\n{comp.get('CONCEPTUALES', '')}\n\n"
        f"Procedimentales:\n{comp.get('PROCEDIMENTALES', '')}\n\n"
        f"Actitudinales:\n{comp.get('ACTITUDINALES', '')}"
    )

    for row in t2.rows:
        for c_idx, cell in enumerate(row.cells):
            if c_idx == 0 and cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].bold = True
                shade_cell(cell, "F1F5F9")
    fijar_anchos_columna(t2, [2.0, 5.5])
    doc.add_paragraph()

    # --- ENUNCIADO E INTENCIÓN ---
    add_table_header("Enunciado de la Actividad e Intención Educativa")
    t3 = doc.add_table(rows=2, cols=1)
    t3.style = 'Table Grid'
    act = datos.get("ACTIVIDAD", {})
    t3.cell(0,0).text = f"Enunciado de la Actividad:\n{act.get('ENUNCIADO', '')}"
    t3.cell(1,0).text = f"Intención Educativa:\n{act.get('INTENCION', '')}"
    for row in t3.rows:
        if row.cells[0].paragraphs[0].runs:
            row.cells[0].paragraphs[0].runs[0].bold = True
    fijar_anchos_columna(t3, [7.5])
    doc.add_paragraph()

    # --- MOMENTOS PEDAGÓGICOS ---
    add_table_header("Momentos Pedagógicos")
    t4 = doc.add_table(rows=3, cols=1)
    t4.style = 'Table Grid'
    
    ini = datos.get("INICIO", {})
    des = datos.get("DESARROLLO", {})
    cie = datos.get("CIERRE", {})

    p_inicio = t4.cell(0,0).paragraphs[0]
    p_inicio.add_run("INICIO (10 min)\n").bold = True
    p_inicio.add_run(
        f"FASE 1 — Motivación y activación de conocimientos previos:\n{ini.get('FASE1', '')}\n\n"
        f"FASE 2 — Recuperación de saberes previos:\n{ini.get('FASE2', '')}\n\n"
        f"FASE 3 — Presentación de la intención educativa:\n{ini.get('FASE3', '')}"
    )

    p_desarrollo = t4.cell(1,0).paragraphs[0]
    p_desarrollo.add_run("DESARROLLO (30 min) — Construcción del aprendizaje\n").bold = True
    p_desarrollo.add_run(
        f"FASE 1:\n{des.get('FASE1', '')}\n\n"
        f"FASE 2:\n{des.get('FASE2', '')}\n\n"
        f"FASE 3:\n{des.get('FASE3', '')}"
    )

    p_cierre = t4.cell(2,0).paragraphs[0]
    p_cierre.add_run("CIERRE (10 min) — Reflexión, consolidación y metacognición\n").bold = True
    p_cierre.add_run(
        f"Fase 1:\n{cie.get('FASE1', '')}\n\n"
        f"Fase 2:\n{cie.get('FASE2', '')}\n\n"
        f"Fase 3 Preguntas de reflexión:\n{cie.get('FASE3', '')}"
    )
    fijar_anchos_columna(t4, [7.5])
    doc.add_paragraph()

    # --- RECURSOS, NEAE Y OBSERVACIONES ---
    add_table_header("Recursos, Adaptaciones y Observaciones")
    t5 = doc.add_table(rows=4, cols=1)
    t5.style = 'Table Grid'
    
    p_rec = t5.cell(0,0).paragraphs[0]
    p_rec.add_run("Recursos:\n").bold = True
    p_rec.add_run(str(datos.get("RECURSOS", "")))
    
    p_eval = t5.cell(1,0).paragraphs[0]
    p_eval.add_run("Instrumento/s de evaluación e Indicadores:\n").bold = True
    p_eval.add_run(str(datos.get("INDICADORES_TEXT", "")))
    
    p_neae = t5.cell(2,0).paragraphs[0]
    p_neae.add_run("Adaptaciones para NEAE:\n").bold = True
    p_neae.add_run(str(datos.get("NEAE", "")))
    
    p_obs = t5.cell(3,0).paragraphs[0]
    p_obs.add_run("Observaciones:\n").bold = True
    p_obs.add_run(str(datos.get("OBSERVACIONES", "")))
    fijar_anchos_columna(t5, [7.5])
    doc.add_paragraph()

    # --- INSTRUMENTO DE EVALUACIÓN: LISTA DE COTEJO ---
    add_table_header("Instrumento de Evaluación: Lista de Cotejo")
    t6 = doc.add_table(rows=1, cols=6)
    t6.style = 'Table Grid'
    hdr = t6.rows[0].cells
    hdr[0].text = "No."
    hdr[1].text = "Criterios de Evaluación"
    hdr[2].text = "L"
    hdr[3].text = "EP"
    hdr[4].text = "NA"
    hdr[5].text = "Observaciones"
    for c_idx in range(6):
        hdr[c_idx].paragraphs[0].runs[0].bold = True
        hdr[c_idx].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        shade_cell(hdr[c_idx], "DBEAFE")

    criterios = datos.get("COTEJO", [])
    while len(criterios) < 5: criterios.append("Criterio pendiente de definir")
    
    for i, crit in enumerate(criterios[:5], start=1):
        row = t6.add_row().cells
        row[0].text = str(i)
        row[1].text = crit
        row[2].text = ""
        row[3].text = ""
        row[4].text = ""
        row[5].text = ""
        row[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        for j in range(2, 5):
            row[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    fijar_anchos_columna(t6, [0.4, 3.8, 0.4, 0.4, 0.4, 2.1])

    # --- ESCALA DE VALORACIÓN ---
    p_esc = doc.add_paragraph()
    p_esc.add_run("Escala de Valoración\n").bold = True
    t7 = doc.add_table(rows=4, cols=2)
    t7.style = 'Table Grid'
    t7.cell(0,0).text = "Sigla"
    t7.cell(0,1).text = "Descripción"
    for c in t7.rows[0].cells:
        c.paragraphs[0].runs[0].bold = True
        shade_cell(c, "F1F5F9")
        
    t7.cell(1,0).text = "L"
    t7.cell(1,1).text = "Logrado (4 o 5, Máximo 2 indicadores en EP y los demás en L)"
    t7.cell(2,0).text = "EP"
    t7.cell(2,1).text = "En proceso (3 o 4, Máximo 2 indicadores en L, pero 3 en EP)"
    t7.cell(3,0).text = "NA"
    t7.cell(3,1).text = "Necesita apoyo (3 o más indicadores en NA, con 1 o 2 EP y sin apenas ninguna L)"
    fijar_anchos_columna(t7, [0.8, 6.7])

    # --- FIRMAS ---
    doc.add_paragraph("\n\n")
    t_firmas = doc.add_table(rows=2, cols=3)
    t_firmas.cell(0,0).text = "__________________________"
    t_firmas.cell(0,1).text = "__________________________"
    t_firmas.cell(0,2).text = "__________________________"
    t_firmas.cell(1,0).text = "Director/a de Centro Educativo"
    t_firmas.cell(1,1).text = "Coordinador/a ETP"
    t_firmas.cell(1,2).text = "Docente ETP"
    for row in t_firmas.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer

# --- CONFIGURACIÓN CENTRALIZADA ---
api_key_usuario = st.session_state.get("api_key_global", "")
proveedor_ia = st.session_state.get("proveedor_ia_global", "Google Gemini")
modelo_seleccionado = st.session_state.get("modelo_global", "gemini-2.5-flash")

with st.sidebar:
    st.markdown("##### ⚡ Plan Diario ETP")
    if not api_key_usuario:
        st.error("🔒 Configura tu API Key en la página de Inicio")
    else:
        st.success(f"✅ {proveedor_ia} · {modelo_seleccionado}")

# --- ENCABEZADO ---
st.markdown('<div class="main-header">Planificación de Clase Diaria ETP</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">Estructura Oficial MINERD · 50 minutos</div>', unsafe_allow_html=True)

# --- FORMULARIO ---
with st.form("form_plandiario", clear_on_submit=False):

    st.markdown('<div class="section-title">🏫 1. Datos Generales</div>', unsafe_allow_html=True)
    col1, col2, col3 = st.columns(3)
    with col1:
        docente = st.text_input("Nombre completo", placeholder="Nombre del Docente")
        regional = st.text_input("Regional", value="06")
        centro = st.text_input("Centro Educativo", placeholder="Nombre del Centro")
        grado = st.text_input("Grado y Sección", value="5to B")
        fecha = st.date_input("Fecha")
    with col2:
        cedula = st.text_input("Cédula", placeholder="000-0000000-0")
        distrito = st.text_input("Distrito", value="06")
        codigo_centro = st.text_input("Código del Centro", placeholder="00000")
        tanda = st.text_input("Tanda", placeholder="Ej: Matutina, Vespertina, JEE")
    with col3:
        modulo = st.text_area("Área / Asignatura / Módulo", height=100, placeholder="Ej: MF 358-3 Impuestos al Consumo y Vehículos de Motor")
        tipo_actividad = st.text_input("Tipo de Actividad", value="En equipos de 4")
        tiempo_estimado = st.text_input("Tiempo Estimado", value="1 hr/clase (50 min)")
        estrategias = st.text_input("Estrategias", value="Estudio de casos · Trabajo colaborativo")
        valor = st.text_input("Valor", value="5 pts.")

    st.markdown('<div class="section-title">🎯 2. Parámetros Curriculares</div>', unsafe_allow_html=True)
    col_c1, col_c2 = st.columns(2)
    with col_c1:
        ra = st.text_area("Resultado de Aprendizaje (RA)", height=100)
        ec = st.text_area("Elemento de Capacidad (EC)", height=100)
    with col_c2:
        ce = st.text_area("Criterio de Evaluación (CE)", height=100)
        caracteristicas = st.text_area("Características del grupo de estudiantes", height=100, placeholder="Describa el perfil sociocognitivo del grupo...")

    st.markdown("<br>", unsafe_allow_html=True)
    submit_button = st.form_submit_button("⚙️ Generar Planificación Diaria")

# --- LÓGICA CORE ---
if submit_button:
    if not api_key_usuario:
        st.error("🔒 Debes ingresar tu API Key en la página de Inicio.")
    elif not docente or not modulo or not ra or not ce or not ec or not caracteristicas:
        st.warning("📝 Por favor, completa los datos básicos y curriculares.")
    else:
        with st.spinner(f'🧠 Diseñando matriz de planificación oficial en {modelo_seleccionado}...'):
            respuesta_ia = None
            try:
                prompt_maestro = f"""Actúa como experto en diseño curricular de la ETP (MINERD).
Objetivo: Diseñar el contenido pedagógico para una "Planificación de Clase Diaria" de 50 minutos para Bachillerato Técnico.

REGLAS:
1. PROHIBIDO DEJAR CAMPOS VACÍOS.
2. Inicio (10 min), Desarrollo (30 min) y Cierre (10 min). Cada momento debe tener 3 fases explícitas.
3. El cierre debe incluir preguntas de metacognición en la Fase 3.
4. Genera exactamente 5 criterios para la lista de cotejo.
5. Formato estricto JSON nativo, sin markdown.
6. Si necesitas salto de línea en un texto, usa la etiqueta: {MARKER_NL}

INSUMOS:
- Características del grupo: {caracteristicas}
- Módulo: {modulo}
- RA: {ra}
- CE: {ce}
- EC: {ec}

FORMATO JSON:
{{
  "COMPONENTES": {{
    "CONCEPTUALES": "[Redacta conceptos clave]",
    "PROCEDIMENTALES": "[Redacta procedimientos]",
    "ACTITUDINALES": "[Redacta actitudes]"
  }},
  "ACTIVIDAD": {{
    "ENUNCIADO": "[Enunciado de la actividad práctica]",
    "INTENCION": "[Intención educativa orientada al perfil]"
  }},
  "INICIO": {{
    "FASE1": "[Motivación y activación - 3 min]",
    "FASE2": "[Recuperación de saberes - 4 min]",
    "FASE3": "[Presentación de intención - 3 min]"
  }},
  "DESARROLLO": {{
    "FASE1": "[Introducción técnica - 10 min]",
    "FASE2": "[Aplicación colaborativa - 15 min]",
    "FASE3": "[Socialización - 5 min]"
  }},
  "CIERRE": {{
    "FASE1": "[Actividad de cierre interactivo - 5 min]",
    "FASE2": "[Metacognición - 3 min]",
    "FASE3": "[3 preguntas de reflexión]"
  }},
  "RECURSOS": "[Lista de recursos PDI, fichas, computadora, etc.]",
  "INDICADORES_TEXT": "[Texto breve sobre el instrumento a usar y sus indicadores]",
  "NEAE": "[Adaptaciones para NEAE]",
  "OBSERVACIONES": "[Plan alternativo u observaciones]",
  "COTEJO": [
    "[Criterio técnico 1]",
    "[Criterio técnico 2]",
    "[Criterio procedimental 3]",
    "[Criterio procedimental 4]",
    "[Criterio actitudinal 5]"
  ]
}}
"""
                if proveedor_ia == "Google Gemini":
                    respuesta_ia = solicitar_gemini_con_reintento(api_key_usuario, modelo_seleccionado, prompt_maestro)
                else:
                    respuesta_ia = solicitar_openai_con_reintento(api_key_usuario, modelo_seleccionado, prompt_maestro)

                datos = parsear_json_robusto(respuesta_ia)
                datos = decodificar_marcadores(datos)

                datos_formulario = {
                    "docente": docente, "cedula": cedula, "regional": regional, "distrito": distrito,
                    "centro": centro, "codigo_centro": codigo_centro, "grado": grado, "tanda": tanda,
                    "modulo": modulo, "tipo_actividad": tipo_actividad, "tiempo_estimado": tiempo_estimado,
                    "estrategias": estrategias, "valor": valor, "fecha": fecha,
                    "ra": ra, "ce": ce, "ec": ec, "caracteristicas": caracteristicas
                }

                buffer_docx = generar_documento_plandiario(datos, datos_formulario)

                st.success("✅ ¡Planificación diaria generada con el formato oficial!")
                st.download_button(
                    label="📥 Descargar Planificación (.docx)",
                    data=buffer_docx,
                    file_name=f"Plan_Diario_{fecha.strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary"
                )

            except ResourceExhausted:
                st.error("❌ Límite de API Gemini alcanzado. Intenta más tarde.")
            except OpenAIRateLimitError:
                st.error("❌ Límite de API OpenAI alcanzado. Intenta más tarde.")
            except ValueError as ve:
                st.error(f"⚠️ Error: {ve}")
                if respuesta_ia:
                    with st.expander("🔍 Ver respuesta cruda de la IA"):
                        st.text(respuesta_ia[:3000])
            except Exception as e:
                st.error(f"⚠️ Error inesperado: {e}")