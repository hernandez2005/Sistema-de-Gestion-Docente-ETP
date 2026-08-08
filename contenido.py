import streamlit as st
import google.generativeai as genai
from openai import OpenAI
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from io import BytesIO
import PyPDF2
import json
import re
from google.api_core.exceptions import ResourceExhausted
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type

# --- ESTILOS CSS ---
st.markdown("""
<style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@300;400;500;600;700&display=swap');
    html, body, [class*="css"] { font-family: 'Inter', sans-serif; background-color: #F8FAFC; color: #1E293B; }
    .main-header { font-size: 2.2rem; font-weight: 700; color: #0F172A; text-align: center; margin-bottom: 5px; line-height: 1.2; }
    .sub-header { text-align: center; color: #475569; font-size: 1.1rem; font-weight: 400; margin-bottom: 35px; }
    [data-testid="stForm"] { background-color: #FFFFFF; border: 1px solid #E2E8F0; border-radius: 8px; box-shadow: 0 4px 6px -1px rgba(0, 0, 0, 0.05); padding: 30px; }
    .section-title { color: #1D4ED8; font-weight: 600; font-size: 1.2rem; border-bottom: 2px solid #DBEAFE; padding-bottom: 8px; margin-top: 20px; margin-bottom: 18px; }
    div.stButton > button:first-child, div.stFormSubmitButton > button:first-child { background-color: #2563EB !important; color: #FFFFFF !important; border: none !important; border-radius: 6px !important; font-weight: 600 !important; padding: 10px 24px !important; width: 100%; }
    div.stButton > button:first-child:hover, div.stFormSubmitButton > button:first-child:hover { background-color: #1D4ED8 !important; }
</style>
""", unsafe_allow_html=True)

# ═══════════════════════════════════════════════════════════
# FUNCIÓN DE LIMPIEZA Y REPARACIÓN JSON (FORTIFICADA)
# ═══════════════════════════════════════════════════════════
def extraer_json_seguro(texto_ia, debug=False):
    """
    Extrae y parsea JSON a prueba de balas.
    Maneja: markdown fences, JSON truncado, strings con saltos de línea,
    comillas escapadas incorrectas, y JSON incompleto.
    Devuelve: (dict_datos, bool_truncado, str_limpio)
    """
    if not texto_ia or not texto_ia.strip():
        return {}, True, ""

    texto = texto_ia.strip()

    # 1. Quitar markdown fences
    if texto.startswith("```json"):
        texto = texto[7:]
    elif texto.startswith("```"):
        texto = texto[3:]
    if texto.endswith("```"):
        texto = texto[:-3]
    texto = texto.strip()

    # 2. Intentar parseo directo
    try:
        return json.loads(texto, strict=False), False, texto
    except json.JSONDecodeError:
        pass

    # 3. Extraer el bloque JSON más externo {...}
    match = re.search(r'(\{.*\})', texto, re.DOTALL)
    if match:
        bloque = match.group(1)
        try:
            return json.loads(bloque, strict=False), False, bloque
        except json.JSONDecodeError:
            pass
    else:
        bloque = texto

    # 4. Reparación de JSON truncado
    truncado = True
    reparado = bloque

    # 4a. Cerrar strings abiertos (busca comillas sin par)
    # Cuenta comillas dobles no escapadas
    comillas = len(re.findall(r'(?<!\\)"', reparado))
    if comillas % 2 != 0:
        # Hay una comilla sin cerrar — cerrar el string
        reparado += '"'

    # 4b. Cerrar corchetes y llaves abiertos
    abiertas_llaves = reparado.count('{') - reparado.count('}')
    abiertas_corchetes = reparado.count('[') - reparado.count(']')

    # Primero cerrar corchetes internos, luego llaves externas
    reparado += ']' * abiertas_corchetes
    reparado += '}' * abiertas_llaves

    # 4c. Quitar comas trailing antes de } o ]
    reparado = re.sub(r',\s*([}\]])', r'\1', reparado)

    # 4d. Quitar caracteres basura después del último }
    ultimo_cierre = reparado.rfind('}')
    if ultimo_cierre > 0 and ultimo_cierre < len(reparado) - 1:
        resto = reparado[ultimo_cierre + 1:].strip()
        if resto and not resto.startswith('}'):
            reparado = reparado[:ultimo_cierre + 1]

    # 5. Intentar parseo del JSON reparado
    try:
        datos = json.loads(reparado, strict=False)
        return datos, truncado, reparado
    except json.JSONDecodeError as e:
        pass

    # 6. Último recurso: reemplazar saltos de línea literales dentro de strings
    reparado2 = re.sub(r'\n', ' ', bloque)
    reparado2 = re.sub(r'\r', ' ', reparado2)
    comillas2 = len(re.findall(r'(?<!\\)"', reparado2))
    if comillas2 % 2 != 0:
        reparado2 += '"'
    abiertas_llaves2 = reparado2.count('{') - reparado2.count('}')
    abiertas_corchetes2 = reparado2.count('[') - reparado2.count(']')
    reparado2 += ']' * abiertas_corchetes2
    reparado2 += '}' * abiertas_llaves2
    reparado2 = re.sub(r',\s*([}\]])', r'\1', reparado2)

    try:
        datos = json.loads(reparado2, strict=False)
        return datos, truncado, reparado2
    except json.JSONDecodeError:
        pass

    # Todo falló — devolver vacío con marca de truncamiento
    return {}, True, reparado


# ═══════════════════════════════════════════════════════════
# LLAMADAS A API CON DETECCIÓN DE TRUNCAMIENTO Y REINTENTO
# ═══════════════════════════════════════════════════════════
@retry(retry=retry_if_exception_type(ResourceExhausted), wait=wait_exponential(multiplier=2, min=4, max=20), stop=stop_after_attempt(5), reraise=True)
def solicitar_gemini_json(api_key, modelo, prompt, max_tokens=16384):
    """Devuelve (texto, fue_cortado_por_tokens)."""
    genai.configure(api_key=api_key)
    model = genai.GenerativeModel(modelo)
    respuesta = model.generate_content(
        prompt,
        generation_config=genai.types.GenerationConfig(
            max_output_tokens=max_tokens,
            temperature=0.1,
            response_mime_type="application/json"
        )
    )
    corte = False
    try:
        finish_reason = respuesta.candidates[0].finish_reason
        corte = (str(finish_reason).upper().find("MAX_TOKENS") != -1) or finish_reason == 2
    except Exception:
        pass
    try:
        texto = respuesta.text
    except Exception:
        texto = ""
    return texto, corte


def solicitar_openai_json(api_key, modelo, prompt, max_tokens=16384):
    """Devuelve (texto, fue_cortado_por_tokens)."""
    client = OpenAI(api_key=api_key)
    response = client.chat.completions.create(
        model=modelo,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.1,
        max_tokens=max_tokens,
        response_format={"type": "json_object"}
    )
    corte = response.choices[0].finish_reason == "length"
    return response.choices[0].message.content, corte


def solicitar_con_reintento_json(proveedor, api_key, modelo, prompt, max_tokens=16384, tope=32000):
    """
    Solicita JSON con detección de truncamiento.
    Si se corta, reintenta UNA vez con el doble de tokens.
    Devuelve: (texto, sigue_truncado, se_reintento)
    """
    texto, corte = (solicitar_gemini_json if proveedor == "Google Gemini" else solicitar_openai_json)(
        api_key, modelo, prompt, max_tokens
    )
    reintentado = False
    if corte and max_tokens < tope:
        nuevos_tokens = min(max_tokens * 2, tope)
        texto2, corte2 = (solicitar_gemini_json if proveedor == "Google Gemini" else solicitar_openai_json)(
            api_key, modelo, prompt, nuevos_tokens
        )
        reintentado = True
        if len(texto2 or "") >= len(texto or ""):
            texto, corte = texto2, corte2
    return texto, corte, reintentado


# --- CONFIGURACIÓN CENTRALIZADA (desde main.py) ---
api_key_usuario = st.session_state.get("api_key_global", "")
proveedor_ia = st.session_state.get("proveedor_ia_global", "Google Gemini")
modelo_seleccionado = st.session_state.get("modelo_global", "gemini-2.5-flash")

with st.sidebar:
    st.markdown("##### ⚡ Contenidos v2.1")
    if not api_key_usuario:
        st.error("🔒 Configura tu API Key en la página de Inicio")
    else:
        st.success(f"✅ {proveedor_ia} · {modelo_seleccionado}")
    
    st.markdown("---")
    max_tokens_contenido = st.slider(
        "🧠 Límite de tokens", 
        8192, 32000, 16384, step=2048,
        help="Si obtienes errores de JSON, sube este valor. El sistema reintenta automáticamente con el doble si se corta."
    )
    modo_debug = st.checkbox("🐛 Modo depuración", value=False, help="Muestra la respuesta cruda de la IA para diagnosticar errores.")

# --- ENCABEZADO ---
st.markdown('<div class="main-header">Generador de Contenidos y Actividades ETP</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-header">Contenido anclado al currículo · Progresión Bloom · Diferenciación pedagógica · Rúbrica multinivel</div>', unsafe_allow_html=True)

# --- FORMULARIO ---
with st.form("form_contenido", clear_on_submit=False):
    
    st.markdown('<div class="section-title">📄 1. Fuente Curricular (Opcional — Anclaje al PDF)</div>', unsafe_allow_html=True)
    archivo_pdf = st.file_uploader(
        "Cargue el PDF del módulo para anclar el contenido (elimina alucinaciones)", 
        type=["pdf"], 
        help="Si subes el PDF, la IA generará contenido 100% fiel. Si no, usará la descripción libre."
    )
    
    st.markdown('<div class="section-title">🏫 2. Datos Institucionales y Curriculares</div>', unsafe_allow_html=True)
    col1, col2 = st.columns(2)
    with col1:
        docente = st.text_input("Nombre del Docente", value="Ing. Bernardo Antonio Hernández Batista")
        asignatura = st.text_input("Módulo / Asignatura", placeholder="Ej: Ofimática, Sistema Operativo, Redes LAN")
    with col2:
        politecnico = st.text_input("Centro Educativo", value="Politécnico Salesiano Arquides Calderón")
        fecha = st.date_input("Fecha de Aplicación")
    
    col_ra1, col_ra2 = st.columns(2)
    with col_ra1:
        ra = st.text_area("Resultado de Aprendizaje (RA)", height=68, placeholder="Pega el RA completo")
    with col_ra2:
        ce_ec = st.text_area("Criterio de Evaluación (CE) / Elemento de Capacidad (EC)", height=68, placeholder="Pega el CE y EC")

    st.markdown('<div class="section-title">📚 3. Base Pedagógica y Contenido</div>', unsafe_allow_html=True)
    contenido = st.text_area("Contenido a Desarrollar", height=90, placeholder="Ej: Configuración de subredes IP y Enrutamiento estático.")
    actividad = st.text_area("Actividad Práctica de Clase", height=90, placeholder="Ej: Los estudiantes simularán una red LAN conectando routers y switches.")

    st.markdown('<div class="section-title">👥 4. Perfil del Grupo y Condiciones</div>', unsafe_allow_html=True)
    col_p1, col_p2, col_p3 = st.columns(3)
    with col_p1:
        caracteristicas_grupo = st.text_area("Características del grupo / NEAE", height=68, placeholder="Ej: Grupo visual-kinestésico. 2 con dislexia.")
    with col_p2:
        duracion_sesion = st.text_input("Duración de la sesión", value="50 minutos")
    with col_p3:
        nivel_bloom_objetivo = st.selectbox("Nivel Bloom objetivo", [
            "Recordar / Comprender",
            "Comprender / Aplicar",
            "Aplicar / Analizar",
            "Analizar / Evaluar",
            "Evaluar / Crear"
        ], index=2)

    st.markdown('<div class="section-title">📋 5. Estrategia de Evaluación</div>', unsafe_allow_html=True)
    col_eval1, col_eval2 = st.columns([3, 1])
    with col_eval1:
        instrumento = st.selectbox("Técnica / Instrumento de Evaluación", [
            "Rúbrica Analítica de Competencias ETP (con niveles L / EP / NA)",
            "Lista de Cotejo Avanzada (Indicadores de Logro)",
            "Escala Estimativa con Descriptores de Desempeño",
            "Guía de Observación Metodológica",
            "Registro de Desempeño Técnico"
        ])
    with col_eval2:
        valor_puntos = st.number_input("Valor (Puntos)", min_value=1, max_value=100, value=100)
        cant_criterios = st.number_input("Criterios", min_value=3, max_value=8, value=5)

    st.markdown("<br>", unsafe_allow_html=True)
    submit_button = st.form_submit_button("⚙️ Generar Material Didáctico Integral (Word)")

# --- LÓGICA CORE ---
if submit_button:
    if not api_key_usuario:
        st.error("🔒 Debes ingresar tu API Key en la página de Inicio (barra lateral).")
    elif not asignatura or not contenido:
        st.warning("📝 Por favor, completa la asignatura y el contenido a desarrollar.")
    else:
        with st.spinner(f'🧠 Generando contenido integral con {modelo_seleccionado} (tokens: {max_tokens_contenido})...'):
            try:
                # --- Extracción de PDF ---
                texto_curriculo = ""
                if archivo_pdf:
                    pdf_reader = PyPDF2.PdfReader(archivo_pdf)
                    texto_curriculo = "".join([pagina.extract_text() for pagina in pdf_reader.pages])
                    if len(texto_curriculo) > 80000:
                        texto_curriculo = texto_curriculo[:80000]

                contexto_curricular = ""
                if texto_curriculo:
                    contexto_curricular = f"""
DOCUMENTO CURRICULAR OFICIAL CARGADO (BASE ÚNICA Y OBLIGATORIA):
{texto_curriculo}

REGLA DE ANCLAJE: Todo el contenido teórico, glosario y actividades deben derivarse EXCLUSIVAMENTE de este documento. PROHIBIDO utilizar conocimiento externo no presente en el PDF.
"""
                else:
                    contexto_curricular = """
REGLA DE CONTENIDO: Genera contenido basado en la descripción proporcionada, con rigor técnico y estándares reales de la industria.
"""

                prompt_maestro = f"""Actúa como un Catedrático Universitario de Alto Nivel, Especialista Curricular ETP del MINERD y Experto en Diseño Instruccional con Taxonomía de Bloom.

INSUMOS:
- Contenido a desarrollar: {contenido}
- Actividad propuesta: {actividad}
- Instrumento seleccionado: {instrumento}
- Valor total: {valor_puntos} puntos
- Cantidad EXACTA de criterios de evaluación: {cant_criterios}
- Duración de la sesión: {duracion_sesion}
- Nivel de Bloom objetivo: {nivel_bloom_objetivo}
- Características del grupo: {caracteristicas_grupo}
- Resultado de Aprendizaje (RA): {ra}
- Criterio de Evaluación / EC: {ce_ec}

{contexto_curricular}

REGLAS DE GENER,ACIÓN:

1. RESUMEN EJECUTIVO: Síntesis de 3 líneas para el docente.

2. CONOCIMIENTOS PREVIOS: 3-4 prerrequisitos.

3. ERRORES COMUNES: 3-4 errores típicos con corrección.

4. CONTENIDO TEÓRICO (PROGRESIÓN BLOOM): Mínimo 3 secciones con nivel Bloom, ayuda visual sugerida.

5. GLOSARIO: Mínimo 5 términos con definición + ejemplo de aplicación.

6. CONEXIONES INTERCURRICULARES: 2-3 vínculos.

7. ACTIVIDAD PRINCIPAL: Pasos con tiempo, modalidad (Individual/Parejas/Grupal) y recurso. Incluir indicador de logro.

8. ACTIVIDAD DE REFUERZO: Para estudiantes que no alcancen el logro. 3-4 pasos.

9. ACTIVIDAD DE EXTENSIÓN: Para estudiantes que superen rápidamente.

10. ADAPTACIONES NEAE: Si aplica, o "Sin adaptaciones adicionales requeridas".

11. TICKET DE SALIDA: 3 preguntas de verificación rápida.

12. TAREA INDEPENDIENTE: Con entregable.

13. SIMULADORES: 3 con URL real y nivel Bloom.

14. CRITERIOS DE EVALUACIÓN: Exactamente {cant_criterios} criterios con descriptores L/EP/NA. Distribuir {valor_puntos} puntos.

15. AUTOEVALUACIÓN: 3-4 ítems de reflexión.

16. GUÍA DE RESPUESTAS: Resultado esperado por paso de la actividad principal.

REGLA CRÍTICA DE JSON: NO uses saltos de línea literales (\\n) dentro de los valores de texto. Une las oraciones con espacios. NO uses comillas dobles dentro de los textos (usa comillas simples si necesitas citar).

FORMATO DE SALIDA ESTRICTO (JSON NATIVO OBLIGATORIO):
{{
  "RESUMEN_EJECUTIVO": "...",
  "CONOCIMIENTOS_PREVIOS": ["...", "..."],
  "ERRORES_COMUNES": [{{"ERROR": "...", "CORRECCION": "..."}}],
  "CONTENIDO_TEORICO": [{{"NIVEL_BLOOM": "...", "TITULO_SECCION": "...", "CONTENIDO": "...", "AYUDA_VISUAL": "..."}}],
  "GLOSARIO": [{{"TERMINO": "...", "DEFINICION": "...", "EJEMPLO_APLICACION": "..."}}],
  "CONEXIONES_INTERCURRICULARES": ["...", "..."],
  "ACTIVIDAD_PRINCIPAL": {{"TITULO": "...", "INDICADOR_LOGRO": "...", "PASOS": [{{"PASO": "...", "TIEMPO": "...", "MODALIDAD": "...", "RECURSO": "..."}}]}},
  "ACTIVIDAD_REFUERZO": {{"TITULO": "...", "PASOS": ["...", "..."]}},
  "ACTIVIDAD_EXTENSION": {{"TITULO": "...", "DESCRIPCION": "...", "ENTREGABLE": "..."}},
  "ADAPTACIONES_NEAE": "...",
  "TICKET_SALIDA": ["...", "...", "..."],
  "TAREA_INDEPENDIENTE": {{"DESCRIPCION": "...", "ENTREGABLE": "..."}},
  "SIMULADORES_RECURSOS": [{{"TIPO": "...", "NOMBRE": "...", "DESCRIPCION": "...", "URL": "...", "NIVEL_BLOOM": "..."}}],
  "WEBGRAFIA": ["...", "..."],
  "CRITERIOS_EVALUACION": [{{"CRITERIO": "...", "INDICADOR": "...", "LOGRADO": "...", "EN_PROCESO": "...", "NECESITA_APOYO": "...", "PUNTOS": 20}}],
  "AUTOEVALUACION": ["...", "..."],
  "GUIA_RESPUESTAS": [{{"PASO&": "...", "RESPUESTA_ESPERADA": "..."}}]
}}
"""
                # ── Petición a la IA con reintento automático ──
                respuesta_ia, corte_tokens, se_reintento = solicitar_con_reintento_json(
                    proveedor_ia, api_key_usuario, modelo_seleccionado, prompt_maestro, max_tokens_contenido
                )

                # ── Depuración ──
                if modo_debug:
                    with st.expander("🐛 Respuesta cruda de la IA (depuración)"):
                        st.text_area("Raw response:", respuesta_ia or "(vacío)", height=300, key="debug_raw")
                        st.write({"Corte por tokens": corte_tokens, "Reintento": se_reintento})

                # ── Parseo JSON con reparación automática ──
                datos, truncado, texto_limpio = extraer_json_seguro(respuesta_ia, debug=modo_debug)

                if not datos:
                    st.error("❌ La IA no devolvió un JSON válido ni reparable.")
                    if modo_debug:
                        with st.expander("🐛 Texto limpio intentado"):
                            st.text_area("Limpio:", texto_limpio, height=300, key="debug_limpio")
                    st.info("💡 **Solución:** Sube el 'Límite de tokens' en la barra lateral (ej. a 24576 o 32000) y vuelve a intentar.")
                    st.stop()

                if truncado:
                    st.warning("⚠️ La respuesta se truncó pero se reparó parcialmente. Algunos campos pueden estar incompletos. Si falta información, sube el límite de tokens y reintent.")

                if se_reintento and not truncado:
                    st.info("ℹ️ La primera llamada se cortó, pero el reintento automático con más tokens tuvo éxito.")

                # ═══════════════════════════════════════════════
                # EXTRACCIÓN DE DATOS
                # ═══════════════════════════════════════════════
                resumen = datos.get("RESUMEN_EJECUTIVO", "")
                conocimientos_previos = datos.get("CONOCIMIENTOS_PREVIOS", [])
                errores_comunes = datos.get("ERRORES_COMUNES", [])
                contenido_teorico = datos.get("CONTENIDO_TEORICO", [])
                glosario = datos.get("GLOSARIO", [])
                conexiones = datos.get("CONEXIONES_INTERCURRICULARES", [])
                actividad_principal = datos.get("ACTIVIDAD_PRINCIPAL", {})
                actividad_refuerzo = datos.get("ACTIVIDAD_REFUERZO", {})
                actividad_extension = datos.get("ACTIVIDAD_EXTENSION", {})
                adapt_neae = datos.get("ADAPTACIONES_NEAE", "")
                ticket_salida = datos.get("TICKET_SALIDA", [])
                tarea_independiente = datos.get("TAREA_INDEPENDIENTE", {})
                simuladores = datos.get("SIMULADORES_RECURSOS", [])
                webgrafia = datos.get("WEBGRAFIA", [])
                criterios = datos.get("CRITERIOS_EVALUACION", [])
                autoevaluacion = datos.get("AUTOEVALUACION", [])
                guia_respuestas = datos.get("GUIA_RESPUESTAS", [])

                # ═══════════════════════════════════════════════
                # CONSTRUCCIÓN DEL DOCUMENTO WORD
                # ═══════════════════════════════════════════════
                doc = Document()
                doc.styles['Normal'].font.name = 'Calibri'
                doc.styles['Normal'].font.size = Pt(11)

                sections = doc.sections
                for section in sections:
                    section.left_margin = Inches(0.75)
                    section.right_margin = Inches(0.75)

                def shade_cell(cell, color):
                    shd = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), color))
                    cell._tc.get_or_add_tcPr().append(shd)

                # ── I. ENCABEZADO ──
                p_enc = doc.add_paragraph()
                p_enc.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_enc.add_run(f"{politecnico}\n").bold = True
                p_enc.add_run("Material Didáctico Integral — Generador de Contenidos ETP v2.1\n").bold = True
                doc.add_paragraph(f"Docente: {docente} | Módulo: {asignatura} | Fecha: {fecha.strftime('%d/%m/%Y')}")
                doc.add_paragraph(f"Duración: {duracion_sesion} | Nivel Bloom: {nivel_bloom_objetivo}")
                doc.add_paragraph("Estudiante: _________________________________________________ | Sección: _______")
                doc.add_paragraph("_" * 70)

                # ── II. RESUMEN EJECUTIVO ──
                if resumen:
                    doc.add_heading("II. Resumen Ejecutivo (Para el Docente)", level=1)
                    p_res = doc.add_paragraph(resumen)
                    p_res.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    p_res.runs[0].italic = True
                    doc.add_paragraph()

                # ── III. ALINEACIÓN CURRICULAR ──
                if ra or ce_ec:
                    doc.add_heading("III. Alineación Curricular", level=1)
                    if ra:
                        p_ra_doc = doc.add_paragraph()
                        p_ra_doc.add_run("Resultado de Aprendizaje (RA): ").bold = True
                        p_ra_doc.add_run(str(ra))
                    if ce_ec:
                        p_ce_doc = doc.add_paragraph()
                        p_ce_doc.add_run("Criterio de Evaluación / EC: ").bold = True
                        p_ce_doc.add_run(str(ce_ec))
                    doc.add_paragraph()

                # ── IV. CONOCIMIENTOS PREVIOS ──
                if conocimientos_previos:
                    doc.add_heading("IV. Conocimientos Previos Requeridos", level=1)
                    doc.add_paragraph("Verifique que los estudiantes dominan estos prerrequisitos antes de iniciar:")
                    for cp in conocimientos_previos:
                        p_cp = doc.add_paragraph(str(cp), style='List Bullet')
                        p_cp.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                    doc.add_paragraph()

                # ── V. ERRORES COMUNES ──
                if errores_comunes:
                    doc.add_heading("V. Errores Comunes y Concepciones Previas", level=1)
                    t_err = doc.add_table(rows=1, cols=2)
                    t_err.style = 'Table Grid'
                    hdr_e = t_err.rows[0].cells
                    hdr_e[0].text = "⚠️ Error / Concepción Errónea"
                    hdr_e[1].text = "✅ Corrección Pedagógica"
                    for i in range(2):
                        hdr_e[i].paragraphs[0].runs[0].bold = True
                        shade_cell(hdr_e[i], "FEE2E2")
                    for err in errores_comunes:
                        row_e = t_err.add_row().cells
                        row_e[0].text = str(err.get("ERROR", ""))
                        row_e[1].text = str(err.get("CORRECCION", ""))
                    doc.add_paragraph()

                # ── VI. CONTENIDO TEÓRICO ──
                if contenido_teorico:
                    doc.add_heading("VI. Desarrollo de Contenido Teórico", level=1)
                    for sec in contenido_teorico:
                        bloom = str(sec.get("NIVEL_BLOOM", ""))
                        titulo_sec = str(sec.get("TITULO_SECCION", ""))
                        contenido_sec = str(sec.get("CONTENIDO", ""))
                        ayuda_vis = str(sec.get("AYUDA_VISUAL", ""))
                        p_sec = doc.add_paragraph()
                        p_sec.add_run(f"[Bloom: {bloom}] ").bold = True
                        p_sec.add_run(f"{titulo_sec}").bold = True
                        p_cont = doc.add_paragraph(contenido_sec)
                        p_cont.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                        if ayuda_vis:
                            p_vis = doc.add_paragraph()
                            p_vis.add_run("💡 Ayuda visual sugerida: ").bold = True
                            p_vis.add_run(str(ayuda_vis))
                            p_vis.runs[-1].italic = True
                        doc.add_paragraph()

                # ── VII. GLOSARIO ──
                if glosario:
                    doc.add_heading("VII. Glosario Técnico", level=1)
                    t_glos = doc.add_table(rows=1, cols=3)
                    t_glos.style = 'Table Grid'
                    hdr_g = t_glos.rows[0].cells
                    for i, txt in enumerate(["Término", "Definición", "Ejemplo de Aplicación"]):
                        hdr_g[i].text = txt
                        hdr_g[i].paragraphs[0].runs[0].bold = True
                        shade_cell(hdr_g[i], "E2E8F0")
                    for item in glosario:
                        row_g = t_glos.add_row().cells
                        row_g[0].text = str(item.get("TERMINO", ""))
                        row_g[0].paragraphs[0].runs[0].bold = True
                        row_g[1].text = str(item.get("DEFINICION", ""))
                        row_g[2].text = str(item.get("EJEMPLO_APLICACION", ""))
                    doc.add_paragraph()

                # ── VIII. CONEXIONES INTERCURRICULARES ──
                if conexiones:
                    doc.add_heading("VIII. Conexiones Intercurriculares", level=1)
                    for conn in conexiones:
                        doc.add_paragraph(str(conn), style='List Bullet')
                    doc.add_paragraph()

                # ── IX. ACTIVIDAD PRINCIPAL ──
                doc.add_heading(f"IX. Actividad Principal: {actividad_principal.get('TITULO', '')}", level=1)
                if actividad_principal.get("INDICADOR_LOGRO"):
                    p_ind = doc.add_paragraph()
                    p_ind.add_run("Indicador de Logro: ").bold = True
                    p_ind.add_run(str(actividad_principal.get("INDICADOR_LOGRO", "")))
                    doc.add_paragraph()
                pasos_ppal = actividad_principal.get("PASOS", [])
                if pasos_ppal:
                    t_act = doc.add_table(rows=1, cols=4)
                    t_act.style = 'Table Grid'
                    hdr_a = t_act.rows[0].cells
                    for i, txt in enumerate(["Paso", "Tiempo", "Modalidad", "Recurso"]):
                        hdr_a[i].text = txt
                        hdr_a[i].paragraphs[0].runs[0].bold = True
                        hdr_a[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        shade_cell(hdr_a[i], "DBEAFE")
                    for p_item in pasos_ppal:
                        row_a = t_act.add_row().cells
                        row_a[0].text = str(p_item.get("PASO", ""))
                        row_a[1].text = str(p_item.get("TIEMPO", ""))
                        row_a[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        row_a[2].text = str(p_item.get("MODALIDAD", ""))
                        row_a[2].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        row_a[3].text = str(p_item.get("RECURSO", ""))
                doc.add_paragraph()

                # ── X. ACTIVIDAD DE REFUERZO ──
                if actividad_refuerzo:
                    doc.add_heading(f"X. Actividad de Refuerzo: {actividad_refuerzo.get('TITULO', '')}", level=1)
                    doc.add_paragraph("Para estudiantes que no alcancen el indicador de logro:")
                    for paso in actividad_refuerzo.get("PASOS", []):
                        doc.add_paragraph(str(paso), style='List Bullet')
                    doc.add_paragraph()

                # ── XI. ACTIVIDAD DE EXTENSIÓN ──
                if actividad_extension:
                    doc.add_heading(f"XI. Actividad de Extensión: {actividad_extension.get('TITULO', '')}", level=1)
                    doc.add_paragraph("Para estudiantes que superen rápidamente el indicador:")
                    p_ext = doc.add_paragraph()
                    p_ext.add_run("Descripción: ").bold = True
                    p_ext.add_run(str(actividad_extension.get("DESCRIPCION", "")))
                    p_ent = doc.add_paragraph()
                    p_ent.add_run("Entregable: ").bold = True
                    p_ent.add_run(str(actividad_extension.get("ENTREGABLE", "")))
                    doc.add_paragraph()

                # ── XII. ADAPTACIONES NEAE ──
                doc.add_heading("XII. Adaptaciones para NEAE", level=1)
                doc.add_paragraph(str(adapt_neae))
                doc.add_paragraph()

                # ── XIII. TICKET DE SALIDA ──
                if ticket_salida:
                    doc.add_heading("XIII. Ticket de Salida", level=1)
                    doc.add_paragraph("Aplique en los últimos 3-5 minutos:")
                    for idx, preg in enumerate(ticket_salida):
                        p_tk = doc.add_paragraph()
                        p_tk.add_run(f"{idx+1}. ").bold = True
                        p_tk.add_run(str(preg))
                        doc.add_paragraph("R: _______________________________________________")
                    doc.add_paragraph()

                # ── XIV. TAREA INDEPENDIENTE ──
                if tarea_independiente:
                    doc.add_heading("XIV. Tarea Independiente", level=1)
                    p_tarea = doc.add_paragraph()
                    p_tarea.add_run("Descripción: ").bold = True
                    p_tarea.add_run(str(tarea_independiente.get("DESCRIPCION", "")))
                    p_entrega = doc.add_paragraph()
                    p_entrega.add_run("Entregable: ").bold = True
                    p_entrega.add_run(str(tarea_independiente.get("ENTREGABLE", "")))
                    doc.add_paragraph()

                # ── XV. SIMULADORES ──
                if simuladores:
                    doc.add_heading("XV. Repositorio de Simuladores y Recursos", level=1)
                    for sim in simuladores:
                        p_sim = doc.add_paragraph(style='List Bullet')
                        p_sim.add_run(f"[{sim.get('TIPO', '')}] ").bold = True
                        p_sim.add_run(f"{sim.get('NOMBRE', '')}: ").bold = True
                        p_sim.add_run(str(sim.get('DESCRIPCION', '')) + " ")
                        p_sim.add_run(f"[Bloom: {sim.get('NIVEL_BLOOM', '')}] ").italic = True
                        url_val = sim.get('URL', '')
                        if url_val:
                            p_sim.add_run(f"🔗 {url_val}").italic = True
                    doc.add_paragraph()

                # ── XVI. WEBGRAFÍA ──
                if webgrafia:
                    doc.add_heading("XVI. Fuentes y Referencias", level=1)
                    for ref in webgrafia:
                        doc.add_paragraph(str(ref), style='List Bullet')
                    doc.add_paragraph()

                # ═══════════════════════════════════════════════
                # SECCIÓN DOCENTE (PÁGINA SEPARADA)
                # ═══════════════════════════════════════════════
                doc.add_page_break()
                p_sol_title = doc.add_paragraph()
                p_sol_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run_st = p_sol_title.add_run("🔑 SECCIÓN PARA EL DOCENTE\n")
                run_st.bold = True
                run_st.font.size = Pt(14)
                p_sol_title.add_run("Rúbrica Multinivel · Guía de Respuestas · Autoevaluación")
                doc.add_paragraph("_" * 70)

                # ── RÚBRICA L/EP/NA ──
                if criterios:
                    doc.add_heading(f"Rúbrica: {instrumento.split('(')[0].strip()}", level=1)
                    col_rub = ["No.", "Criterio e Indicador", "L (Logrado)", "EP (En Proceso)", "NA (Necesita Apoyo)", "Pts"]
                    t_rub = doc.add_table(rows=1, cols=len(col_rub))
                    t_rub.style = 'Table Grid'
                    hdr_r = t_rub.rows[0].cells
                    for i, txt in enumerate(col_rub):
                        hdr_r[i].text = txt
                        hdr_r[i].paragraphs[0].runs[0].bold = True
                        hdr_r[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        shade_cell(hdr_r[i], "E2E8F0")
                    for idx, crit in enumerate(criterios):
                        row_r = t_rub.add_row().cells
                        row_r[0].text = str(idx + 1)
                        row_r[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        row_r[1].text = f"• {str(crit.get('CRITERIO', ''))}\n• Ind: {str(crit.get('INDICADOR', ''))}"
                        row_r[2].text = str(crit.get("LOGRADO", ""))
                        shade_cell(row_r[2], "D1FAE5")
                        row_r[3].text = str(crit.get("EN_PROCESO", ""))
                        shade_cell(row_r[3], "FEF3C7")
                        row_r[4].text = str(crit.get("NECESITA_APOYO", ""))
                        shade_cell(row_r[4], "FEE2E2")
                        row_r[5].text = f"{crit.get('PUNTOS', 0)}"
                        row_r[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    row_tot = t_rub.add_row().cells
                    row_tot[1].text = "TOTAL"
                    row_tot[1].paragraphs[0].runs[0].bold = True
                    row_tot[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    row_tot[5].text = str(valor_puntos)
                    row_tot[5].paragraphs[0].runs[0].bold = True
                    row_tot[5].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()

                # ── AUTOEVALUACIÓN ──
                if autoevaluacion:
                    doc.add_heading("Autoevaluación del Estudiante", level=2)
                    t_auto = doc.add_table(rows=len(autoevaluacion) + 1, cols=4)
                    t_auto.style = 'Table Grid'
                    hdr_au = t_auto.rows[0].cells
                    for i, txt in enumerate(["Reflexión", "Sí", "Parcialmente", "No"]):
                        hdr_au[i].text = txt
                        hdr_au[i].paragraphs[0].runs[0].bold = True
                        if i > 0: hdr_au[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        shade_cell(hdr_au[i], "F1F5F9")
                    for idx, item in enumerate(autoevaluacion):
                        t_auto.cell(idx+1, 0).text = str(item)
                        for c in range(1, 4):
                            t_auto.cell(idx+1, c).text = "☐"
                            t_auto.cell(idx+1, c).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()

                # ── GUÍA DE RESPUESTAS ──
                if guia_respuestas:
                    doc.add_heading("Guía de Respuestas (Solucionario)", level=2)
                    t_resp = doc.add_table(rows=1, cols=2)
                    t_resp.style = 'Table Grid'
                    hdr_rp = t_resp.rows[0].cells
                    hdr_rp[0].text = "Paso"
                    hdr_rp[1].text = "Respuesta / Resultado Esperado"
                    for i in range(2):
                        hdr_rp[i].paragraphs[0].runs[0].bold = True
                        shade_cell(hdr_rp[i], "DBEAFE")
                    for resp in guia_respuestas:
                        row_rp = t_resp.add_row().cells
                        row_rp[0].text = str(resp.get("PASO", ""))
                        row_rp[0].paragraphs[0].runs[0].bold = True
                        row_rp[1].text = str(resp.get("RESPUESTA_ESPERADA", ""))

                # ── Guardar ──
                buffer = BytesIO()
                doc.save(buffer)
                buffer.seek(0)

                # ── Métricas ──
                col_m1, col_m2, col_m3, col_m4 = st.columns(4)
                with col_m1: st.metric("📄 Secciones", sum(1 for x in [resumen, conocimientos_previos, errores_comunes, contenido_teorico, glosario, conexiones, actividad_principal, actividad_refuerzo, actividad_extension, ticket_salida, tarea_independiente, simuladores, webgrafia, criterios] if x))
                with col_m2: st.metric("📚 Secciones Bloom", len(contenido_teorico))
                with col_m3: st.metric("🎯 Criterios", len(criterios))
                with col_m4: st.metric("🔗 Simuladores", len(simuladores))

                if archivo_pdf: st.info("📌 Contenido anclado al PDF curricular.")
                else: st.warning("⚠️ Sin PDF — contenido desde descripción libre.")

                st.success("✅ ¡Material Didáctico Integral v2.1 generado con éxito!")

                st.download_button(
                    label="📥 Descargar Material Didáctico Integral (.docx)",
                    data=buffer,
                    file_name=f"Material_Integral_v2_{asignatura[:10]}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    type="primary" 
                )
                
            except ResourceExhausted:
                st.error("❌ Se alcanzó el límite de API.")
            except Exception as e:
                st.error(f"⚠️ Error de procesamiento: {e}")